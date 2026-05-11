"""
docx_generator_tcc.py  — Generador DOCX con formato PJF oficial

Convierte el markdown del estudio de fondo (salida del pipeline v3)
en un documento DOCX con formato de sentencias del PJF:

  • Texto principal:  Arial 14, interlineado 1.5, justificado
  • Títulos:          Arial 14, negrita, centrado
  • Citas (rubros):   Arial 12, cursiva, interlineado 1.0, bloque indentado
  • Artículos:        Arial 12, cursiva, interlineado 1.0, indentado
  • Notas al pie:     Arial 10, referencia APA automática

Uso:
    from docx_generator_tcc import generate_docx_bytes
    docx_bytes = generate_docx_bytes(markdown_text, meta={...})
"""

from __future__ import annotations
import re
import io
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════════
# Constants
# ═══════════════════════════════════════════════════════════════════════════════

FONT_NAME = "Arial"
FONT_SIZE_MAIN = Pt(14)
FONT_SIZE_CITE = Pt(12)
FONT_SIZE_FOOTNOTE = Pt(10)
LINE_SPACING_MAIN = 1.5
LINE_SPACING_CITE = 1.0

MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(2)

# Regex to detect thesis/jurisprudence citation blocks
# Matches patterns like "Registro digital: 12345" or "Registro: 2012345"
RE_REGISTRO = re.compile(
    r"(?:Registro\s*(?:digital)?\s*:?\s*#?\s*)(\d{4,8})",
    re.IGNORECASE,
)

# Detect thesis rubro (typically ALL CAPS lines after a quote block)
RE_RUBRO_LINE = re.compile(r"^[A-ZÁÉÍÓÚÑÜ\s,\.\-:;¿¡\(\)]{20,}$")

# Detect article references
RE_ARTICULO = re.compile(
    r"(?:artículo|art\.)\s+\d+",
    re.IGNORECASE,
)


# ═══════════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _set_paragraph_format(paragraph, line_spacing: float, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Apply line spacing and alignment to a paragraph."""
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.alignment = alignment
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)


def _add_run(paragraph, text: str, bold=False, italic=False,
             font_size=FONT_SIZE_MAIN, font_name=FONT_NAME):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.italic = italic
    # Ensure Arial is used (set for East Asian and Complex Script too)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    return run


def _add_footnote(paragraph, footnote_text: str, doc: Document):
    """
    Add a footnote reference to a paragraph.
    Uses a simplified approach: adds superscript number inline and
    appends footnote text at the bottom of the document.
    """
    # We track footnotes as a list on the document object
    if not hasattr(doc, '_tcc_footnotes'):
        doc._tcc_footnotes = []
    doc._tcc_footnotes.append(footnote_text)
    fn_num = len(doc._tcc_footnotes)

    # Add superscript number
    run = paragraph.add_run(str(fn_num))
    run.font.superscript = True
    run.font.size = Pt(10)
    run.font.name = FONT_NAME

    return fn_num


def _set_indent(paragraph, left_cm: float = 1.5):
    """Set left indent for citation blocks."""
    paragraph.paragraph_format.left_indent = Cm(left_cm)


def _build_apa_citation(registro: str, rubro: str = "", instancia: str = "",
                        epoca: str = "") -> str:
    """Build APA-style footnote text for a thesis."""
    parts = []
    if rubro:
        # Clean and truncate rubro for footnote
        clean = rubro.strip().rstrip(".")
        if len(clean) > 120:
            clean = clean[:117] + "..."
        parts.append(f'"{clean}"')
    if registro:
        parts.append(f"Registro digital: {registro}")
    if instancia:
        parts.append(instancia)
    if epoca:
        parts.append(epoca)
    return ". ".join(parts) + "." if parts else ""


# ═══════════════════════════════════════════════════════════════════════════════
# Markdown Parser
# ═══════════════════════════════════════════════════════════════════════════════

def _classify_line(line: str):
    """
    Classify a markdown line into a type for DOCX formatting.
    Returns: (type, content)
    Types: 'h1', 'h2', 'h3', 'quote', 'bold_line', 'empty', 'text'
    """
    stripped = line.strip()

    if not stripped:
        return ('empty', '')

    if stripped.startswith('### '):
        return ('h3', stripped[4:].strip())

    if stripped.startswith('## '):
        return ('h2', stripped[3:].strip())

    if stripped.startswith('# '):
        return ('h1', stripped[2:].strip())

    if stripped.startswith('> '):
        return ('quote', stripped[2:].strip())

    # Line that is entirely bold: **text**
    if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
        return ('bold_line', stripped[2:-2].strip())

    return ('text', stripped)


def _parse_inline(text: str):
    """
    Parse inline markdown (bold **text**) into segments.
    Returns list of (text, is_bold) tuples.
    """
    segments = []
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], True))
        elif part:
            segments.append((part, False))
    return segments


# ═══════════════════════════════════════════════════════════════════════════════
# Main Generator
# ═══════════════════════════════════════════════════════════════════════════════

def generate_docx_bytes(
    markdown_text: str,
    meta: Optional[dict] = None,
) -> bytes:
    """
    Convert pipeline v3 markdown output to a DOCX with PJF formatting.

    Args:
        markdown_text: The study markdown from Pass 3
        meta: Optional metadata dict with keys like 'tipo_asunto', 'materia', 'circuito'

    Returns:
        bytes of the DOCX file
    """
    doc = Document()
    doc._tcc_footnotes = []

    # ── Page margins ──
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_MAIN
    style.paragraph_format.line_spacing = LINE_SPACING_MAIN
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Process lines ──
    lines = markdown_text.split('\n')
    in_quote_block = False
    quote_buffer = []
    current_registro = None

    for line in lines:
        line_type, content = _classify_line(line)

        # Handle quote blocks (accumulate consecutive > lines)
        if line_type == 'quote':
            if not in_quote_block:
                in_quote_block = True
                quote_buffer = []
            quote_buffer.append(content)
            # Check for registro in the quote
            reg_match = RE_REGISTRO.search(content)
            if reg_match:
                current_registro = reg_match.group(1)
            continue
        elif in_quote_block:
            # Flush quote block
            _flush_quote_block(doc, quote_buffer, current_registro)
            in_quote_block = False
            quote_buffer = []
            current_registro = None

        if line_type == 'empty':
            # Add an empty paragraph with reduced spacing
            p = doc.add_paragraph()
            _set_paragraph_format(p, LINE_SPACING_MAIN)
            p.paragraph_format.space_after = Pt(4)
            continue

        if line_type in ('h1', 'h2'):
            p = doc.add_paragraph()
            _set_paragraph_format(p, LINE_SPACING_MAIN, WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            _add_run(p, content.upper(), bold=True, font_size=FONT_SIZE_MAIN)
            continue

        if line_type == 'h3':
            p = doc.add_paragraph()
            _set_paragraph_format(p, LINE_SPACING_MAIN, WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            _add_run(p, content, bold=True, font_size=FONT_SIZE_MAIN)
            continue

        if line_type == 'bold_line':
            p = doc.add_paragraph()
            _set_paragraph_format(p, LINE_SPACING_MAIN, WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.paragraph_format.space_before = Pt(8)
            _add_run(p, content, bold=True, font_size=FONT_SIZE_MAIN)
            continue

        # Regular text with inline bold
        if line_type == 'text':
            p = doc.add_paragraph()
            _set_paragraph_format(p, LINE_SPACING_MAIN, WD_ALIGN_PARAGRAPH.JUSTIFY)

            # Check if this line contains a registro reference in running text
            reg_match = RE_REGISTRO.search(content)

            segments = _parse_inline(content)
            for text, is_bold in segments:
                _add_run(p, text, bold=is_bold, font_size=FONT_SIZE_MAIN)

            # If a registro was found in running text, add footnote
            if reg_match:
                registro = reg_match.group(1)
                fn_text = _build_apa_citation(registro)
                if fn_text:
                    _add_footnote(p, fn_text, doc)
            continue

    # Flush any remaining quote block
    if in_quote_block and quote_buffer:
        _flush_quote_block(doc, quote_buffer, current_registro)

    # ── Add footnotes section ──
    if doc._tcc_footnotes:
        _add_footnotes_section(doc)

    # ── Serialize ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _flush_quote_block(doc: Document, lines: list[str], registro: Optional[str]):
    """
    Render a quote block as an indented, italic paragraph with single spacing.
    This is used for thesis rubros, article transcriptions, etc.
    """
    full_text = " ".join(lines)
    p = doc.add_paragraph()
    _set_paragraph_format(p, LINE_SPACING_CITE, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_indent(p, 1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Check if it looks like a rubro (ALL CAPS)
    is_rubro = RE_RUBRO_LINE.match(full_text[:100].strip()) is not None

    _add_run(p, full_text, italic=True, font_size=FONT_SIZE_CITE)

    # Add APA footnote if we have a registro
    if registro:
        rubro_text = full_text[:120] if is_rubro else ""
        fn_text = _build_apa_citation(registro, rubro=rubro_text)
        if fn_text:
            _add_footnote(p, fn_text, doc)


def _add_footnotes_section(doc: Document):
    """Add a footnotes section at the end of the document."""
    # Separator
    p = doc.add_paragraph()
    _set_paragraph_format(p, LINE_SPACING_MAIN)
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("_" * 40)
    run.font.size = Pt(8)
    run.font.color.rgb = None  # default color

    # Title
    p = doc.add_paragraph()
    _set_paragraph_format(p, LINE_SPACING_CITE, WD_ALIGN_PARAGRAPH.LEFT)
    _add_run(p, "NOTAS AL PIE", bold=True, font_size=FONT_SIZE_FOOTNOTE)

    # Each footnote
    for i, fn_text in enumerate(doc._tcc_footnotes, 1):
        p = doc.add_paragraph()
        _set_paragraph_format(p, LINE_SPACING_CITE, WD_ALIGN_PARAGRAPH.LEFT)
        _add_run(p, f"{i}. ", bold=True, font_size=FONT_SIZE_FOOTNOTE)
        _add_run(p, fn_text, italic=True, font_size=FONT_SIZE_FOOTNOTE)


# ═══════════════════════════════════════════════════════════════════════════════
# CLI test
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    test_md = """## ESTUDIO DE LOS CONCEPTOS DE VIOLACIÓN

El **primer concepto de violación** es **fundado** por las razones siguientes.

### Problema jurídico 1: Indebida valoración probatoria

La parte quejosa sostiene que la responsable valoró indebidamente las pruebas documentales ofrecidas en el juicio natural.

> PRUEBA DOCUMENTAL EN EL JUICIO ORAL MERCANTIL. SU VALORACIÓN DEBE REALIZARSE CONFORME A LAS REGLAS QUE PARA TAL EFECTO ESTABLECE EL CÓDIGO DE COMERCIO.
> Registro digital: 2018234. Tribunales Colegiados de Circuito. Undécima Época.

Al respecto, el artículo 1205 del Código de Comercio establece:

> Son admisibles como medios de prueba todos aquellos elementos que puedan producir convicción en el ánimo del juzgador acerca de los hechos controvertidos o dudosos.

**Conclusión del problema 1:** En virtud de lo anterior, resulta **fundado** el concepto de violación analizado.
"""
    result = generate_docx_bytes(test_md)
    with open("test_formato_pjf.docx", "wb") as f:
        f.write(result)
    print(f"[OK] DOCX generado: {len(result):,} bytes -> test_formato_pjf.docx")
