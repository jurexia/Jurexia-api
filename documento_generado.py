"""EL DOCUMENTO SIN PLANTILLA — se escribe entero, no se rellena.

David, 30-ago-2026: «El documento tiene que ser creado completamente por el
modelo, no rellenar huecos porque la herramienta no es exclusivamente para
secretarios del Tercer Tribunal Colegiado de Circuito en Materias
Administrativa y Civil del Vigésimo Segundo Circuito, sino de todo el país».

Y tiene razón por una razón que se vio medida: la plantilla NO es un formato,
es una sentencia real. Lleva dentro su tribunal, su magistrado, su quejoso, su
toca y sus fórmulas de sesión. Rellenar sus huecos funciona para quien la
escribió y produce basura para cualquier otro: un secretario de Yucatán
firmaría «Resolución del Tercer Tribunal Colegiado… del Vigésimo Segundo
Circuito» sin darse cuenta.

Aquí no hay huecos que rellenar porque no hay plantilla. Se compone:
  · lo que el asunto dice          → las fases (antecedentes, resúmenes, estudio)
  · lo que la ley obliga a decir   → lo escribe el modelo con los datos del caso
  · lo que el cómputo calcula      → la tabla, dibujada
  · lo que el tribunal es          → viene del encargo, NO del código

LO ÚNICO QUE QUEDA EN BLANCO es la fecha de la sesión, porque la sesión no ha
ocurrido. Eso no es un hueco de plantilla: es un dato que todavía no existe.

FORMATO, medido sobre los engroses reales y no inventado:
    papel oficio 21.59 × 34.03 cm · márgenes 5/2/3/3
    cuerpo Arial 14, justificado, sangría de primera línea 1.25 cm, 1.5 líneas
    cita  Arial 12, sin sangría, un espacio
"""

from __future__ import annotations

import json

# EL CATÁLOGO SE IMPORTA UNA VEZ, ARRIBA. Estaba importado dentro de las
# funciones que lo usaban, y al añadir un uso NUEVO más arriba que el
# `import` local, `_ta` quedaba sin ligar: el nombre es local a la función
# desde que aparece un import suyo en cualquier punto de ella, así que
# usarlo antes revienta con UnboundLocalError. No es un riesgo de ciclo:
# tipos_asunto no importa nada de este módulo.
import tipos_asunto as _ta

import os
import re
from dataclasses import dataclass, field

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

MODELO_ESTRUCTURA = os.getenv("MODELO_ESTRUCTURA",
                              os.getenv("MODELO_ESTUDIO", "gpt-5.6-luna"))
ESFUERZO_ESTRUCTURA = os.getenv("ESFUERZO_ESTRUCTURA", "medium")
MAX_TOKENS_ESTRUCTURA = int(os.getenv("MAX_TOKENS_ESTRUCTURA", "16000"))

FUENTE = "Arial"
TAMANO = Pt(14)
TAMANO_CITA = Pt(12)
TAMANO_TABLA = Pt(11)
SANGRIA = Cm(1.25)
# La sangría del bloque transcrito —artículo o tesis—, medida en el corpus.
SANGRIA_CITA = Cm(1.25)
INTERLINEADO = 1.5
INTERLINEADO_CITA = 1.0

# Negro y gris, como pidió David. Nada de color: es una sentencia.
NEGRO = "000000"
GRIS_CABECERA = "3B3B3B"      # fondo de la fila de encabezado
GRIS_ALTERNO = "F2F2F2"       # bandeado de filas
GRIS_LINEA = "9A9A9A"         # bordes
BLANCO = "FFFFFF"


# ═══════════════════════════════════════════════════════════════════════════
# Utillaje de formato
# ═══════════════════════════════════════════════════════════════════════════

def _fmt(p, sangria=True, tamano=TAMANO, interlineado=INTERLINEADO,
         alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.alignment = alineacion
    pf.line_spacing = interlineado
    pf.space_after = Pt(6)
    pf.first_line_indent = SANGRIA if sangria else Cm(0)
    for r in p.runs:
        r.font.name = FUENTE
        r.font.size = tamano
    return p


def parrafo(doc, texto, sangria=True, negrita=False, tamano=TAMANO,
            interlineado=INTERLINEADO, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = negrita
    return _fmt(p, sangria, tamano, interlineado, alineacion)


def tramos(doc, piezas, sangria=True, tamano=TAMANO,
           interlineado=INTERLINEADO, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """[(texto, {'bold':True}), …] en un solo párrafo."""
    p = doc.add_paragraph()
    for texto, est in piezas:
        if not texto:
            continue
        r = p.add_run(texto)
        r.bold = bool(est.get("bold"))
        r.italic = bool(est.get("italic"))
    return _fmt(p, sangria, tamano, interlineado, alineacion)


def rotulo(doc, texto):
    """«R E S U L T A N D O:» — centrado, en negrita y espaciado."""
    p = doc.add_paragraph()
    r = p.add_run(" ".join(texto.upper()))
    r.bold = True
    return _fmt(p, sangria=False, alineacion=WD_ALIGN_PARAGRAPH.CENTER)


def _sombrear(celda, color):
    tc = celda._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), color)
    tc.append(sh)


def _bordes(tabla, color=GRIS_LINEA, grosor="6"):
    tbl = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{lado}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), grosor)
        e.set(qn("w:color"), color)
        bordes.append(e)
    tbl.append(bordes)


def _celda(celda, texto, negrita=False, color=NEGRO, fondo=None,
           alineacion=WD_ALIGN_PARAGRAPH.LEFT):
    celda.text = ""
    p = celda.paragraphs[0]
    r = p.add_run(str(texto))
    r.bold = negrita
    r.font.name = FUENTE
    r.font.size = TAMANO_TABLA
    r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.alignment = alineacion
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if fondo:
        _sombrear(celda, fondo)


# ═══════════════════════════════════════════════════════════════════════════
# LAS NOTAS AL PIE
# ═══════════════════════════════════════════════════════════════════════════
# Un documento nuevo NO trae la parte `word/footnotes.xml`, y el utillaje del
# ensamblador clona una nota existente de la plantilla para heredar su estilo.
# Aquí no hay plantilla, así que la parte se escribe entera: el XML, su
# relación y su tipo de contenido. Sin esto, la localización de cada tesis
# —«Gaceta S.J.F., Undécima Época, Libro 52, tomo III, página 2489»— se
# quedaba en el cuerpo o se perdía, que es lo que David detectó.

_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NOTAS = ("http://schemas.openxmlformats.org/officeDocument/2006/"
              "relationships/footnotes")
_TIPO_NOTAS = ("application/vnd.openxmlformats-officedocument."
               "wordprocessingml.footnotes+xml")


def _run_llamada(parrafo, ident: int):
    """El numerito volado que llama a la nota."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    est = OxmlElement("w:rStyle")
    est.set(qn("w:val"), "FootnoteReference")
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(est)
    rpr.append(va)
    r.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(ident))
    r.append(ref)
    parrafo._p.append(r)


def _xml_notas(notas: list) -> bytes:
    """`word/footnotes.xml` con las dos notas de sistema y las nuestras."""
    def _esc(x):
        return (str(x).replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;"))
    piezas = [f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
              f'<w:footnotes xmlns:w="{_NS_W}">']
    for ident, tipo in ((-1, "separator"), (0, "continuationSeparator")):
        marca = "separator" if tipo == "separator" else "continuationSeparator"
        piezas.append(
            f'<w:footnote w:type="{tipo}" w:id="{ident}"><w:p><w:pPr>'
            f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            f'<w:r><w:{marca}/></w:r></w:p></w:footnote>')
    for i, texto in enumerate(notas, start=1):
        piezas.append(
            f'<w:footnote w:id="{i}"><w:p><w:pPr>'
            f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
            f'<w:jc w:val="both"/></w:pPr>'
            f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/>'
            f'<w:vertAlign w:val="superscript"/></w:rPr>'
            f'<w:footnoteRef/></w:r>'
            f'<w:r><w:rPr><w:rFonts w:ascii="{FUENTE}" w:hAnsi="{FUENTE}"/>'
            f'<w:sz w:val="18"/></w:rPr>'
            f'<w:t xml:space="preserve"> {_esc(texto)}</w:t></w:r>'
            f'</w:p></w:footnote>')
    piezas.append("</w:footnotes>")
    return "".join(piezas).encode("utf8")


def _inyectar_notas(ruta: str, notas: list) -> None:
    """Mete la parte de notas en el .docx ya guardado.

    python-docx no sabe crear notas al pie, así que se añaden sobre el paquete:
    el XML, la relación desde document.xml y el Override del tipo. Se reescribe
    el zip entero porque no se puede modificar una entrada en su sitio.
    """
    import shutil
    import zipfile as _zp
    if not notas:
        return
    tmp = ruta + ".tmp"
    with _zp.ZipFile(ruta) as z:
        nombres = z.namelist()
        datos = {n: z.read(n) for n in nombres}

    datos["word/footnotes.xml"] = _xml_notas(notas)

    rels = datos["word/_rels/document.xml.rels"].decode("utf8")
    if "footnotes.xml" not in rels:
        ids = re.findall(r'Id="rId(\d+)"', rels)
        nuevo = max((int(x) for x in ids), default=0) + 1
        rels = rels.replace(
            "</Relationships>",
            f'<Relationship Id="rId{nuevo}" Type="{_REL_NOTAS}" '
            f'Target="footnotes.xml"/></Relationships>')
        datos["word/_rels/document.xml.rels"] = rels.encode("utf8")

    ct = datos["[Content_Types].xml"].decode("utf8")
    if "footnotes+xml" not in ct:
        ct = ct.replace("</Types>",
                        f'<Override PartName="/word/footnotes.xml" '
                        f'ContentType="{_TIPO_NOTAS}"/></Types>')
        datos["[Content_Types].xml"] = ct.encode("utf8")

    with _zp.ZipFile(tmp, "w", _zp.ZIP_DEFLATED) as z:
        for n in list(datos):
            z.writestr(n, datos[n])
    shutil.move(tmp, ruta)


# ═══════════════════════════════════════════════════════════════════════════
# LA CITA DE UNA TESIS
# ═══════════════════════════════════════════════════════════════════════════
# Como la dictó David y como quedó medida en su corpus:
#
#     …de rubro y texto siguientes:          ← anuncio, fin de párrafo
#     «RUBRO EN NEGRITA.»                    ← párrafo aparte
#     Texto íntegro de la tesis…             ← cursiva, sangrado, 12pt, a uno
#     ÓRGANO EMISOR.¹                        ← y ahí la nota con la localización
#
# El modelo escribe el rubro EMBEBIDO en la prosa —«…siguientes: «RUBRO.» La
# responsable…»— y así la cita queda partida y sin transcripción. Esto la
# rehace desde el ACERVO, que es de donde tiene que salir el texto: palabra por
# palabra, no de la memoria del modelo.

_RX_RUBRO = re.compile(r"[“«\"]([A-ZÁÉÍÓÚÑ][^”»\"]{18,}?)[”»\"]")


# La coleta con que el modelo cierra el anuncio, para no escribirla dos veces.
_RX_COLA_ANUNCIO = re.compile(
    r"\s*,?\s*(?:cuy[oa]s?\s+|de\s+|del\s+)?"
    r"rubro(?:\s+y\s+(?:texto|registro|contenido))?"
    r"(?:\s+(?:es|son|siguientes?|los\s+siguientes?))*\s*[:.]?\s*$", re.I)

# Un estudio invoca entre tres y seis criterios; más transcripciones que eso
# convierten la sentencia en un compendio.
MAX_CITAS_DOCUMENTO = 8


def _normaliza_rubro(x: str) -> str:
    import unicodedata
    x = unicodedata.normalize("NFKD", (x or "").upper())
    x = "".join(c for c in x if not unicodedata.combining(c))
    return re.sub(r"[^A-Z0-9]+", " ", x).strip()


def tesis_del_rubro(texto: str, tesis: list):
    """La tesis del acervo cuyo rubro se cita en este párrafo, si alguna."""
    m = _RX_RUBRO.search(texto or "")
    if not m:
        return None, None
    citado = _normaliza_rubro(m.group(1))
    if len(citado) < 25:
        return None, None
    for t in (tesis or []):
        real = _normaliza_rubro(t.get("rubro", ""))
        if real and (real.startswith(citado[:70]) or citado.startswith(real[:70])):
            return t, m
    return None, m


# ═══════════════════════════════════════════════════════════════════════════
# EL ANUNCIO DE LA CITA SE COMPONE, NO SE COPIA
# ═══════════════════════════════════════════════════════════════════════════
# Una auditoría del proyecto 380/2025 encontró que TRES de las cuatro citas
# llamaban «jurisprudencia» a lo que son tesis aisladas, y una de ellas
# atribuía al Pleno como si fuera de la Primera Sala. Lo grave no es el error:
# es que el documento SE DESMENTÍA A SÍ MISMO tres párrafos después, porque la
# nota al pie —que sale del acervo— decía «[TA]; 9a. Época; Pleno». Cuerpo y
# nota, en la misma página, diciendo cosas distintas.
#
# Y la causa no era del modelo. Era MÍA, en dos sitios:
#
#   1. El prompt le daba este ejemplo literal de cómo se cita:
#         «Sirve de apoyo la jurisprudencia de la Primera Sala de la Suprema
#          Corte de Justicia de la Nación, de registro 2022074…»
#      El modelo lo copió y sólo cambió el número. Hizo lo que le pedí.
#   2. El bloque de material le enseñaba si el criterio VINCULA —«JURISPRUDENCIA
#      OBLIGATORIA» o «tesis orientadora»— pero nunca le enseñaba si es
#      jurisprudencia o tesis aislada. Dos cosas distintas que yo había fundido
#      en una etiqueta.
#
# Los dos se arreglan, pero ninguno de los dos es la defensa. La defensa es
# ésta: el anuncio se construye aquí, con los campos del acervo que ya llegan
# hasta la nota al pie. Del modelo se conserva SÓLO el verbo de enlace —«Sirve
# de apoyo», «Resulta aplicable»—, que es lo que ata la cita al razonamiento y
# es lo único que él sabe y yo no. Compuesto así el fallo es imposible: la
# frase y la nota nacen del mismo dato.

_RX_VERBO = re.compile(
    r"^(.{0,90}?)\s*(?:,\s*)?(?:resulta[n]?\s+)?(?:la|el|las|los)?\s*"
    r"(?:jurisprudencia|tesis|criterio)\b", re.I | re.S)

# Los verbos de enlace que sabemos leer. Si el modelo escribe otra cosa, se usa
# el suyo tal cual mientras no nombre instancia ni tipo; y si no hay nada
# aprovechable, «Sirve de apoyo», que es la fórmula del oficio.
_POR_DEFECTO = "Sirve de apoyo"


# Lo que puede quedar colgando al recortar el sintagma: el modelo escribe
# «Sirve de apoyo, como criterio orientador, la jurisprudencia…» y el recorte se
# lleva desde «criterio», dejando «Sirve de apoyo, como». Compuesto luego da
# «Sirve de apoyo, como, la jurisprudencia…», con la palabra huérfana entre
# comas. Se poda la cola.
_RX_COLA_HUERFANA = re.compile(
    r"[\s,;:]*\b(como|en\s+calidad\s+de|a\s+t[íi]tulo\s+de|con\s+car[áa]cter"
    r"|por|de|del|la|el|los|las|y|e)\s*$", re.I)


def _verbo_de_enlace(anuncio: str) -> str:
    """Lo único del anuncio que escribe el modelo y merece conservarse."""
    a = " ".join((anuncio or "").split())
    if not a:
        return _POR_DEFECTO
    m = _RX_VERBO.match(a)
    if m and m.group(1).strip():
        v = m.group(1).strip().rstrip(",;:")
        # Se poda dos veces: «Sirve de apoyo, como» → «Sirve de apoyo»; y un
        # «Resulta aplicable, en calidad de» → «Resulta aplicable».
        for _ in range(2):
            v = _RX_COLA_HUERFANA.sub("", v).strip()
        return v or _POR_DEFECTO
    # Sin sustantivo reconocible: se conserva sólo si es corto y no nombra
    # órgano ni tipo, que es lo que no puede venir de él.
    if len(a) <= 60 and not re.search(
            r"sala|pleno|colegiado|jurisprudencia|tesis aislada", a, re.I):
        return a.rstrip(" ,;:")
    return _POR_DEFECTO


def anuncio_de(t: dict, anuncio_del_modelo: str = "") -> str:
    """«Sirve de apoyo la tesis aislada del Pleno…, de registro 191358».

    Cada pieza sale del acervo. `tipo` decide el sustantivo, `instancia` el
    órgano y `obligatoria` el calificativo —esa parte YA funcionaba: las tres
    aisladas iban «como criterio orientador» y la única jurisprudencia real «de
    carácter obligatorio»—; lo que fallaba era el sustantivo y el nombre del
    órgano, que venían del modelo.
    """
    tipo = str(t.get("tipo") or "").strip().upper()
    inst = " ".join(str(t.get("instancia") or "").strip().split())
    reg = str(t.get("registro") or "").strip()
    verbo = _verbo_de_enlace(anuncio_del_modelo)

    if "AISLAD" in tipo:
        sustantivo = "la tesis aislada"
    elif "JURISPRUDENCIA" in tipo:
        sustantivo = "la jurisprudencia"
    else:
        # SIN EL DATO NO SE AFIRMA NINGUNO DE LOS DOS. «El criterio» es cierto
        # de la jurisprudencia y de la tesis aislada, así que no miente; decir
        # «jurisprudencia» sin saberlo, sí.
        sustantivo = "el criterio"

    # EL ORDEN ES EL DEL OFICIO, no el que salga. Un tribunal escribe «Sirve de
    # apoyo, como criterio orientador, la tesis aislada del Pleno de la Suprema
    # Corte de Justicia de la Nación, de registro digital 191358, de rubro y
    # texto siguientes:». El calificativo va entre comas detrás del verbo, no
    # colgando del órgano, donde suena a que el órgano es el orientador.
    # EL CALIFICATIVO SÓLO CUANDO APORTA. En una jurisprudencia la
    # obligatoriedad va de suyo y escribirla suena a énfasis de quien no está
    # seguro; en una tesis aislada, en cambio, decir que sólo orienta es
    # información necesaria y es lo que evita que se lea como vinculante.
    if t.get("obligatoria"):
        # Un verbo con inciso propio —«Es aplicable, además»— pide cerrar la
        # coma antes del sustantivo, o queda «además la jurisprudencia».
        frase = (f"{verbo}, {sustantivo}" if "," in verbo
                 else f"{verbo} {sustantivo}")
    else:
        frase = f"{verbo}, como criterio orientador, {sustantivo}"
    if inst:
        # El artículo, según el órgano. «del Tribunales Colegiados» no es
        # español: los órganos en plural piden «de los».
        if re.match(r"^(primera|segunda|tercera|cuarta)\s+sala", inst, re.I):
            de = "de la"
        elif re.search(r"^(tribunales|plenos|salas)\b", inst.strip(), re.I):
            de = "de los" if not inst.strip().lower().startswith("salas") else "de las"
        else:
            de = "del"
        # El Pleno y las Salas son de la Corte y así se nombran en un engrose;
        # los colegiados y los plenos regionales traen su nombre completo en el
        # propio campo y no se les añade nada.
        completo = inst
        if re.match(r"^(pleno|primera\s+sala|segunda\s+sala)$", inst.strip(), re.I):
            completo = f"{inst} de la Suprema Corte de Justicia de la Nación"
        frase += f" {de} {completo}"
    if reg:
        frase += f", de registro digital {reg}"
    return frase + ", de rubro y texto siguientes:"


def escribir_cita(doc, t: dict, anuncio: str, notas: list) -> None:
    """El bloque entero de la cita, con su nota al pie."""
    # EL ANUNCIO NO SE COPIA: se compone del acervo. Ver `anuncio_de`.
    parrafo(doc, anuncio_de(t, anuncio))

    # El rubro, solo y en negrita. Sin párrafo vacío delante: la cita va
    # pegada a su anuncio y el aire lo pone el espaciado.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"«{(t.get('rubro') or '').strip().rstrip('.')}.»")
    r.bold = True
    _fmt(p, sangria=False, tamano=TAMANO_CITA, interlineado=INTERLINEADO_CITA)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.keep_with_next = True

    # EL TEXTO DE LA TESIS BAJA A LA NOTA AL PIE. Iba íntegro en el cuerpo, y
    # eso es lo que hace que un estudio con seis criterios invocados quede
    # sepultado bajo sus propias citas: medido en los engroses de referencia,
    # las tres transcripciones de tesis del ARA 17/2025 se llevan la mitad del
    # considerando, y en la queja 233/2025 la transcripción de una ejecutoria
    # de la Corte ocupa 2,260 de 3,798 palabras.
    #
    # El rubro se queda arriba —identifica el criterio y se lee de un vistazo—
    # y el texto va abajo, donde quien firma lo comprueba si quiere. Es un
    # cambio sobre el corpus, no una imitación suya, y está pedido: «nuestro
    # redactor debe ser mejor que el secretario que redactó esos proyectos».
    #
    # SE CONSERVA EN EL CUERPO CUANDO ES CORTO. Una tesis de cuatro renglones
    # leída al pie es una molestia sin ganancia; el problema son las de
    # trescientas palabras.
    cuerpo = (t.get("texto") or "").strip()
    _al_pie = len(cuerpo.split()) > MAX_PALABRAS_TESIS_CUERPO
    if cuerpo and not _al_pie:
        q = doc.add_paragraph()
        rq = q.add_run(cuerpo)
        rq.italic = True
        _fmt(q, sangria=False, tamano=TAMANO_CITA,
             interlineado=INTERLINEADO_CITA)
        q.paragraph_format.left_indent = Cm(1.25)
        # EL BLOQUE LARGO NO SE ATA A LO SIGUIENTE. Con `keep_with_next` en
        # toda la cadena —anuncio, rubro, texto y órgano— Word empuja el
        # conjunto entero a la página siguiente y deja media hoja en blanco:
        # ése era el «espacio enorme antes de citar una tesis». El rubro sí
        # sigue atado a su texto, que es lo que no puede partirse.
        q.paragraph_format.keep_with_next = False

    # El órgano, y ahí cuelga la nota con la localización.
    inst = (t.get("instancia") or "").strip()
    loc = (t.get("localizacion") or "").strip()
    reg = str(t.get("registro") or "").strip()
    if inst or loc:
        z = doc.add_paragraph()
        rz = z.add_run(inst.upper() + ("." if inst else ""))
        rz.bold = True
        _fmt(z, sangria=False, tamano=TAMANO_CITA,
             interlineado=INTERLINEADO_CITA)
        z.paragraph_format.left_indent = Cm(1.25)
        z.paragraph_format.keep_with_next = False
        if loc or reg or _al_pie:
            pie = loc if loc else ""
            if reg and reg not in pie:
                pie = (pie + ", " if pie else "") + f"registro digital {reg}"
            if _al_pie:
                pie = (pie + ". " if pie else "") + f"Texto: {cuerpo}"
            if pie in notas:
                _run_llamada(z, notas.index(pie) + 1)   # se reusa la existente
            else:
                notas.append(pie)
                _run_llamada(z, len(notas))


# ═══════════════════════════════════════════════════════════════════════════
# LA TABLA DEL CÓMPUTO
# ═══════════════════════════════════════════════════════════════════════════

def _sin_partir(tabla) -> None:
    """`cantSplit`: la fila entera va a la página donde quepa."""
    from docx.oxml.ns import qn
    for fila in tabla.rows:
        trPr = fila._tr.get_or_add_trPr()
        el = trPr.makeelement(qn("w:cantSplit"), {})
        trPr.append(el)


def tabla_computo(doc, computo, fecha_en_letra,
                  tipo_asunto: str = "amparo_directo") -> None:
    """El cómputo del plazo, en negro y gris.

    No es adorno: es la parte de la sentencia que más se revisa y la que peor
    se lee en prosa. Una fila por hito, la fecha al lado, y el resultado
    destacado abajo. Quien la revisa comprueba en diez segundos lo que en un
    párrafo corrido cuesta releer tres veces.
    """
    # LOS RÓTULOS SON DEL TIPO, NO DEL AMPARO DIRECTO. Una queja mostraba
    # «Notificación de la sentencia reclamada» y «Presentación de la demanda»
    # cuando lo que se notificó fue un auto y lo que se presentó, un recurso.
    # El resto del documento ya hablaba con el vocabulario correcto: sólo la
    # tabla seguía anclada al tipo con el que nació.
    _v = _ta.vocabulario_de(tipo_asunto)
    # El «recurrido» del catálogo trae artículo —«el auto recurrido»— y el
    # «escrito» no —«recurso de queja»—: cada uno necesita su contracción.
    from fase0_oportunidad import _del as _del_

    def _complemento(x: str) -> str:
        return "del " + x[3:] if x.startswith("el ") else "de " + x

    _recurrido = _complemento(_v["recurrido"])
    _escrito = _del_(_v["escrito"])

    filas = [
        (f"Notificación {_recurrido}",
         fecha_en_letra(computo.notificacion)),
        (f"Surtimiento de efectos ({computo.regla.descripcion})",
         fecha_en_letra(computo.surtio)),
    ]
    # SIN PLAZO NO HAY VENCIMIENTO QUE ENSEÑAR. Con el cómputo ya blindado
    # contra el plazo cero, la tabla habría escrito «Plazo legal: 0 días
    # hábiles» y un vencimiento igual al día de inicio: una fecha inventada
    # para un asunto en el que la ley dice que no vence nada.
    if getattr(computo, "en_cualquier_tiempo", False):
        filas.append(("Plazo legal", "no hay: procede en cualquier tiempo"))
    else:
        filas += [
            ("Inicio del plazo", fecha_en_letra(computo.inicio)),
            ("Plazo legal", f"{computo.plazo} días hábiles"),
            ("Vencimiento del plazo", fecha_en_letra(computo.vencimiento)),
        ]
    if computo.presentacion is not None:
        filas.append((f"Presentación {_escrito}",
                      fecha_en_letra(computo.presentacion)))
    if computo.inhabiles_en_medio:
        # LA FILA MENTÍA. Decía «Días inhábiles descontados: 1» mientras el
        # párrafo explicaba que se descontaron seis sábados y domingos más un
        # festivo. Contaba SÓLO los inhábiles declarados, no los fines de
        # semana. Se dice lo que de verdad cuenta.
        filas.append(("Días inhábiles declarados (además de sábados y domingos)",
                      str(len(computo.inhabiles_en_medio))))

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _bordes(t)

    _celda(t.rows[0].cells[0], "CÓMPUTO DEL PLAZO", negrita=True,
           color=BLANCO, fondo=GRIS_CABECERA)
    _celda(t.rows[0].cells[1], "FECHA", negrita=True, color=BLANCO,
           fondo=GRIS_CABECERA, alineacion=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (concepto, valor) in enumerate(filas):
        fila = t.add_row()
        fondo = GRIS_ALTERNO if i % 2 == 0 else None
        _celda(fila.cells[0], concepto, fondo=fondo)
        _celda(fila.cells[1], valor, fondo=fondo,
               alineacion=WD_ALIGN_PARAGRAPH.CENTER)

    if computo.oportuna is not None:
        fila = t.add_row()
        veredicto = ("PRESENTADA ANTES DEL INICIO DEL PLAZO"
                     if computo.anticipada
                     else "PRESENTADA EN TIEMPO" if computo.oportuna
                     else "PRESENTADA FUERA DE PLAZO")
        _celda(fila.cells[0], "Resultado", negrita=True, color=BLANCO,
               fondo=GRIS_CABECERA)
        _celda(fila.cells[1], veredicto, negrita=True, color=BLANCO,
               fondo=GRIS_CABECERA, alineacion=WD_ALIGN_PARAGRAPH.CENTER)

    # NINGUNA FILA SE PARTE ENTRE PÁGINAS, y se declara AL FINAL: puesto tras
    # `add_table` sólo alcanzaba a la única fila que existía entonces, y las
    # ocho que vienen después —las que de verdad se parten— se quedaban fuera.
    # Cada fila es de un renglón; partida, Word deja media y su sombreado al
    # pie de una hoja, que es la franja negra suelta.
    _sin_partir(t)

    for fila in t.rows:
        fila.cells[0].width = Cm(9.5)
        fila.cells[1].width = Cm(5.0)


# ═══════════════════════════════════════════════════════════════════════════
# Lo que la ley obliga a decir — lo escribe el modelo, con los datos del caso
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class Estructura:
    apertura: str = ""
    visto: str = ""
    resultandos: list = field(default_factory=list)   # [{"titulo","texto"}]
    competencia: str = ""
    existencia: str = ""
    procedencia: str = ""
    avisos: list = field(default_factory=list)


_RX_JSON = re.compile(r"\{.*\}", re.S)


# CADA ASUNTO SE IDENTIFICA CON LO QUE TIENE, y no todos tienen lo mismo. El
# prompt ordenaba identificar el acto por «fecha, SALA, TOCA y expediente de
# origen, y qué confirmó, modificó o revocó»: eso es un amparo directo contra
# una sentencia de segunda instancia. En una QUEJA se recurre un auto de un juez
# de distrito y en una REVISIÓN una sentencia de amparo indirecto: no hay sala,
# no hay toca y no hay nada que confirmar ni revocar. El modelo, obligado a
# decirlo, escribía «dentro del toca y expediente de origen que constan en
# autos, acto que no confirmó, modificó ni revocó otra resolución», que es la
# frase defensiva que un dictamen ya nos reprochó: una sentencia no explica al
# lector por qué NO hay toca; sencillamente no lo menciona.
_IDENTIFICA_ACTO = {
    "amparo_directo": ("fecha, sala, toca y expediente de origen, y qué "
                       "confirmó, modificó o revocó"),
    "amparo_revision": ("fecha, juzgado de distrito y número del juicio de "
                        "amparo indirecto en que se dictó"),
    "revision_fiscal": ("fecha, sala del Tribunal Federal de Justicia "
                        "Administrativa y número del juicio de nulidad"),
    "queja": ("fecha, juzgado de distrito y número del juicio de amparo en que "
              "se dictó, y qué proveyó"),
}
_NOMBRE_ASUNTO = {
    "amparo_directo": "amparo directo",
    "amparo_revision": "amparo en revisión",
    "revision_fiscal": "revisión fiscal",
    "queja": "recurso de queja",
}


def prompt_estructura(datos: dict) -> str:
    q = "agravios" if datos.get("es_recurso") else "conceptos de violación"
    _tipo = str(datos.get("tipo_asunto") or "amparo_directo").strip().lower()
    _clase = _NOMBRE_ASUNTO.get(_tipo, "amparo directo")
    # LOS RESULTANDOS SON DEL TIPO. Estaban escritos a mano aquí —los cuatro
    # del amparo directo— y salían iguales en una queja y en una revisión
    # fiscal. Ahora se arman del catálogo, que es donde están medidos.
    import tipos_asunto as _ta_r
    _rs = []
    for _i, (_rot, _que) in enumerate(_ta_r.resultandos_de(_tipo)):
        _extra = ("" if _i else
                  " Si alguno de esos datos NO consta, NO lo menciones ni "
                  "expliques que no consta: se omite y ya. PROHIBIDO resumir "
                  "aquí su razonamiento: eso va en el estudio.")
        _rs.append('     {{"titulo": %s, "texto": "<%s>"}}'
                   % (json.dumps(_rot, ensure_ascii=False), _que + _extra))
    _resultandos = ",\n".join(_rs)

    # LA FÓRMULA DEL PROEMIO SALE DEL CATÁLOGO, NO DE UN EJEMPLO. Este campo
    # traía ESCRITA la del amparo directo —«para resolver el juicio de amparo
    # directo…»— y un modelo con un ejemplo concreto delante lo copia y le
    # cambia los datos: de ahí salían «V I S T O, para resolver el juicio de
    # amparo directo relacionado con el recurso de queja civil» y el mismo
    # apócrifo en la revisión fiscal. Es la tercera vez en este proyecto que un
    # ejemplo del prompt se firma literal —antes fue «la jurisprudencia de la
    # Primera Sala» y «si lo traes, es por analogía»—, así que la regla ya no
    # es una sospecha: en un prompt, lo que se escribe entero se copia entero.
    _molde_visto = _ta_r.proemio_de(_tipo)["molde"]

    # LA HOJA DE DATOS TAMBIÉN ENSEÑABA EL AMPARO DIRECTO. Le decía al motor
    # «QUEJOSO: …» y «AUTORIDAD RESPONSABLE: …» en los cuatro tipos, así que
    # aunque los rótulos de los resultandos ya vinieran bien del catálogo, la
    # PROSA seguía hablando de quejoso y de autoridad responsable dentro de una
    # queja. Era el canal por el que la plantilla única volvía a entrar después
    # de haberla quitado de la estructura.
    _vc = _ta_r.vocabulario_de(_tipo)
    _ficha_partes = "\n".join(
        f"{_et}: {datos.get(_cl, '')}"
        for _et, _cl, _ob in _ta_r.caratula_de(_tipo)
        if datos.get(_cl) or _ob)
    _rotulo_acto = {"amparo_directo": "ACTO RECLAMADO"}.get(
        _tipo, _vc["sub_recurrido"].upper())
    _de_escrito = ("DE LA DEMANDA" if _tipo == "amparo_directo"
                   else "DEL " + _vc["escrito"].upper())

    # ERA CÓDIGO MUERTO Y ERA JUSTO LA INSTRUCCIÓN QUE FALTABA: se calculaba y
    # no se interpolaba en ningún sitio del prompt. Dice cómo se identifica el
    # acto en cada tipo —«fecha, sala, toca y expediente de origen» en el
    # amparo directo— y ahora entra donde sirve, junto a la ficha de datos.
    _identifica = _IDENTIFICA_ACTO.get(_tipo, _IDENTIFICA_ACTO["amparo_directo"])
    return f"""Eres el secretario de un Tribunal Colegiado de Circuito y escribes las
partes ESTRUCTURALES de una sentencia de {_clase}. No escribes el estudio
de fondo —ese ya está hecho—: escribes lo que la ley obliga a decir antes de
llegar a él, con los datos de ESTE asunto y de ESTE tribunal.

EL TRIBUNAL QUE RESUELVE: {datos.get('tribunal','')}
CIUDAD: {datos.get('ciudad','')}
EXPEDIENTE: {datos.get('encabezado','')}
{_ficha_partes}
{_rotulo_acto} —léelo de aquí e IDENTIFÍCALO por {_identifica}—:
{(datos.get('acto') or '').strip() or '(no se aportó: NO lo describas, omite la mención)'}
FECHA DE PRESENTACIÓN {_de_escrito}: {datos.get('presentacion','')}
MAGISTRADO PONENTE: {datos.get('magistrado','')}
SECRETARIO: {datos.get('secretario','')}

ANTECEDENTES DEL ASUNTO, ya redactados
{datos.get('antecedentes','')[:4000]}

REGLAS:
- ESTOS SON LOS ROTULOS MEDIDOS EN 26 ENGROSES DE ESTE TRIBUNAL, no una
  propuesta: escríbelos tal cual. El resultando lleva esos cuatro apartados y
  el documento añade solo el de la sesión.
- NO ESCRIBAS ORDINALES. Nada de «PRIMERO.» ni «SEGUNDO.»: el documento los
  calcula y ponerlos aquí los duplica.
- NO INVENTES DATOS. Si no sabes una fecha, un número de toca o un nombre, NO
  lo pongas y NO lo sustituyas por uno verosímil: redacta la frase de modo que
  no lo necesite, o deja constancia de que consta en autos. Un dato inventado
  en un resultando se firma.
- EL TRIBUNAL ES EL DE ARRIBA, no otro. La competencia se funda en los
  artículos 103, fracción I, y 107, fracción V, de la Constitución; 33,
  fracción II, 34 y 170, fracción I, de la Ley de Amparo; y 37 de la Ley
  Orgánica del Poder Judicial de la Federación, con el acuerdo general que
  fije la jurisdicción territorial de ese tribunal —si no sabes cuál es, no
  cites acuerdo alguno—.
- FRASE de unas 35 palabras, subordinada. Voz impersonal: «se estima», «este
  Tribunal Colegiado considera». Nunca primera persona del singular.
- Sin Markdown, sin viñetas.

Devuelve SÓLO este JSON:
{{"apertura": "<ciudad y fórmula de resolución del tribunal; la FECHA DE LA SESIÓN se deja como ___ porque aún no ocurre>",
  "visto": "<{_molde_visto}. SIN repetir el rótulo V I S T O / V I S T O S, que lo pone el compositor>",
  "resultandos": [
{_resultandos}
  ],
  "competencia": "<por qué este tribunal es competente, con su fundamento>",
  "existencia": "<la existencia del acto reclamado, acreditada con el informe justificado y los autos>",
  "procedencia": "<que el juicio es procedente y no se advierte causa de improcedencia, o cuál>"}}"""


async def redactar_estructura(cliente, datos: dict) -> Estructura:
    kw = dict(model=MODELO_ESTRUCTURA,
              max_completion_tokens=MAX_TOKENS_ESTRUCTURA,
              messages=[{"role": "user", "content": prompt_estructura(datos)}])
    if ESFUERZO_ESTRUCTURA:
        kw["reasoning_effort"] = ESFUERZO_ESTRUCTURA
    import llamada_modelo as _lm
    r = await _lm.crear(cliente, **kw)
    crudo = (r.choices[0].message.content or "").strip()
    m = _RX_JSON.search(crudo)
    if not m:
        return Estructura(avisos=[
            "El motor no devolvió las partes estructurales "
            f"({'sin texto' if not crudo else 'JSON ilegible'}). El documento "
            "sale sin resultandos ni competencia: revísalo antes de firmar."])
    try:
        d = json.loads(m.group(0))
    except Exception as e:
        return Estructura(avisos=[f"El JSON de la estructura no se pudo leer: {e}"])
    return Estructura(
        apertura=str(d.get("apertura", "")),
        visto=str(d.get("visto", "")),
        resultandos=[{"titulo": str(x.get("titulo", "")),
                      "texto": str(x.get("texto", ""))}
                     for x in (d.get("resultandos") or [])],
        competencia=str(d.get("competencia", "")),
        existencia=str(d.get("existencia", "")),
        procedencia=str(d.get("procedencia", "")))


# ═══════════════════════════════════════════════════════════════════════════
# EL MARCO JURÍDICO — se compone, no se pide
# ═══════════════════════════════════════════════════════════════════════════
# Dos intentos por prompt fallaron: el material del bloque de
# constitucionalidad llegaba al estudio —6,400 caracteres con el artículo 4º y
# la Convención sobre los Derechos del Niño— y el modelo no lo escribía. Se
# movió al 93% del prompt y siguió sin escribirlo.
#
# Lo que tiene que aparecer se compone. El marco pasa a ser un apartado propio
# del CONSIDERANDO, escrito por su propia llamada y colocado por el compositor
# ANTES del estudio. Sigue dependiendo del caso —si el acervo no devuelve nada
# constitucional ni convencional, el apartado no existe—, que es lo que David
# pidió: «no fijar un marco constitucional para todos los casos, sino sobre la
# solución en función del problema jurídico».

# Lo que se nombra tiene que estar en el material. En el ADC 380/2025 el
# estudio citó la Convención sobre los Derechos del Niño SEIS veces sin que la
# capa convencional se hubiera buscado siquiera: el modelo la escribió de
# memoria. Una cita que nadie puede comprobar es lo que este sistema existe
# para evitar, y da igual que la Convención exista: lo que no consta, no consta.
_RX_TRATADO = re.compile(
    r"(Convenci[óo]n\s+(?:sobre|Americana|Interamericana|de)[^.,;]{0,60}"
    r"|Pacto\s+(?:de\s+San\s+Jos[ée]|Internacional[^.,;]{0,40})"
    r"|Protocolo\s+de\s+San\s+Salvador)", re.I)
_RX_COIDH = re.compile(r"Corte\s+Interamericana|CoIDH|caso\s+[A-ZÁÉÍÓÚÑ][\w]+\s+Vs\.",
                       re.I)


def revisar_marco(marco_escrito: str, material_marco: str) -> list:
    """Lo que el marco nombra y el acervo no respalda."""
    fuera = []
    if not (marco_escrito or "").strip():
        return fuera
    mat = material_marco or ""
    nombrados = {m.group(0).strip() for m in _RX_TRATADO.finditer(marco_escrito)}
    sin_respaldo = [x for x in nombrados
                    if x.split()[0].lower() not in mat.lower()
                    or not any(p.lower() in mat.lower()
                               for p in x.split()[:4] if len(p) > 4)]
    if sin_respaldo:
        fuera.append(
            f"El marco nombra instrumentos que NO están en el material "
            f"recuperado: {sorted(sin_respaldo)[:4]}. El modelo los escribió de "
            f"memoria; compruébalos antes de firmar.")
    if _RX_COIDH.search(marco_escrito) and not _RX_COIDH.search(mat):
        fuera.append(
            "El marco invoca a la Corte Interamericana y el acervo no devolvió "
            "ni un fragmento suyo para este asunto. Sin la fuente no se cita.")
    return fuera


async def redactar_marco(cliente, material_marco: str, problemas: list,
                         es_recurso: bool = False,
                         tipo_asunto: str = "") -> str:
    """El marco jurídico, escrito. Devuelve texto vacío si no hay material."""
    if not (material_marco or "").strip():
        return ""
    # La bisagra de cierre —«…dar solución a los planteamientos de la parte
    # quejosa»— estaba escrita a mano y duplicada literal en el prompt del
    # estudio: es una fórmula que el modelo copia palabra por palabra, así que
    # metía «la parte quejosa» en el punto de bisagra de los cuatro tipos.
    _t = tipo_asunto or ("amparo_revision" if es_recurso else "amparo_directo")
    _voc_m = _ta.vocabulario_de(_t)
    _parte_prosa = _voc_m["parte"]
    q = _voc_m["combate"]
    lista = "\n".join(
        f"- {p.get('pregunta','') if isinstance(p, dict) else str(p)}"
        for p in (problemas or []))
    prompt = f"""Eres el secretario de un Tribunal Colegiado y escribes el apartado
de MARCO JURÍDICO de una sentencia de amparo: la premisa mayor con la que
después se resolverán los planteamientos.

LOS PROBLEMAS QUE HAY QUE RESOLVER
{lista}

MATERIAL RECUPERADO DEL BLOQUE DE CONSTITUCIONALIDAD Y DE LA LEY LOCAL
{material_marco[:12000]}

CÓMO SE ESCRIBE, medido sobre los engroses de este tribunal:
- ARRANCA POR LA FIGURA JURÍDICA discutida —los alimentos, la convivencia, la
  acción—, NO por los derechos humanos en abstracto.
- LA CONSTITUCIÓN SE PARAFRASEA, no se transcribe: «el artículo 4º de la
  Constitución reconoce el derecho de la niñez a…». Nómbrala expresamente, con
  su número de artículo.
- Y NOMBRA TODOS LOS ARTÍCULOS CONSTITUCIONALES QUE TE DIERON. Están arriba
  porque el asunto los tocó; el que no escribas queda sin premisa. Si uno de
  verdad no viene al caso, DILO en una frase —«el artículo X no rige aquí
  porque…»— en vez de callarlo: el silencio no se distingue del olvido.
- EL PRECEPTO LOCAL O SECUNDARIO decisivo SÍ se transcribe, entre comillas y
  con su número al frente.
- LA FUENTE CONVENCIONAL —Convención sobre los Derechos del Niño, Convención
  Americana— y los criterios de la CORTE INTERAMERICANA entran SÓLO si el
  problema los exige. Cuando entran, se dice qué obligación imponen, no que
  existen.
- SI ARRIBA HAY MATERIAL DE LA CORTE INTERAMERICANA, ÚSALO. Está ahí porque el
  acervo lo encontró para ESTOS problemas, y desaprovecharlo deja el marco a
  medias. Se cita por el caso y el párrafo que trae la ficha —«Caso X Vs. Y,
  párr. N»— y se dice qué estándar fija, no que existe. Si de veras no viene al
  caso, no lo pongas; pero entonces tampoco cites la Corte de memoria.
- NO INVENTES NADA que no esté en el material de arriba. Ni un artículo, ni un
  caso, ni un párrafo de cuadernillo.
- EXTENSIÓN: entre 300 y 600 palabras. Frase de unas 35 palabras, subordinada,
  voz impersonal. Sin Markdown ni viñetas.
- CIERRA con la bisagra que devuelve al expediente: «Con ese marco jurídico, es
  posible dar solución a los planteamientos de {_parte_prosa}.»

Devuelve SÓLO el texto del apartado, en párrafos separados por una línea en
blanco. Sin rótulo ni encabezado: el documento se lo pone."""
    kw = dict(model=MODELO_ESTRUCTURA,
              max_completion_tokens=MAX_TOKENS_ESTRUCTURA,
              messages=[{"role": "user", "content": prompt}])
    if ESFUERZO_ESTRUCTURA:
        kw["reasoning_effort"] = ESFUERZO_ESTRUCTURA
    import llamada_modelo as _lm
    r = await _lm.crear(cliente, **kw)
    return (r.choices[0].message.content or "").strip()


# ═══════════════════════════════════════════════════════════════════════════
# La composición
# ═══════════════════════════════════════════════════════════════════════════

_ORDINALES = ("PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO",
              "SÉPTIMO", "OCTAVO", "NOVENO", "DÉCIMO")

_AMPARA = "ampara y protege"
_NO_AMPARA = "no ampara ni protege"

# ═══════════════════════════════════════════════════════════════════════════
# EL RESOLUTIVO NO ES EL MISMO EN TODOS LOS ASUNTOS
# ═══════════════════════════════════════════════════════════════════════════
# Se generaron cinco asuntos reales del corpus del secretario y se compararon
# con sus engroses. Tres de los cinco salieron con un resolutivo JURÍDICAMENTE
# IMPOSIBLE, porque esta función escribía la fórmula del amparo directo pasara
# lo que pasara:
#
#   queja QA 143/2026    engrose: «ÚNICO. Es fundado el recurso de queja.»
#                        motor:   «La Justicia de la Unión ampara y protege…»
#   revisión fiscal 6/25 engrose: «ÚNICO. Se confirma la sentencia de cuatro de
#                                  octubre…, dictada en el expediente 293/24…»
#                        motor:   «La Justicia de la Unión ampara y protege…»
#
# Una QUEJA no ampara: se declara fundada o infundada. Una REVISIÓN no ampara:
# confirma, revoca o modifica la sentencia recurrida. Sólo el amparo directo
# —y el resolutivo de fondo de una revisión que ampara— usan la fórmula de la
# Justicia de la Unión. Escribirla en una queja no es un defecto de estilo: es
# una resolución que no existe en derecho.
#
# Las fórmulas salen LITERALES de los engroses del propio tribunal, no de mi
# idea de cómo se redactan.
RESOLUTIVO = {
    "queja": {
        "punto": "Es {calificacion} el recurso de queja.",
        "calif": ("fundado", "infundado"),
        "notif": ("Notifíquese; publíquese y anótese en el libro de control de "
                  "este tribunal, hágase la captura correspondiente en el "
                  "Sistema Integral de Seguimiento de Expedientes, envíese "
                  "testimonio de esta resolución al juzgado de origen y, en su "
                  "oportunidad archívese como asunto concluido."),
    },
    "revision_fiscal": {
        # LA FÓRMULA MEDIDA, no la mía. `banco_formulas_medidas.json` la cuenta
        # en 16 de 28 revisiones fiscales del tribunal: «Se confirma la
        # sentencia DE {fecha}, dictada EN EL EXPEDIENTE {expediente}, por la
        # Sala…». Aquí se escribía «la sentencia recurrida, dictada por la
        # Sala», que no dice CUÁL sentencia se revoca —y en un tribunal donde
        # la misma Sala dicta cientos, identificarla no es un adorno—.
        #
        # Los dos datos se leen: el expediente con `fase_origen.numero_de`, que
        # acierta 5 de 5 sobre los engroses reales, y la fecha con `fecha_de`,
        # que ahora calla cuando duda. Lo que no se pudo leer sale en hueco, a
        # la vista, con su aviso.
        "punto": ("Se {calificacion} la sentencia de {fecha_sentencia}, dictada "
                  "en el expediente {expediente_origen}, por {responsable}."),
        "calif": ("revoca", "confirma"),
        "notif": ("Notifíquese; publíquese y anótese en el Libro de control de "
                  "este Tribunal, hágase la captura correspondiente en el "
                  "Sistema Integral de Seguimiento de Expedientes, con "
                  "testimonio de esta resolución vuelvan los autos a su lugar "
                  "de origen y, en su oportunidad archívese como asunto "
                  "concluido."),
    },
    "amparo_revision": {
        "punto": "Se {calificacion} la sentencia recurrida, dictada por {responsable}.",
        "calif": ("revoca", "confirma"),
        "notif": ("Notifíquese; publíquese y anótese en el libro de control de "
                  "este tribunal, hágase la captura correspondiente en el "
                  "Sistema Integral de Seguimiento de Expedientes, con "
                  "testimonio de esta resolución vuelvan los autos a su lugar "
                  "de origen y, en su oportunidad archívese como asunto "
                  "concluido."),
    },
    "amparo_directo": {
        "punto": None,          # lleva la fórmula de la Justicia de la Unión
        "notif": ("Notifíquese; publíquese y anótese en el libro de control de "
                  "este tribunal, hágase la captura correspondiente en el "
                  "Sistema Integral de Seguimiento de Expedientes, con "
                  "testimonio de esta resolución vuelvan los autos a su lugar "
                  "de origen y, en su oportunidad archívese como asunto "
                  "concluido."),
    },
}


# LO QUE TECLEA EL SECRETARIO TAMBIÉN SE COMPONE. En el proyecto 382/2024 el
# resolutivo salió diciendo «contra el acto que reclamó de la Junta especial 50
# de la federal de arbitraje en el estado de querétaro.,, precisado en el primer
# resultando». Tres defectos en una línea, y los tres del mismo origen: el campo
# se copiaba verbatim.
#
#   · el punto final que escribió el usuario, seguido de la coma de la
#     plantilla, da «.,» —y con la segunda coma de la frase, «.,,»—;
#   · las minúsculas, que en un resolutivo se leen como descuido;
#   · y el nombre convive en el mismo documento con la grafía correcta que el
#     modelo sacó del laudo, «Junta Especial Número Cincuenta de la Federal de
#     Conciliación y Arbitraje», así que el documento se contradice a sí mismo.
#
# No se cambia lo que el secretario escribió —eso es suyo y puede tener razones
# para nombrarla así—: se le quita la puntuación final y se le arreglan las
# mayúsculas si vino todo en minúsculas. Nada más.
_CONECTIVAS = {"de", "del", "la", "las", "el", "los", "y", "en", "e", "al"}


def _sin_articulo(x: str) -> str:
    """«el Juzgado Segundo» → «Juzgado Segundo». Lo pone la plantilla.

    EL CONTRATO ESTABA ESCRITO Y NO IMPLEMENTADO: el comentario de `_datos_bk`
    dice «se entrega el nombre limpio y la plantilla lo enmarca», y las
    fórmulas del banco dicen «dictado por el {responsable}». Pero nada quitaba
    el artículo, así que en cuanto el secretario teclea «el Juzgado Segundo de
    Distrito» en el formulario —que es como se dice— sale «por el el Juzgado».
    No pasaba con la autoridad LEÍDA del acto, que viene sin artículo; pasa con
    la que se escribe a mano, que es la que existe para corregir la leída.
    """
    return re.sub(r"^\s*(?:el|la|los|las)\s+", "", x or "", flags=re.I)


def _normalizar_autoridad(nombre: str) -> str:
    n = " ".join((nombre or "").split()).strip(" \t.,;:")
    if not n:
        return ""
    # Si trae mayúsculas propias, se respeta tal cual: el secretario sabe cómo
    # se llama la autoridad de su expediente mejor que yo.
    if any(c.isupper() for c in n[1:]):
        return n
    # LOS NÚMEROS ROMANOS NO SE CAPITALIZAN. `"II".capitalize()` devuelve «Ii»,
    # y el resolutivo salió diciendo «la Sala Regional del Centro Ii». Se
    # reconocen y se dejan como están.
    _romano = re.compile(r"^[IVXLCDM]{1,7}$")
    partes = []
    for i, w in enumerate(n.split()):
        limpio = w.strip(".,;:()")
        if _romano.match(limpio.upper()) and limpio.upper() == limpio:
            partes.append(w)                       # ya viene en versales
        elif _romano.match(limpio.upper()) and len(limpio) > 1:
            partes.append(w.upper())               # «ii» → «II»
        elif i and w.lower() in _CONECTIVAS:
            partes.append(w)
        else:
            partes.append(w.capitalize())
    return " ".join(partes)


# CONTRA QUÉ SE RECURRE, en las palabras del oficio. El banco lo midió: «del
# acuerdo que desechó la demanda de amparo», «de un auto dictado en un juicio de
# amparo indirecto». Si no se puede decir con precisión, se dice lo genérico
# —que es cierto— en vez de dejar un hueco: un considerando de competencia con
# un agujero en mitad de la frase no se puede leer en sesión.
_GENERICO_ACTO = {
    "queja": "del auto recurrido",
    "amparo_revision": "de la sentencia recurrida",
    "revision_fiscal": "de la sentencia recurrida",
    "amparo_directo": "de la sentencia reclamada",
}


def _descripcion_del_acto(datos: dict, tipo: str) -> str:
    d = " ".join(str(datos.get("descripcion_acto") or "").split()).strip(" .,;")
    if d:
        return d if d.lower().startswith(("del ", "de ", "de la ")) else f"del {d}"
    return _GENERICO_ACTO.get(str(tipo or "").strip().lower(), "del acto recurrido")


def _de_la(nombre: str) -> str:
    """«de la Sala Regional…», «del Tribunal Unitario…».

    El resolutivo decía «contra el acto que reclamó de el Tribunal Unitario
    Agrario», porque la plantilla ponía «de » delante de lo que `_con_articulo`
    devolvía con su artículo. «De el» no es español: se contrae.
    """
    n = _con_articulo(nombre)
    if not n:
        return ""
    if n.lower().startswith("el "):
        return "del " + n[3:]
    return "de " + n


# EL ARTÍCULO CONTRAÍDO. «contra el acto reclamado a el Director» no es
# español, y sale en cuanto una plantilla escribe la preposición y otra función
# pone el artículo: ninguna de las dos ve a la otra. Se arregla al final, sobre
# el texto ya armado, que es donde las dos se juntan.
def _contraer(t: str) -> str:
    t = re.sub(r"\ba\s+el\b(?!\s+que\b)", "al", t)
    return re.sub(r"\bde\s+el\b(?!\s+que\b)", "del", t)


def _con_articulo(nombre: str) -> str:
    """«Primera Sala Civil…» → «la Primera Sala Civil…».

    Sin esto el resolutivo dice «reclamó de Primera Sala Civil», que no es
    español. El artículo se elige por la primera palabra, y si ya viene con él
    no se duplica.
    """
    n = _normalizar_autoridad(nombre)
    if not n:
        return ""
    if re.match(r"^(?:el|la|los|las)\s", n, re.I):
        return n
    primera = n.split()[0].lower()
    femeninas = ("sala", "junta", "primera", "segunda", "tercera", "cuarta",
                 "quinta", "sexta", "séptima", "octava", "novena", "décima",
                 "autoridad", "comisión", "procuraduría", "secretaría",
                 "dirección", "delegación", "subdelegación")
    return f"{'la' if primera in femeninas else 'el'} {n}"


HUECO = "*********"

_PLURAL = {"fundado": "fundados", "infundado": "infundados",
           "inoperante": "inoperantes", "ineficaz": "ineficaces",
           # INNECESARIO NO ES UNA CALIFICACIÓN DEL PLANTEAMIENTO. No se dice
           # que sea infundado —eso sería contestarlo— sino que no hace falta
           # entrar: queda sin materia porque el principal ya resolvió el
           # asunto. Va aquí para que la calificativa del rótulo lo diga.
           "innecesario": "innecesarios de estudiar"}


def _calificacion_plural(cs: list) -> str:
    """«fundados», «en parte fundados y en parte inoperantes»…

    La calificativa va PEGADA al rótulo del Estudio en 17 de 26 engroses:
    «SEXTO. Estudio. Los conceptos de violación son infundados.»
    """
    limpios = [c for c in cs if c in _PLURAL]
    if not limpios:
        return ""
    unicos = []
    for c in limpios:
        if c not in unicos:
            unicos.append(c)
    if len(unicos) == 1:
        return _PLURAL[unicos[0]]
    if len(unicos) == 2:
        return f"en parte {_PLURAL[unicos[0]]} y en parte {_PLURAL[unicos[1]]}"
    return ", ".join(_PLURAL[c] for c in unicos[:-1]) + f" y {_PLURAL[unicos[-1]]}"


# ═══ LAS MARCAS DE CITA DEL RESUMEN ════════════════════════════════════════
# Las fases 1-3 anotan de dónde sale cada afirmación del resumen del acto
# reclamado: «[[p.82 §3]]». En la plantilla se convertían en nota al pie; el
# generador nuevo no las tocaba y salían LITERALES en el cuerpo —quince en el
# proyecto que leyó David—. Una referencia visible entre corchetes dobles no es
# una cita rota por descuido: es el andamio del redactor asomando en el papel.
# UN CORCHETE DE CIERRE O DOS. El modelo escribió «[[p.38 §3-4; p.39 §1]» con
# uno solo y la marca sobrevivió entera en el papel. Exigir la forma perfecta
# de algo que escribe un modelo es garantizar que un día no case.
_MARCA_CITA = re.compile(r"\s*\[\[([^\[\]]{2,120})\]\]?")
_UNA_CITA = re.compile(
    r"p{1,2}\.?\s*(\d{1,4})(?:\s*[-–]\s*\d{1,4})?"
    r"(?:\s*§+\s*(\d{1,3}(?:\s*[-–]\s*\d{1,3})?))?", re.I)

# Tres por párrafo. Veintidós llamadas apiladas en el último punto es lo que
# pasa cuando el modelo devuelve el apartado entero sin saltos de línea.
MAX_CITAS_POR_PARRAFO = 3

# Y UN TOPE TOTAL. Medido contra los engroses reales: el ARC 448-2025 firmado
# lleva TRES notas al pie y el proyecto generado llevaba VEINTIOCHO. Una nota
# por cada afirmación del resumen no es rigor, es ruido: el secretario anota
# las que sostienen lo que se discute, no todas las que podría. Seis deja
# margen sobre su media sin convertir el pie en un segundo documento.
MAX_NOTAS_DEL_RESUMEN = 6


def _citas_de(marca: str) -> list:
    fuera = []
    for trozo in re.split(r"[;,]", marca):
        m = _UNA_CITA.search(trozo)
        if m:
            fuera.append((m.group(1), m.group(2) or ""))
    return fuera


def _texto_de_nota(pagina: str, parrafo: str) -> str:
    if not parrafo:
        return f"Cfr. página {pagina} de la sentencia reclamada."
    p = re.sub(r"\s*[-–]\s*", " a ", parrafo)
    return (f"Cfr. página {pagina}, párrafos {p}, de la sentencia reclamada."
            if " a " in p else
            f"Cfr. página {pagina}, párrafo {p}, de la sentencia reclamada.")


# ═══ EL ARTÍCULO CITADO, CON SU TEXTO AL PIE ═══════════════════════════════
# David: «cuando el redactor cita artículos (de cualquier fuente) debería citar
# su contenido textual a pie de página; eso incrementará dramáticamente el
# valor argumentativo del proyecto». Tiene razón y es barato: el precepto ya
# está en el acervo, palabra por palabra. Quien revisa deja de tener que ir a
# buscarlo, y quien firma ve de un vistazo si el artículo dice lo que se le
# atribuye —que es donde se cuelan los errores que nadie detecta—.
_RX_ARTICULO_CITADO = re.compile(
    r"art[íi]culos?\s+(\d{1,4})(?:\s*(?:bis|ter|qu[áa]ter))?"
    r"(?:[^.;]{0,90}?(c[óo]digo|ley|constituci[óo]n|reglamento)[^.;,]{0,60})?",
    re.I)


# Las palabras que no distinguen una ley de otra.
_VACIAS_LEY = {"de", "del", "la", "el", "los", "las", "y", "en", "para", "por",
               "sobre", "estado", "estados", "unidos", "nacional", "general"}
_RX_NOMBRA_LEY = re.compile(
    r"constituci[óo]n|constitucional|c[óo]digo|\bley\b|reglamento|convenci[óo]n|"
    r"pacto|tratado|contrato\s+colectivo|condiciones\s+generales", re.I)


def _sin_tildes(x: str) -> str:
    import unicodedata
    x = unicodedata.normalize("NFKD", x or "")
    return "".join(c for c in x if not unicodedata.combining(c))


def _norma_del_texto(frag: str, num: str, normas: list):
    """El precepto del acervo que se está citando, si lo hay.

    Se exige que coincidan el NÚMERO y, cuando el texto nombra una ley, que la
    fuente comparta palabras con ella: el «artículo 296» del código civil de
    Querétaro y el «artículo 296» de otro cuerpo no son el mismo, y poner al
    pie el texto equivocado es peor que no poner nada.
    """
    ley = _sin_tildes((frag or "").lower())
    # ¿La cita nombra una ley, o dice «el artículo 17» a secas?
    nombra_ley = bool(_RX_NOMBRA_LEY.search(frag or ""))
    mejor, puntos = None, -99
    for n in (normas or []):
        if str(n.get("articulo", "")).strip() != str(num):
            continue
        # El acervo llama al campo `cuerpo_legal`; sólo algunas fuentes usan
        # `fuente`. Leer una sola de las dos daba CERO coincidencias y ninguna
        # nota de artículo salía, sin que nada avisara.
        fuente = _sin_tildes(str(n.get("cuerpo_legal") or n.get("fuente") or "").lower())
        suyas = {w for w in re.findall(r"[a-z]{4,}", fuente) if w not in _VACIAS_LEY}
        acierta = len([w for w in suyas if w in ley])
        # SE PENALIZA LO QUE SOBRA, igual que al traer los artículos por número.
        # Contando sólo aciertos, «Ley Federal de Responsabilidad Patrimonial
        # del Estado» empata con cualquier cosa que diga «Estado».
        p = acierta - len(suyas - {w for w in suyas if w in ley})
        if p > puntos:
            mejor, puntos = n, p
    # Y LA REGLA QUE FALTABA, QUE ES LA QUE COSTÓ UN PROYECTO. Antes, si ninguna
    # ley coincidía, esta función se quedaba con la PRIMERA norma que tuviera
    # ese número —`mejor is None and p == 0`—, viniera de donde viniera. Así el
    # documento transcribió, DENTRO DE COMILLAS y presentándolo como el artículo
    # 17 de la Constitución, el artículo 17 de la Ley Federal de Responsabilidad
    # Patrimonial del Estado: «Las resoluciones que se dicten con motivo de las
    # reclamaciones deberán contener… relación de causalidad entre el
    # funcionamiento del servicio público…». La prosa del modelo era correcta;
    # lo que mentía era la transcripción que yo le pegaba debajo.
    #
    # Si la cita nombra una ley y ninguna norma del material es de esa ley, NO
    # SE TRANSCRIBE NADA. Un artículo sin su texto se queda sin nota al pie y
    # quien firma lo comprueba a mano; un artículo con el texto de otra ley se
    # firma sin comprobar, y eso es lo que no se perdona.
    # Y SIN NOMBRE DE LEY TAMPOCO SE ADIVINA. El prompt exige desde hace
    # semanas nombrar la ley en la misma frase que el número —«el artículo 296
    # del Código Civil del Estado de Querétaro», nunca «el 296» a secas—; si
    # aun así llega pelado, elegir por él es apostar. Se queda sin nota y quien
    # firma lo comprueba, que es exactamente lo que la nota existe para
    # ahorrarle cuando SÍ se puede saber.
    if mejor is None or puntos < 1:
        return None
    return mejor


# CUÁNTO SE TRANSCRIBE DE UN ARTÍCULO. Medido en el proyecto de la queja civil
# 233/2025: el artículo 107 CONSTITUCIONAL salió transcrito ENTERO, 1,924
# palabras, con sus dieciocho fracciones y sus incisos, cuando lo que se
# discutía era una garantía de suspensión. Eso solo se llevaba una quinta parte
# del documento, disparaba la medida de transcripción por encima del engrose y
# generaba media docena de «pasajes duplicados» que eran trozos del mismo
# artículo repetidos.
#
# El secretario no hace eso: transcribe «en la parte conducente». Y el corpus lo
# dice con esas palabras —está en la dispensa de los cinco engroses—.
MAX_PALABRAS_PRECEPTO = 180

# A partir de aquí, la tesis se lee al pie. Ochenta palabras son unos cinco
# renglones: lo que cabe sin romper la lectura del razonamiento.
MAX_PALABRAS_TESIS_CUERPO = 80


def _en_lo_conducente(cuerpo: str, fraccion: str = "") -> str:
    """El artículo, o su parte conducente si es largo.

    Primero se intenta quedarse con la FRACCIÓN que el párrafo citó, que es lo
    que el secretario haría. Si no consta cuál, se corta en frontera de frase y
    se dice «en lo conducente», que es como se anuncia una transcripción
    parcial: fingir que es íntegra cuando no lo es sería peor que cortarla.
    """
    pal = cuerpo.split()
    if len(pal) <= MAX_PALABRAS_PRECEPTO:
        return cuerpo
    if fraccion:
        # «IV.» o «fracción IV» dentro del texto del artículo.
        m = re.search(rf"(?:^|[;.]\s*){re.escape(fraccion)}\.\s", cuerpo)
        if m:
            resto = cuerpo[m.start():]
            fin = re.search(r"[;.]\s+[IVXLC]+\.\s", resto[3:])
            trozo = resto[:fin.start() + 3] if fin else resto
            if 10 <= len(trozo.split()) <= MAX_PALABRAS_PRECEPTO * 2:
                # EL RÓTULO SE CORTABA EN LA ABREVIATURA. `split(".")[0]`
                # sobre «Art. 104.- Los Tribunales…» devuelve «Art», y el
                # precepto salía encabezado por «Art. […]», que no dice qué
                # artículo es. Se toma el rótulo entero, con su número.
                _r = re.match(r"\s*(Art[íi]culos?|Art)\.?\s*(\d{1,3}\s*"
                              r"(?:bis|ter)?)", cuerpo, re.I)
                cab = (f"Artículo {_r.group(2).strip()}" if _r
                       else cuerpo.split(".")[0])
                # EL PUNTO Y COMA VIAJABA PEGADO AL TROZO y salía «[…] ; III.
                # De los recursos…», con el espacio delante del signo. Se
                # quitan los signos de puntuación con que empieza el corte:
                # pertenecen a la frase anterior, que es la que se elidió.
                trozo = trozo.strip().lstrip(".;,: ")
                # Y EL RÓTULO NO SE ESCRIBE DOS VECES. El acervo guarda el
                # texto empezando por «Art. 104.-», así que anteponerle el
                # encabezado producía «Artículo 104. Art. […]».
                if re.match(r"^Art[íi]?c?u?l?o?\.?\s*\d", trozo):
                    return f"[…] {trozo}"
                return f"{cab}. […] {trozo}"
    corte = " ".join(pal[:MAX_PALABRAS_PRECEPTO])
    ult = max(corte.rfind(". "), corte.rfind("; "))
    if ult > len(corte) * 0.5:
        corte = corte[:ult + 1]
    # UN ROMANO SUELTO AL FINAL no es una fracción: es media fracción cortada.
    corte = re.sub(r"\s+[IVXLC]{1,6}\.?\s*$", "", corte.rstrip(" ;,."))
    return corte.rstrip(" ;,.") + " […]"


def escribir_precepto(doc, texto_articulo: str, ley: str, num: str,
                      fraccion: str = ""):
    """El artículo transcrito, como lo hace el secretario.

    David: «Cuando citamos un artículo hay que hacerlo con interlineado uno y
    con sangría en todo el artículo… No se dice en el mismo párrafo, se abren
    dos puntos, se cita textualmente el artículo y luego se sigue con la
    redacción». Así queda:

        …conforme al artículo 296 del Código Civil del Estado de Querétaro,
        que dispone:
            «Artículo 296. Los alimentos han de ser…»        ← sangrado, a uno
        Como se advierte del precepto transcrito, …          ← sigue la prosa

    El precepto sale del acervo, palabra por palabra. Un artículo transcrito de
    memoria es el error que nadie detecta al revisar.
    """
    cuerpo = " ".join(str(texto_articulo or "").split())
    if not cuerpo:
        return None
    cuerpo = _en_lo_conducente(cuerpo, fraccion)
    # EL ACERVO GUARDA UNA MIGAJA DELANTE: «[Ley de Amparo | CAPÍTULO X
    # Sentencias | Disposiciones Fundamentales] Artículo 79. La autoridad…».
    # Es su índice interno, no el precepto, y transcrita queda ridícula en una
    # sentencia. Se quita, y con ella el «Artículo N.» duplicado que viene
    # detrás.
    cuerpo = re.sub(r"^\s*\[[^\]]{0,200}\]\s*", "", cuerpo)
    # Y LA GRAFÍA ABREVIADA, que es la del acervo constitucional. El patrón
    # exigía la palabra «ARTÍCULO» entera y el corpus escribe «Art. 14.-», así
    # que el encabezado duplicado sobrevivía y salía «Artículo 14. Art. 14.- A
    # ninguna ley se dará efecto retroactivo…». Se ve en la nota al pie de cada
    # precepto constitucional, que es donde quien firma va a comprobar.
    for _ in range(2):
        cuerpo = re.sub(r"^\s*ART(?:[ÍI]CULO)?\.?\s*\d+[^.]{0,14}\.?\s*[-–]?\s*",
                        "", cuerpo, flags=re.I)
    q = doc.add_paragraph()
    r = q.add_run(f"«Artículo {num}. {cuerpo}»")
    _fmt(q, sangria=False, tamano=TAMANO_CITA,
         interlineado=INTERLINEADO_CITA)
    q.paragraph_format.left_indent = SANGRIA_CITA
    q.paragraph_format.right_indent = Cm(0.5)
    q.paragraph_format.space_before = Pt(6)
    q.paragraph_format.space_after = Pt(6)
    q.paragraph_format.keep_with_next = False
    return q


def _preceptos_del_parrafo(texto: str, normas: list) -> list:
    """[(número, norma)] de los artículos que este párrafo cita y tenemos."""
    fuera, vistos = [], set()
    for m in _RX_ARTICULO_CITADO.finditer(texto or ""):
        num = m.group(1)
        if num in vistos:
            continue
        n = _norma_del_texto(m.group(0), num, normas)
        if n and str(n.get("texto") or "").strip():
            vistos.add(num)
            fuera.append((num, n))
    return fuera


def notas_de_articulos(doc, p, texto: str, normas: list, notas: list) -> int:
    """Cuelga del párrafo el texto de los artículos que cita. Devuelve cuántos."""
    if not normas:
        return 0
    puestos = 0
    for m in _RX_ARTICULO_CITADO.finditer(texto or ""):
        if puestos >= MAX_ARTICULOS_POR_PARRAFO:
            break
        num = m.group(1)
        n = _norma_del_texto(m.group(0), num, normas)
        if not n:
            continue
        cuerpo = " ".join(str(n.get("texto", "")).split())[:900]
        if not cuerpo:
            continue
        _ley = n.get("cuerpo_legal") or n.get("fuente") or ""
        # El texto del acervo ya suele venir con «Artículo N.» delante.
        cuerpo = re.sub(r"^\s*\[[^\]]{0,200}\]\s*", "", cuerpo)
        cuerpo = re.sub(r"^\s*ART[ÍI]CULO\s+\d+[^.]{0,12}\.?\s*", "", cuerpo,
                        flags=re.I)
        pie = f"«Artículo {num}. {cuerpo}» — {_ley}".strip()
        if pie in notas:
            continue
        notas.append(pie)
        _run_llamada(p, len(notas))
        puestos += 1
    return puestos


MAX_ARTICULOS_POR_PARRAFO = 1
MAX_NOTAS_DE_ARTICULOS = 8


def parrafo_con_citas(doc, texto: str, notas: list):
    """Escribe el párrafo y baja sus marcas a notas al pie."""
    citas = [c for m in _MARCA_CITA.finditer(texto) for c in _citas_de(m.group(1))]
    limpio = _MARCA_CITA.sub("", texto).strip()
    if not limpio:
        return None
    p = parrafo(doc, limpio)
    # UNA LLAMADA POR PÁRRAFO, Y NUNCA DOS PEGADAS. Dos referencias seguidas sin
    # texto en medio se leen como un solo número: las notas 3 y 4 aparecían como
    # «34». Lo vio David. Y la MISMA página se citaba cinco veces seguidas
    # porque el modelo repite la marca párrafo tras párrafo; una nota repetida
    # no aporta nada y ensucia el pie.
    if not citas:
        return p
    pagina, parr = citas[0]
    texto_nota = _texto_de_nota(pagina, parr)
    if texto_nota in notas:
        return p                       # ya está al pie: no se repite
    if len([x for x in notas if x.startswith("Cfr.")]) >= MAX_NOTAS_DEL_RESUMEN:
        return p
    notas.append(texto_nota)
    _run_llamada(p, len(notas))
    return p


def _subtitulo(doc, texto: str):
    """Subtítulo en negrita SIN ordinal, como los del Estudio.

    Medido: «Sentencia reclamada», «Conceptos de violación», «Solución»,
    «Conclusión». No llevan número: no son considerandos, son las partes de
    uno solo.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(texto)
    r.bold = True
    _fmt(p, sangria=True)
    p.paragraph_format.keep_with_next = True
    return p


# ANDAMIO DEL MODELO QUE NO DEBE LLEGAR AL PAPEL. En la queja salió, dentro de
# una jurisprudencia, «la jurisprudencia 2a./J. 58/2010,[NOTA 2] emitida por la
# Segunda Sala…». Nadie le enseñó esa etiqueta: la inventó imitando una
# convención de nota al pie, y el documento YA lleva notas de verdad —las
# inserta el compositor con su XML— así que ese corchete es un marcador
# huérfano que el secretario tiene que borrar a mano.
#
# El filtro es DELIBERADAMENTE ESTRECHO. Un limpiador de corchetes a secas se
# llevaría «[sic]», «[…]» y los incisos «[a]» de una transcripción, que sí son
# del documento. Sólo caen las etiquetas editoriales con su palabra clave.
_RX_ANDAMIO = re.compile(
    r"\s*\[\s*(?:NOTA|NOTAS|FOOTNOTE|CITA|REF|PIE)\s*[:#]?\s*\d*\s*\]", re.I)


def sin_andamio(texto: str) -> str:
    """Quita los marcadores que el modelo escribe para sí mismo."""
    t = _RX_ANDAMIO.sub("", texto or "")
    # El corchete se come el espacio de delante; si iba pegado a una coma,
    # queda «58/2010, emitida», que es lo que debía decir.
    return re.sub(r"\s+([,.;:])", r"\1", t)


def _sin_eco(texto: str, cuerpo_tesis: str) -> str:
    """Quita del párrafo las frases que repiten la tesis ya transcrita.

    Pedirlo en el prompt no basta: se dijo en el cuerpo y al final, y aun así
    el modelo vuelve a contar la tesis en una de cada dos corridas. Se hace
    aquí, que es donde no falla, y QUIRÚRGICAMENTE: se borran las frases que
    repiten y se conserva lo que aplica el criterio al caso, que es lo único
    que el lector no tiene ya delante.
    """
    if not (cuerpo_tesis or "").strip():
        return texto
    voc = [set(_norm_palabras(f)) for f in _frases_de(cuerpo_tesis)]
    if not voc:
        return texto
    quedan = []
    for frase in re.split(r"(?<=[.])\s+", texto or ""):
        p = set(_norm_palabras(frase))
        if len(p) >= 8 and any(len(p & v) / max(1, len(p)) > 0.72 for v in voc):
            continue                       # el lector acaba de leerla arriba
        quedan.append(frase)
    return " ".join(x for x in quedan if x.strip()).strip()


# EL MISMO ECO, PERO CON LOS PRECEPTOS. El dictamen del 382/2024 v6 lo contó
# cinco veces seguidas: el modelo escribe un extracto entrecomillado del
# artículo 840 —«fracciones IV, VI y VII»— y el compositor pega debajo el
# artículo íntegro; luego el 47, el 815, el 784 y el 48, todos dos veces. El
# lector se encuentra lo mismo dos renglones después y el proyecto engorda de
# paja legislativa.
#
# No se arregla en el prompt: al modelo hay que dejarle citar el trozo que le
# interesa, porque es su razonamiento. Se arregla al componer, y en este orden:
# si el documento va a transcribir el artículo entero justo debajo, el extracto
# entrecomillado de arriba SOBRA y se borra.
_RX_ENTRECOMILLADO = re.compile(r"[«\"“]([^»\"”]{40,1800})[»\"”]")


def _sin_extracto_repetido(texto: str, preceptos: list) -> str:
    """Quita del párrafo los entrecomillados del artículo que se va a transcribir."""
    if not preceptos or not (texto or "").strip():
        return texto
    vocablos = []
    for _num, n in preceptos:
        pal = set(_norm_palabras(str(n.get("texto") or "")))
        if len(pal) >= 12:
            vocablos.append(pal)
    if not vocablos:
        return texto
    fuera = texto
    for m in list(_RX_ENTRECOMILLADO.finditer(texto)):
        p = set(_norm_palabras(m.group(1)))
        if len(p) < 8:
            continue
        # El extracto está CONTENIDO en el artículo: casi todas sus palabras
        # aparecen en el texto que se va a transcribir. Ese es el eco.
        if any(len(p & v) / max(1, len(p)) > 0.80 for v in vocablos):
            fuera = fuera.replace(m.group(0), "")
    if fuera == texto:
        return texto
    # LA FRASE QUE INTRODUCÍA EL EXTRACTO SE QUEDA COJA. Al borrar el
    # entrecomillado, «…, que dispone: «…».» se convierte en «…, que dispone.»,
    # que es peor que la repetición: no dice nada y se nota. Se poda el
    # introductor entero, esté al final o en medio del párrafo, y la frase se
    # cierra donde acababa el sujeto.
    fuera = re.sub(
        r"[,;]?\s*(?:en\s+la\s+parte\s+conducente[,\s]*)?"
        r"(?:el\s+cual|la\s+cual|que|y\s+que|donde)?\s*"
        r"(?:dispone|establece|se[ñn]ala|prev[ée]|dice|reza|indica|prescribe)"
        r"(?:\s+lo\s+siguiente)?\s*:?\s*(?=[.;]|$)",
        "", fuera, flags=re.I)
    fuera = re.sub(r"\s*[,:;]\s*(?=[.;])", "", fuera)
    fuera = re.sub(r"\s{2,}", " ", fuera).strip()
    fuera = re.sub(r"\s+\.", ".", fuera)
    return fuera.strip(" ,;:")


def _frases_de(t: str) -> list:
    return [f for f in re.split(r"(?<=[.])\s+", t or "") if len(f.split()) >= 8]


def _norm_palabras(t: str) -> list:
    import unicodedata
    t = unicodedata.normalize("NFKD", (t or "").lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    return re.sub(r"[^a-z0-9 ]+", " ", t).split()


# ═══ LOS EFECTOS LOS ESCRIBE EL MODELO, NO LA PLANTILLA ════════════════════
# Auditando el proyecto salió el defecto más caro de todos: el modelo había
# redactado SIETE efectos concretos —«el quinto efecto consiste en que obtenga
# información oficial sobre el importe, periodicidad y condiciones de pago de
# la pensión por orfandad»— y el compositor los tiraba para poner en su lugar
# «dicte otra en la que atienda los lineamientos de esta ejecutoria».
#
# Es justo el efecto que NO se puede ejecutar sin interpretarlo, que es lo que
# el corpus prohíbe. La responsable recibe la ejecutoria y no sabe qué hacer.
# Se usan los del modelo; la fórmula de plantilla queda de respaldo para cuando
# no los haya escrito.
_RX_INICIO_EFECTOS = re.compile(
    r"(?:^|\s)(?:por\s+tanto,?\s+)?la\s+concesi[óo]n\s+del\s+amparo\s+debe\s+"
    r"producir\s+los\s+efectos|^\s*los\s+efectos\s+de\s+la\s+concesi[óo]n|"
    r"^\s*el\s+primer\s+efecto\s+consiste", re.I)
_RX_UN_EFECTO = re.compile(
    r"^\s*(?:el\s+)?(?:primer|segundo|tercer|cuarto|quinto|sexto|s[ée]ptimo|"
    r"octavo)\s+efecto\b", re.I)


def partir_efectos(estudio: list) -> tuple:
    """(estudio sin los efectos, párrafos de efectos). Vacío si no los escribió."""
    if not estudio:
        return list(estudio or []), []
    corte = None
    for i, t in enumerate(estudio):
        if _RX_INICIO_EFECTOS.search(t or "") or _RX_UN_EFECTO.match(t or ""):
            corte = i
            break
    if corte is None:
        return list(estudio), []
    cuerpo = list(estudio[:corte])
    efectos = [x for x in estudio[corte:] if (x or "").strip()]
    # El párrafo final de cierre —«por lo que procede conceder el amparo…»— no
    # es un efecto: cierra el estudio y ya lo pone el resolutivo.
    while efectos and re.search(r"procede\s+conceder\s+el\s+amparo",
                                efectos[-1], re.I) and len(efectos) > 1:
        efectos.pop()
    return cuerpo, efectos


# LA PREGUNTA QUE FIJA LA CUESTIÓN NO SE DESCARTA POR CORTA. El compositor
# tira los párrafos de menos de seis palabras —restos de un corte, coletillas
# sueltas— y «1. ¿Era aplicable ese criterio?» son cinco. Habría pedido la
# pregunta expresa y la habría borrado acto seguido.
_RX_ES_PREGUNTA = re.compile(r"^\s*(?:\d{1,2}[.)]\s*)?¿.{10,}\?\s*$", re.S)


def _es_pregunta(t: str) -> bool:
    return bool(_RX_ES_PREGUNTA.match((t or "").strip()))


def _escribir_estudio(doc, estudio, tesis, notas, normas=None) -> int:
    """Los párrafos del estudio, con sus citas rehechas desde el acervo."""
    citadas = 0
    ultima_tesis = None
    transcritos = set()
    transcritas_tesis = set()
    for t in (estudio or []):
        t = (t or "").strip()
        if not t:
            continue
        # El encabezado ordinal que el modelo se pone a sí mismo sobra: el
        # ordinal lo calcula el compositor y ya está escrito arriba.
        t = re.sub(r"^(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|"
                   r"S[ÉE]PTIMO|OCTAVO|NOVENO)\.\s*(?:Estudio(?:\s+de\s+fondo)?\.\s*)?",
                   "", t)
        t = re.sub(r"^Los\s+(?:conceptos\s+de\s+violaci[óo]n|agravios)\s+son\s+"
                   r"[^.]{3,60}\.\s*", "", t)
        if not t.strip():
            continue
        hallada, m_r = tesis_del_rubro(t, tesis or [])
        # Y UNA TESIS TAMBIÉN. La 169606 se transcribió dos veces en el mismo
        # considerando —una en el marco y otra al contestar el concepto—: el
        # texto íntegro repetido no aporta y alarga la sentencia sin decir nada.
        if hallada and str(hallada.get("registro") or "") in transcritas_tesis:
            hallada = None
        if hallada and m_r and citadas < MAX_CITAS_DOCUMENTO:
            transcritas_tesis.add(str(hallada.get("registro") or ""))
            antes = _RX_COLA_ANUNCIO.sub("", t[:m_r.start()].rstrip(" ,;:"))
            cola = t[m_r.end():].lstrip(" ,;:.")
            escribir_cita(doc, hallada, antes.rstrip(" ,;:"), notas)
            citadas += 1
            cola = _sin_eco(cola, hallada.get("texto") or "")
            if len(cola.split()) > 6 or _es_pregunta(cola):
                parrafo_con_citas(doc, cola, notas)
            ultima_tesis = hallada
            continue
        if ultima_tesis is not None:
            t = _sin_eco(t, ultima_tesis.get("texto") or "")
            ultima_tesis = None
            if len(t.split()) < 6 and not _es_pregunta(t):
                continue
        # Se decide ANTES de escribir qué artículos van a transcribirse, para
        # poder quitar del párrafo el extracto que quedaría repetido debajo.
        _del_parrafo = _preceptos_del_parrafo(t, normas)[:MAX_ARTICULOS_POR_PARRAFO]
        _preceptos = [(n_, x) for n_, x in _del_parrafo if n_ not in transcritos]
        # EL ECO SOBREVIVÍA CUANDO EL ARTÍCULO YA ESTABA TRANSCRITO. La poda
        # miraba sólo los preceptos que este párrafo va a transcribir DEBAJO;
        # si el compositor ya lo había puesto páginas antes, `_preceptos`
        # quedaba vacío y la copia que el modelo escribió aquí se quedaba.
        #
        # Medido en el ARA 17/2025 generado: el artículo 76 de la Ley de Amparo
        # aparece dos veces, una en el bloque del compositor —entre comillas
        # angulares— y otra en el cuerpo, escrita por el modelo. Es el defecto
        # que el detector de duplicación marcaba y la regla del prompt no
        # bastaba para evitar, porque no es del modelo solo: es de los dos.
        t = _sin_extracto_repetido(t, _del_parrafo)
        if len(t.split()) < 6 and not _es_pregunta(t):
            continue
        p_ = parrafo_con_citas(doc, t, notas)
        # EL PRECEPTO SE TRANSCRIBE, NO SE RESUME. Va en bloque aparte, con
        # sangría y a un espacio, detrás del párrafo que lo anuncia.
        if p_ is not None:
            for num, n_ in _preceptos:
                # UN ARTÍCULO SE TRANSCRIBE UNA VEZ. La clave era (número,
                # ley) y el 48 salió DOS veces porque llegó por dos caminos con
                # el nombre de la ley escrito distinto —«Artículo 48.-» y
                # «Artículo 48.»—. Al lector le da igual de dónde vino: lee lo
                # mismo dos veces seguidas.
                if num in transcritos:
                    continue
                transcritos.add(num)
                # LA FRACCIÓN QUE EL PÁRRAFO CITÓ. Si el texto que anuncia
                # el precepto dice «artículo 107, fracción X», se transcribe
                # ESA fracción y no el artículo entero: es lo que el
                # secretario hace y lo que hace legible el bloque.
                _fr = re.search(
                    rf"art[íi]culo\s+{re.escape(str(num))}\s*,?\s*"
                    rf"fracci[óo]n\s+([IVXLC]+)", t, re.I)
                escribir_precepto(doc, n_.get("texto"),
                                  n_.get("cuerpo_legal") or n_.get("fuente") or "",
                                  num, _fr.group(1) if _fr else "")
    return citadas


# ═══ EL ESQUELETO DE CADA TIPO, MEDIDO ═════════════════════════════════════
# amparo directo civil 26 · administrativo 16 · queja 20 (y 153 de recuento) ·
# revisión civil 31 · administrativa 16 · fiscal 28. La regla que vale para
# TODOS: el resumen de lo recurrido NO es considerando; el considerando que
# lleva su nombre es la DISPENSA de transcribirlo.
#
# Lo que cambia de un tipo a otro y rompería una plantilla única:
#   · Los RECURSOS no tienen «Existencia del acto reclamado»: es del amparo.
#   · En la QUEJA el cómputo va en PROSA, sin tabla, y la procedencia lleva
#     UNA nota al pie con el artículo 97 de la Ley de Amparo.
#   · El secretario escribe «Trascripción» sin la n: 104 veces contra 12.
#   · El estudio no tiene ordinal fijo: es el último, y su número sale de
#     cuántos apartados le preceden.
ESQUELETO = {
    # ═══════════════════════════════════════════════════════════════════
    # LA TABLA VA EN LOS CUATRO, Y ES UNA DECISIÓN DE PRODUCTO
    # ═══════════════════════════════════════════════════════════════════
    # Yo la había quitado de tres tipos porque el corpus casi no la usa —tabla
    # en 1 de 45 amparos directos, 0 de 21 quejas—. David: «cometí un error al
    # pedirte que coincidieran con mis adelantos. Un plus que tenía el generador
    # era la tabla. Ésa quiero conservarla para todos».
    #
    # Tiene razón y el criterio es distinto del que yo aplicaba: contra el
    # corpus se mide lo que HAY QUE IMITAR —los rótulos, las fórmulas, el orden,
    # el vocabulario— porque ahí el corpus es la autoridad. Pero el corpus no
    # manda sobre lo que el producto puede MEJORAR: si el secretario no dibuja
    # la tabla es porque hacerla a mano cuesta, no porque sobre. La máquina
    # tiene el calendario y la aritmética; regalarle el desglose es justamente
    # lo que ella aporta.
    #
    # La prosa sí se queda corta —«resultó oportuna, a la luz del artículo 17»—
    # porque con la tabla debajo, repetir el cómputo en palabras es decir dos
    # veces lo mismo.
    "amparo_directo": {
        "q": "conceptos de violación",
        "recurrido": "la sentencia reclamada",
        "tabla_computo": True,
        "dispensa": "Acto reclamado y {q}.",
        "legitimacion": "Legitimación y oportunidad.",
        "existencia": True,
        "sub_recurrido": "Sentencia reclamada",
        "adhesivo": "Amparo adhesivo.",
    },
    "amparo_revision": {
        "q": "agravios",
        "recurrido": "la resolución recurrida",
        "tabla_computo": True,
        "dispensa": "Resolución recurrida y {q} de la parte recurrente.",
        "legitimacion": "Legitimación y oportunidad para interponer el recurso.",
        "existencia": False,
        "sub_recurrido": "Resolución recurrida",
        "adhesivo": "Revisión adhesiva.",
    },
    "queja": {
        "q": "agravios",
        "recurrido": "el auto recurrido",
        "tabla_computo": True,
        "dispensa": "Trascripción innecesaria del auto recurrido y {q}.",
        "legitimacion": "Legitimación y oportunidad.",
        "existencia": False,
        "procedencia_propia": True,
        "sub_recurrido": "Auto recurrido",
        "adhesivo": "",
    },
    "revision_fiscal": {
        "q": "agravios",
        "recurrido": "la sentencia impugnada",
        "tabla_computo": True,
        "dispensa": "Consideraciones de la sentencia impugnada y {q}.",
        "legitimacion": "Legitimación y oportunidad.",
        "existencia": False,
        "procedencia_propia": True,
        "sub_recurrido": "Sentencia impugnada",
        "adhesivo": "",
    },
}


def esqueleto_de(tipo: str) -> dict:
    return ESQUELETO.get((tipo or "").strip().lower(),
                         ESQUELETO["amparo_directo"])


def _pagina(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.59), Cm(34.03)
    s.left_margin, s.right_margin = Cm(5), Cm(2)
    s.top_margin, s.bottom_margin = Cm(3), Cm(3)
    normal = doc.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = TAMANO
    return s


def _encabezado(doc, texto):
    for s in doc.sections:
        for parte in (s.header, s.first_page_header, s.even_page_header):
            if parte is None:
                continue
            p = parte.paragraphs[0] if parte.paragraphs else parte.add_paragraph()
            p.text = ""
            r = p.add_run(texto)
            r.bold = True
            r.font.name = FUENTE
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor.from_string(GRIS_CABECERA)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _caratula(doc, datos, tipo_asunto: str = "") -> list:
    """La ficha de identificación. Del asunto, no de ningún otro."""
    # LAS FIGURAS SON DEL TIPO. Esta lista tenía las tres del amparo directo
    # escritas a mano y se imprimía igual en los cuatro, teniendo el tipo en la
    # mano: una revisión rotulaba «QUEJOSO» a quien es recurrente y
    # «AUTORIDAD RESPONSABLE» al Juez de Distrito, que no es parte del recurso
    # sino el órgano de control cuya sentencia se revisa.
    import tipos_asunto as _ta_c
    _t = str(datos.get("tipo_asunto") or tipo_asunto or "amparo_directo")
    # LA PRIMERA LÍNEA NO SE ROTULA. El corpus escribe la clase del asunto como
    # etiqueta —«REVISIÓN FISCAL: 87/2025»—, así que anteponerle «EXPEDIENTE: »
    # la rotula dos veces; salía «EXPEDIENTE: REVISIÓN FISCAL».
    # LOS AVISOS SE DEVUELVEN, NO SE ACUMULAN EN EL MÓDULO. Una lista global
    # la comparten las peticiones que atiende el mismo worker de gunicorn: el
    # aviso de un asunto acabaría en la sentencia de otro. Es la misma clase de
    # error que la de mutar el calendario del cómputo.
    _avisos: list = []
    # UN ENCABEZADO SIN NÚMERO NO IDENTIFICA EL ASUNTO. Salía «REVISIÓN
    # FISCAL» y «RECURSO DE QUEJA CIVIL» a secas, y de ahí el proemio decía
    # «cuyo número consta en autos». Se avisa; no se inventa.
    _enc = str(datos.get("encabezado", ""))
    if _enc and not re.search(r"\b\d{1,5}\s*/\s*\d{2,4}\b", _enc):
        _avisos.append(
            f"EL ENCABEZADO NO TRAE NÚMERO DE EXPEDIENTE: «{_enc[:60]}». Sin él "
            f"el proemio no puede citar el asunto y la carátula no lo "
            f"identifica.")
    campos = [("", _enc)]
    campos += [(et, datos.get(clave, "")) for et, clave, _ob
               in _ta_c.caratula_de(_t)]
    campos += [("MAGISTRADO PONENTE", datos.get("magistrado", "")),
               ("SECRETARIA/O", datos.get("secretario", ""))]
    # LOS DATOS DE IDENTIFICACIÓN VAN A LA IZQUIERDA Y EN MAYÚSCULAS, como en
    # sus proyectos. Justificados y en caja mixta parecían prosa; son una ficha
    # y se leen de un golpe de vista.
    for etiqueta, valor in campos:
        if not valor:
            continue
        p = doc.add_paragraph()
        if etiqueta:
            r1 = p.add_run(f"{etiqueta}: ")
            r1.bold = True
        r2 = p.add_run(str(valor).upper())
        r2.bold = True
        _fmt(p, sangria=False, interlineado=1.0,
             alineacion=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    return _avisos


def _bloque_firmas(doc, datos):
    for etiqueta, quien in (("MAGISTRADO PONENTE", datos.get("magistrado", "")),
                            ("SECRETARIA/O DE TRIBUNAL", datos.get("secretario", ""))):
        if not quien:
            continue
        parrafo(doc, "", sangria=False)
        parrafo(doc, etiqueta, sangria=False, negrita=True,
                alineacion=WD_ALIGN_PARAGRAPH.CENTER, interlineado=1.0)
        parrafo(doc, str(quien).upper(), sangria=False, negrita=True,
                alineacion=WD_ALIGN_PARAGRAPH.CENTER, interlineado=1.0)


def componer(datos: dict, estructura: Estructura, computo, fecha_en_letra,
             ruta_salida: str, antecedentes=None, resumen_acto=None,
             resumen_conceptos=None, problemas=None, estudio=None,
             calificaciones=None, tesis=None, marco_escrito="",
             tipo_asunto="amparo_directo", normas=None, criterios=None) -> str:
    """Escribe el .docx entero. No hay plantilla de la que partir."""
    doc = docx.Document()
    notas: list = []
    _pagina(doc)
    _encabezado(doc, datos.get("encabezado", ""))
    # Los avisos deterministas del documento, en un solo sitio y declarados
    # ANTES del primero que los usa: la lista se llenaba en dos puntos y se
    # declaraba entre ellos, que en Python es un UnboundLocalError esperando.
    _avisos_bk: list = []

    # ═══════════════════════════════════════════════════════════════════════
    # UN SOLO EMBUDO PARA TODO LO QUE ESCRIBIÓ EL MODELO
    # ═══════════════════════════════════════════════════════════════════════
    # El meta-lenguaje se limpiaba en las fases de lectura, que es donde salió,
    # pero al documento llegan CINCO salidas de modelo por caminos distintos:
    # los resúmenes, la estructura, el estudio de fondo, el marco jurídico y la
    # propuesta. Parchear cada fase deja siempre una puerta abierta —ya van
    # tres rondas encontrando la siguiente—, así que se filtra aquí, que es por
    # donde pasa todo lo que acaba en el .docx, sea cual sea su origen.
    try:
        import meta_lenguaje as _mlc

        def _limpio(x):
            t, q = _mlc.limpiar(x or "")
            for f in q:
                _avisos_bk.append(
                    f"SE QUITÓ UNA FRASE QUE HABLABA DEL ARCHIVO, NO DEL "
                    f"ASUNTO: «{f[:200]}». Va entera aquí por si el filtro se "
                    f"equivocó y hay que devolverla.")
            return t

        estudio = _limpio(estudio) if isinstance(estudio, str) else estudio
        marco_escrito = _limpio(marco_escrito)
        if estructura is not None:
            estructura.apertura = _limpio(getattr(estructura, "apertura", ""))
            estructura.visto = _limpio(getattr(estructura, "visto", ""))
            estructura.competencia = _limpio(getattr(estructura, "competencia", ""))
            estructura.existencia = _limpio(getattr(estructura, "existencia", ""))
            estructura.procedencia = _limpio(getattr(estructura, "procedencia", ""))
            for _r in (estructura.resultandos or []):
                if isinstance(_r, dict):
                    _r["texto"] = _limpio(_r.get("texto", ""))
        if isinstance(antecedentes, list):
            antecedentes = [_limpio(x) for x in antecedentes]
        elif isinstance(antecedentes, str):
            antecedentes = _limpio(antecedentes)
    except Exception:
        pass

    # LA SUPLENCIA POR UNA MATERIA QUE NO ES LA DEL ASUNTO. Ver la nota en
    # tipos_asunto: no es un defecto de la máquina —la parte lo invocó— pero el
    # proyecto no lo decía, y quien lee no puede distinguir el disparate de la
    # parte del de la máquina sin ir al escrito. Además, esa petición hay que
    # contestarla.
    try:
        _sup = _ta.suplencia_de_otra_materia(
            " ".join(str(x) for x in [resumen_conceptos, estudio] if x),
            str(datos.get("materia") or ""))
        for _m in _sup:
            _avisos_bk.append(
                f"LA PARTE PIDE SUPLENCIA POR MATERIA {_m.upper()} y este "
                f"asunto no lo es. Está en SU escrito, no lo puso el sistema; "
                f"pero es una petición que hay que contestar, normalmente "
                f"declarándola improcedente, y el proyecto no la contesta.")
    except Exception:
        pass

    avisos_doc = _caratula(doc, datos, tipo_asunto)

    if estructura.apertura:
        parrafo(doc, estructura.apertura, sangria=True)
    if estructura.visto:
        # El rótulo lo pone la composición; el modelo lo repite igual aunque se
        # le pida que no —«V I S T O, VISTO, para resolver…»—. Se le quita.
        _v = re.sub(r"^\s*V\s*I\s*S\s*T\s*O\s*S?\s*,?\s*", "",
                    estructura.visto.strip(), flags=re.I)
        # «V I S T O S» en la revisión fiscal: 28 de 28 en el corpus. Se
        # imprimía en singular a máquina y encima se limpiaba con una regex que
        # admite la S, así que aunque el modelo acertara se le borraba: este
        # tipo era incorregible por prompt.
        # «V I S T O, Para resolver…»: el modelo escribe su trozo como si
        # empezara una frase, porque para él lo es. Detrás de la coma del
        # rótulo va minúscula, y el corpus lo escribe así en los cuatro.
        if _v[:1].isupper() and _v[:3] not in ("V I",):
            _v = _v[:1].lower() + _v[1:]
        tramos(doc, [(_ta.proemio_de(tipo_asunto)["rotulo"], {"bold": True}),
                     (_v, {})], sangria=True)

    # ═══ LA ESTRUCTURA, MEDIDA SOBRE 58 ENGROSES ═══════════════════════
    # 26 amparos directos civiles, 16 administrativos y 16 revisiones. En
    # NINGUNO existe un considerando llamado «Consideraciones de la sentencia
    # reclamada», ni «Planteamientos de la parte quejosa», ni «Marco jurídico»
    # —éste aparece 1 vez en 58 y como SUBTÍTULO dentro del Estudio—. Todo eso
    # vive DENTRO del considerando de Estudio, en subtítulos en negrita SIN
    # ordinal. El redactor los emitía como considerandos propios y por eso el
    # estudio acababa en OCTAVO, chocando con su propio encabezado.
    #
    # LOS ORDINALES SE CALCULAN, NO SE ESCRIBEN. Se numera al final la lista de
    # apartados realmente emitidos. Así el Estudio cae en QUINTO cuando no hay
    # Antecedentes —que es lo que dijo David— y en SEXTO cuando sí los hay, sin
    # que nadie tenga que decidirlo.
    def _emitir(apartados):
        """Numera y escribe. Cada apartado es (rótulo, escritor)."""
        for k, (rot, escribir_cuerpo) in enumerate(apartados):
            # SEPARACIÓN POR ESPACIADO, NO POR PÁRRAFO VACÍO. Un párrafo en
            # blanco entre cada apartado y cada cita dejaba huecos enormes en
            # el papel —22 de 173 párrafos eran aire— y además se arrastran al
            # editar. Word tiene `space_before` justo para esto.
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            r1 = p.add_run(f"{_ORDINALES[min(k, 9)]}. ")
            r1.bold = True
            r2 = p.add_run(f"{rot} ")
            r2.bold = True
            # TODO EL PROYECTO CON SANGRÍA, también el párrafo que abre cada
            # apartado: el ordinal en negrita no lo exime. Lo pidió David y es
            # lo que hace el corpus.
            _fmt(p, sangria=True)
            # EL RÓTULO SE ATA A SU CUERPO; UN CUERPO NO SE ATA A LO SIGUIENTE.
            # `keep_with_next` quedaba puesto en este párrafo SIEMPRE, y como
            # el cuerpo del apartado se escribe DENTRO de él, lo que Word leía
            # era «no separes este párrafo largo del que viene detrás»: si el
            # siguiente no cabía, empujaba media página en blanco. Es el hueco
            # que David vio detrás de la tabla del cómputo, y es exactamente el
            # mismo fallo que ya se corrigió en el bloque de las tesis —donde
            # atar la cadena entera dejaba media hoja vacía antes de cada
            # jurisprudencia— sin que la lección llegara hasta aquí.
            #
            # Se ata sólo mientras el párrafo es un rótulo solo: si el cuerpo
            # entra en él, deja de estarlo. Un párrafo que ya lleva su propio
            # texto no puede quedar huérfano.
            _solo_rotulo = len(p.text)
            p.paragraph_format.keep_with_next = True
            escribir_cuerpo(p)
            if len(p.text) > _solo_rotulo + 40:
                p.paragraph_format.keep_with_next = False

    def _texto_en(p, texto, resto=None):
        """El primer párrafo continúa el rótulo; el resto van aparte.

        Es el embudo: todo lo que el modelo escribe entra al documento por
        aquí, así que el andamio se quita aquí una vez y no en catorce sitios.
        """
        texto = sin_andamio(texto or "")
        if texto.strip():
            r = p.add_run(texto.strip())
            r.font.name = FUENTE
            r.font.size = TAMANO
        for x in (resto or []):
            x = sin_andamio(x)
            if x.strip():
                parrafo(doc, x.strip())

    # ── RESULTANDO ──
    rotulo(doc, "Resultando")
    res_apartados = []
    # LA PERÍFRASIS EN UN RESULTANDO ES EL APARTADO SIN CUMPLIR. Se avisa —no
    # se borra: la frase ocupa el sitio de algo que debe escribirse, y quitarla
    # dejaría el resultando mudo—. Y arrastra un segundo daño que conviene
    # nombrar en el mismo aviso: el número de expediente del considerando de
    # existencia se lee de estos resultandos.
    try:
        import meta_lenguaje as _ml
        _ev = _ml.perifrasis(" ".join(str(r.get("texto") or "")
                                      for r in (estructura.resultandos or [])))
        for _e in _ev:
            _avisos_bk.append(
                f"UN RESULTANDO EVADE EL DATO: «{_e}». Los resultandos existen "
                f"para individualizar —fecha, órgano, expediente, nombre—, y "
                f"una perífrasis ahí deja además sin número el expediente de "
                f"origen del considerando de existencia.")
    except Exception:
        pass

    for res in (estructura.resultandos or []):
        cuerpo = (res.get("texto") or "").strip()
        if not cuerpo:
            continue
        rot = (res.get("titulo") or "").strip().rstrip(".") + "."
        res_apartados.append((rot, (lambda c: lambda p: _texto_en(p, c))(cuerpo)))
    # La sesión SIEMPRE cierra el resultando y enlaza con el considerando.
    # LA FÓRMULA ESTABA MEDIDA EN EL BANCO Y NADIE LA LEÍA. Otra vez: 43 de 44
    # documentos del propio tribunal escriben «El presente asunto se listó el
    # {fecha}, para verse en sesión ordinaria de {fecha} siguiente; lo
    # anterior, de conformidad con el Acuerdo General…». Aquí se escribía una
    # versión corta de mi cosecha —«se listó para la sesión de *********, la
    # cual se celebró conforme a las disposiciones aplicables»— que fundía LAS
    # DOS FECHAS en un solo hueco y se comía la cita del acuerdo, que es lo que
    # da fundamento a que la sesión sea remota.
    #
    # LOS DOS HUECOS SE QUEDAN, y son honestos: el proyecto se redacta ANTES de
    # la sesión, así que ni el secretario sabe todavía esas fechas. Lo que se
    # arregla es que sean DOS huecos con su forma y su fundamento alrededor, y
    # no un agujero que se traga medio párrafo.
    #
    # LA COLA NORMATIVA: el banco mide dos, y la frontera es abril de 2026. La
    # del Acuerdo 6/2026 rige los asuntos listados a partir de entonces (20/44)
    # y es la vigente; la cola COVID (23/44) es la de los anteriores. Como la
    # fecha de sesión no consta, se escribe la vigente y se avisa.
    _rot_sesion = ("Verificación de la sesión vía remota."
                   if _ta.normalizar(tipo_asunto) in ("queja", "amparo_revision")
                   else "Celebración de la sesión vía remota.")
    _cola_sesion = (
        "lo anterior, de conformidad con el Acuerdo General 6/2026, del Pleno "
        "del Órgano de Administración Judicial que regula la integración y "
        "trámite del expediente electrónico y el uso de videoconferencias en "
        "todos los asuntos competencia de los órganos jurisdiccionales a cargo "
        "del Órgano, publicado en el Diario Oficial de la Federación el "
        "diecisiete de abril de dos mil veintiséis; y,")
    res_apartados.append((
        _rot_sesion,
        lambda p: _texto_en(p, f"El presente asunto se listó el "
                               f"{datos.get('fecha_lista') or HUECO}, para "
                               f"verse en sesión ordinaria de "
                               f"{datos.get('fecha_sesion') or HUECO} "
                               f"siguiente; {_cola_sesion}")))
    if not (datos.get("fecha_lista") and datos.get("fecha_sesion")):
        # Y NO SE PIDEN. David las retiró del formulario el mismo día que se
        # añadieron: «no me sirven porque estas, al final, quedarán hasta el
        # momento en que se revisen por los magistrados». El hueco no es un
        # fallo del sistema, es el estado real del asunto cuando se redacta.
        _avisos_bk.append(
            "LAS DOS FECHAS DE LA SESIÓN VAN EN HUECO, y así tienen que ir: "
            "cuándo se listó el asunto y para qué sesión se fijan cuando los "
            "magistrados lo revisan, después de que esto se escriba. Se "
            "rellenan al engrosar.")
    _avisos_bk.append(
        "COMPRUEBA LA COLA NORMATIVA del párrafo de la sesión: se escribió la del "
        "Acuerdo General 6/2026, que rige los asuntos listados desde abril de "
        "dos mil veintiséis; si éste se listó antes, la cola es la de los "
        "Acuerdos 16/2009 y 12/2020 del otrora Consejo de la Judicatura "
        "Federal.")
    _emitir(res_apartados)

    # ── CONSIDERANDO ──
    rotulo(doc, "Considerando")
    cs = [str(c or "").strip().lower() for c in (calificaciones or []) if c]
    # «INNECESARIO» NO CONCEDE NI NIEGA. Empieza por «i» como «infundado» e
    # «inoperante», pero significa otra cosa: que no se entra al planteamiento.
    # Lo que decide el sentido del fallo es el PRINCIPAL, y si ése es fundado
    # el asunto se concede aunque los accesorios queden sin materia.
    concede = any(c.startswith("fundad") for c in cs)
    esq = esqueleto_de(tipo_asunto)
    q = esq["q"]
    con_apartados = []

    # ═══ EL BANCO MANDA DONDE LA REDACCIÓN ES FORMAL ═══════════════════
    # La competencia, la existencia, la legitimación y la dispensa tienen UNA
    # frase en el oficio, con su cadena de fundamentos en un orden que no es
    # casual. El modelo escribía una versión correcta y anodina, y el
    # secretario la reescribía entera: entonces el adelanto no le ahorró nada.
    # Medido sobre 363 documentos. El estudio y los antecedentes NO llevan
    # plantilla: ahí no hay fórmula que valga.
    import banco as _bk
    _datos_bk = dict(datos)
    _datos_bk.setdefault("q", q)
    # LOS MARCADORES QUE SÍ SE DEDUCEN. Salieron en hueco la primera vez y no
    # tenían por qué: la materia está en el encabezado —«AMPARO DIRECTO CIVIL
    # 380/2025»—, la concordancia depende del género de la autoridad, y el
    # inciso lo manda la materia: c) en civil y mercantil, b) en administrativa
    # y agraria. Medido en el corpus. Un hueco que se puede rellenar es trabajo
    # que se le deja al secretario sin motivo.
    _resp = str(datos.get("responsable", ""))
    _fem = bool(re.match(r"\s*(?:la|las)\s", _resp, re.I) or
                re.search(r"\b(sala|junta|secretar[íi]a|comisi[óo]n|"
                          r"procuradur[íi]a|direcci[óo]n)\b", _resp, re.I))
    _mat = (datos.get("materia") or "").strip().lower()
    if not _mat:
        _enc = str(datos.get("encabezado", "")).lower()
        for m_, clave in (("civil", "civil"), ("administrativ", "administrativa"),
                          ("mercantil", "mercantil"), ("laboral", "laboral"),
                          ("familiar", "familiar"), ("agrari", "agraria")):
            if m_ in _enc:
                _mat = clave
                break
    # LA MATERIA NO PUEDE QUEDAR EN HUECO en mitad de la competencia. Si no se
    # dedujo del encabezado ni del tribunal, se dice la del tipo de asunto, que
    # es cierta: una revisión fiscal es administrativa por definición.
    _datos_bk.setdefault("materia", _mat or {
        "revision_fiscal": "administrativa", "queja": "administrativa",
    }.get(str(tipo_asunto or "").strip().lower(), "") or HUECO)
    _datos_bk.setdefault("inciso", "b" if _mat in ("administrativa", "agraria") else "c")
    # EL INCISO DEL 97 NO ES EL DE LA MATERIA. Ver la nota en tipos_asunto: un
    # solo marcador servía a dos preceptos que se reparten por cosas distintas,
    # y la queja acababa fundando su procedencia en el supuesto equivocado.
    if _ta.normalizar(tipo_asunto) == "queja":
        # DE DÓNDE SE LEE. `descripcion_acto` cae a menudo al genérico «del
        # auto recurrido», que no dice QUÉ se desechó y deja al inciso sin
        # materia. Los resultandos sí lo describen —los acaba de escribir el
        # modelo con el auto delante— y son cuatro párrafos, no el OCR entero:
        # la heurística de una palabra sólo explota cuando se le da un
        # expediente pegado, y esto no lo es.
        _para_inciso = " ".join([
            str(_datos_bk.get("descripcion_acto") or ""),
            " ".join(str(r.get("texto") or "")
                     for r in (estructura.resultandos or []))[:1200],
        ])
        _i97 = _ta.inciso_97(_para_inciso)
        _datos_bk["inciso"] = _i97 or HUECO
        # Y LA COLA VA CON SU INCISO. Si el inciso no se pudo afirmar, la cola
        # tampoco: emitir la del desechamiento de la demanda junto a un inciso
        # en hueco sería afirmar el hecho y callar el fundamento.
        _datos_bk["cola_97"] = _ta.cola_97(_i97) or HUECO
        if not _i97:
            _avisos_bk.append(
                "EL INCISO DEL ARTÍCULO 97, FRACCIÓN I, SALE EN HUECO: no se "
                "pudo afirmar cuál corresponde a este acto. Escríbelo: es el "
                "supuesto que hace procedente la queja, y uno equivocado se "
                "caza en sesión.")
    _datos_bk.setdefault("concordancia", "localizada" if _fem else "localizado")
    # EL ARTÍCULO LO PONE LA PLANTILLA, NO YO. Las fórmulas del banco ya dicen
    # «por el {responsable}» y «dictado por el {responsable}», así que
    # anteponerle el artículo aquí producía «por el el Juez de Distrito». Se
    # entrega el nombre limpio y la plantilla lo enmarca; donde hace falta
    # artículo —el resolutivo, los efectos— se pone en ese sitio.
    _datos_bk["responsable"] = _sin_articulo(_normalizar_autoridad(_resp))
    # Y LOS DATOS QUE LA PLANTILLA PIDE Y NADIE LLENABA. `{descripcion_acto}`
    # salía como hueco «*********» en la competencia de toda queja: es la única
    # frase que dice CONTRA QUÉ se recurre, y sin ella el considerando primero
    # no se sostiene. Sale del propio acto, que el secretario ya subió.
    _datos_bk.setdefault("descripcion_acto", _descripcion_del_acto(datos, tipo_asunto))
    # EL EXPEDIENTE DE ORIGEN. La plantilla de «Existencia del acto reclamado»
    # pide `{expediente}` y nadie lo alimentaba: salía «los autos del
    # expediente *********» en el considerando SEGUNDO. Se lee de lo que el
    # modelo YA escribió en los resultandos —que salió del OCR— y no se le
    # vuelve a preguntar: una llamada más es una ocasión más de inventarlo.
    if not str(_datos_bk.get("expediente") or "").strip():
        try:
            import fase_origen as _fo
            _res_txt = " ".join(str(r.get("texto") or "")
                                for r in (estructura.resultandos or []))
            _datos_bk["expediente"] = (_fo.numero_de(_res_txt,
                                                     str(datos.get("numero") or ""))
                                       or HUECO)
            # Y LA FECHA DE LO RECURRIDO, del mismo sitio: la plantilla de
            # procedencia dice «se impugna el auto de {fecha_acto}» y salía en
            # asteriscos justo al lado del inciso que sí se dedujo.
            if not str(_datos_bk.get("fecha_acto") or "").strip().strip("*"):
                _datos_bk["fecha_acto"] = _fo.fecha_de(_res_txt) or HUECO
        except Exception:
            pass
    # LA FRACCIÓN DEL ACUERDO GENERAL ES LA DE CADA CIRCUITO, Y SE DEDUCE.
    # El banco traía escrita la XXII, que reparte la jurisdicción del Vigésimo
    # Segundo, así que un secretario de Mérida nombraba bien a su tribunal y
    # fundaba su competencia en la fracción de otro. Se dejó en hueco visible.
    #
    # Se dejó de más. Lo que no se deduce es del EXPEDIENTE —ahí no consta—,
    # pero sí del TRIBUNAL, que el secretario declara en el formulario: el
    # punto tercero enumera los circuitos en orden, de modo que la fracción es
    # el número del circuito en romanos. El adelanto real del Vigésimo Segundo
    # lo confirma: fracción XXII. Los cuatro proyectos salían con «fracción
    # *********» en el considerando de competencia, que es el primero que se
    # lee, y ese asterisco era evitable.
    #
    # La regla de la casa se respeta: si el tribunal no dice su circuito, sigue
    # saliendo hueco. Y el aviso a los secretarios de otro circuito sigue
    # pidiéndoles revisar la cadena de fundamentos entera antes de firmar,
    # porque la fracción es una pieza de esa cadena, no la cadena.
    # LOS MARCADORES DE LA PLANTILLA DE REVISIÓN, que nadie llenaba: salían como
    # «en materia *********, por el *********» en mitad del considerando de
    # competencia. El juez de distrito ES la autoridad responsable —en una
    # revisión de amparo indirecto lo recurrido es su sentencia— y la materia ya
    # está calculada tres líneas más arriba.
    # CON SU ARTÍCULO, porque la plantilla ya no lo pone: «por el Sala Regional»
    # no es español y la concordancia depende del órgano, no de la frase.
    _datos_bk.setdefault("juez_distrito", _con_articulo(_resp) or HUECO)
    _datos_bk.setdefault("juzgado", _con_articulo(_resp) or HUECO)
    _datos_bk.setdefault("recurrente", str(datos.get("quejoso") or "").strip() or HUECO)
    _datos_bk.setdefault(
        "fraccion_acuerdo",
        str(datos.get("fraccion_acuerdo") or "").strip()
        or _bk.fraccion_del_acuerdo(str(datos.get("tribunal") or ""))
        or HUECO)
    _datos_bk.setdefault("fecha_acto", str(datos.get("fecha_acto") or "").strip()
                         or HUECO)
    _datos_bk.setdefault("objeto", (f"una sentencia definitiva en materia {_mat}"
                                    if _mat else "una sentencia definitiva"))
    _huecos_bk = []

    def _del_banco(ident, respaldo):
        """La frase del oficio si el banco la tiene; si no, la del modelo."""
        t, faltan = _bk.texto_de(tipo_asunto, ident, _datos_bk)
        if t:
            _huecos_bk.extend(faltan)
            return t
        return respaldo


    _comp = _del_banco("competencia", estructura.competencia)
    if (_comp or "").strip():
        con_apartados.append((_bk.rotulo_de(tipo_asunto, "competencia", "Competencia."),
                              (lambda c: lambda p: _texto_en(p, c))(_comp)))
    _exi = _del_banco("existencia", estructura.existencia)
    if esq["existencia"] and (_exi or "").strip():
        con_apartados.append((_bk.rotulo_de(tipo_asunto, "existencia",
                                            "Existencia del acto reclamado."),
                              (lambda c: lambda p: _texto_en(p, c))(_exi)))

    # Legitimación y oportunidad, con LA TABLA detrás.
    from fase0_oportunidad import parrafo_oportunidad
    # EL FUNDAMENTO Y EL VOCABULARIO SALEN DEL CATÁLOGO. Antes decía siempre
    # «el precepto 17 del mencionado ordenamiento», que es el del amparo: una
    # queja se fundaba en el artículo del amparo y una revisión fiscal, en la
    # Ley de Amparo cuando la suya es la LFPCA.
    _op = parrafo_oportunidad(
        computo,
        _ta.plazo_de(tipo_asunto, "").get("fundamento") or "artículo 17 de la Ley de Amparo",
        tipo_asunto)

    def _legitimacion(p):
        _texto_en(p, _op)
        # La tabla va en los cuatro tipos. En el corpus casi no aparece
        # —el secretario la dibuja a mano y cuesta—, pero eso mide lo que hoy
        # es caro hacer, no lo que sobra: la máquina tiene el calendario.
        if esq["tabla_computo"]:
            tabla_computo(doc, computo, fecha_en_letra, tipo_asunto)

    _legit = (_bk.rotulo_de(tipo_asunto, "legitimacion", esq["legitimacion"]),
              _legitimacion)

    # EL ORDEN ENTRE PROCEDENCIA Y LEGITIMACIÓN LO FIJA EL CATÁLOGO, no este
    # archivo. En la queja el corpus resuelve primero si el recurso procede
    # —21 competencias, 20 procedencias, 16 legitimaciones, en ese orden— y el
    # adelanto real de la QC 259/2025 hace lo mismo: SEGUNDO Procedencia,
    # TERCERO Legitimación. Aquí salía al revés porque el orden estaba escrito
    # a mano, y con razón: en el amparo directo la legitimación va primero.
    # Es la misma pregunta de siempre —¿quién manda, el código o lo medido?—
    # y la respuesta no cambia.
    _cons = [c for c, _ in
             (_ta.estructura_de(tipo_asunto).get("considerandos") or [])]

    def _puesto(clave: str) -> int:
        for i, c in enumerate(_cons):
            if clave in c.lower():
                return i
        return 99

    _procedencia_primero = _puesto("procedencia") < _puesto("legitim")
    if not _procedencia_primero:
        con_apartados.append(_legit)

    # PROCEDENCIA NO ES UN CONSIDERANDO PROPIO cuando hay «Existencia del acto
    # reclamado»: medido, es su ALTERNATIVA —3 de 26, en asuntos venidos de
    # juez y no de sala—, no un apartado más. Emitirlos los dos corría todo un
    # ordinal y dejaba el Estudio en SÉPTIMO donde el corpus lo tiene SEXTO.
    # PROCEDENCIA sólo donde el corpus la tiene como apartado propio —queja
    # (14 de 20) y revisión fiscal (16 de 28)— o en el amparo cuando sustituye
    # a «Existencia del acto reclamado». En revisión civil aparece en 2 de 31:
    # emitirla por defecto ahí corría un ordinal contra la medida.
    # LA PROCEDENCIA DE LA REVISIÓN FISCAL SE MOTIVA DE OFICIO. Salía «El
    # juicio es procedente y no se advierte causa de improcedencia» —fórmula
    # del amparo, y encima llamando JUICIO a un recurso— porque este apartado
    # tomaba SIEMPRE el texto del modelo y nunca el del banco, aunque el banco
    # tenga medida la suya. Es la tercera vez que aparece el mismo patrón: una
    # fórmula del corpus que existe y es inalcanzable.
    #
    # Y en revisión fiscal ni siquiera basta la del banco: el 63 de la LFPCA
    # obliga al Colegiado a decir POR QUÉ procede, y la vía más común es la
    # cuantía. Eso es aritmética, no criterio, así que se calcula.
    _proc = estructura.procedencia or ""
    if _ta.normalizar(tipo_asunto) == "revision_fiscal":
        try:
            import fase_procedencia_rf as _pf
            _fuente = " ".join([
                str(datos.get("antecedentes") or ""),
                " ".join(str(r.get("texto") or "")
                         for r in (estructura.resultandos or [])),
                str(datos.get("acto") or "")])
            # El año es el de la resolución recurrida. Se toma el de su
            # notificación, que es el dato duro que el secretario tecleó: entre
            # emisión y notificación median días, no años, salvo en el cambio
            # de ejercicio —y ahí el aviso lo dirá, porque la cifra saldrá.
            _anio = getattr(computo, "notificacion", None)
            _anio = _anio.year if _anio else 0
            _p_rf, _av_rf = _pf.parrafo(_fuente, _anio)
            _avisos_bk.extend(_av_rf)
            if _p_rf:
                _proc = _p_rf
        except Exception:
            pass
    # Y EN LOS CUATRO TIPOS, LA DEL BANCO ANTES QUE LA DEL MODELO. Éste era el
    # apartado que tomaba SIEMPRE el texto libre, y por eso la queja decía «El
    # recurso de queja es procedente y no se advierte causa de improcedencia»
    # —una fórmula que no funda nada— teniendo el banco medida la suya, que
    # cita el artículo 97, fracción I, con su inciso y con el auto que se
    # impugna. Es el mismo patrón por tercera vez: una fórmula del corpus que
    # existe y a la que nadie llama.
    if not (_ta.normalizar(tipo_asunto) == "revision_fiscal" and _proc
            != (estructura.procedencia or "")):
        _del_banco_proc = _bk.texto_de(tipo_asunto, "procedencia", _datos_bk)[0]
        if (_del_banco_proc or "").strip():
            _proc = _del_banco_proc
    if (_proc or "").strip() and (
            esq.get("procedencia_propia")
            or (esq["existencia"] and not (estructura.existencia or "").strip())):
        con_apartados.append(("Procedencia.",
                              (lambda c: lambda p: _texto_en(p, c))(_proc)))

    if _procedencia_primero:
        con_apartados.append(_legit)

    # LA DISPENSA. El rótulo promete el acto reclamado y los conceptos, y el
    # contenido es justamente que NO hace falta transcribirlos: 21 de 26.
    def _dispensa(p):
        _texto_en(
            p,
            _del_banco("dispensa", "") or
            f"Es innecesario transcribir el contenido de "
            f"{esq['recurrido']} y los {q} hechos valer, pues el deber formal "
            f"y material de "
            f"exponer los argumentos legales que sustenten esta resolución no "
            f"depende de la reproducción literal de los aspectos que conforman "
            f"la litis, sino de su adecuado análisis.")

    con_apartados.append((esq["dispensa"].format(q=q), _dispensa))

    if antecedentes:
        def _antecedentes(p):
            _texto_en(p,
                      "Previo al análisis de los planteamientos que se proponen, "
                      "es menester relatar los hechos relevantes del asunto.")
            for x in (antecedentes or []):
                if x.strip():
                    parrafo_con_citas(doc, x.strip(), notas)
        con_apartados.append(("Antecedentes.", _antecedentes))

    # EL ESTUDIO. Es el ÚLTIMO considerando salvo que detrás vaya Efectos, y
    # lleva dentro los subtítulos en negrita, sin ordinal.
    _cuerpo_estudio, _efectos_escritos = partir_efectos(estudio or [])

    def _estudio(p):
        calif = _calificacion_plural(cs)
        _texto_en(p, f"Los {q} son {calif}." if calif else "")
        if resumen_acto:
            _subtitulo(doc, esq["sub_recurrido"])
            for x in resumen_acto:
                if x.strip():
                    parrafo_con_citas(doc, x.strip(), notas)
        if resumen_conceptos:
            _subtitulo(doc, q[0].upper() + q[1:])
            for x in resumen_conceptos:
                if x.strip():
                    parrafo_con_citas(doc, x.strip(), notas)
        if (marco_escrito or "").strip():
            _subtitulo(doc, "Marco jurídico")
            for x in re.split(r"\n\s*\n", marco_escrito):
                if x.strip():
                    parrafo(doc, x.strip())
        _subtitulo(doc, "Solución")
        _escribir_estudio(doc, _cuerpo_estudio, tesis, notas, normas)
        # EL CIERRE ES DEL TIPO. Aquí decía «lo procedente es negar el amparo
        # solicitado» en los cuatro, incluida la queja, que además decretaba
        # «Es infundado el recurso de queja» treinta líneas más abajo: el mismo
        # documento afirmaba dos desenlaces incompatibles.
        #
        # Y CUANDO CONCEDE, EL CIERRE VA AQUÍ TAMBIÉN salvo en el amparo
        # directo, que lleva su apartado de «Efectos». Ponerlo en un apartado
        # propio para los recursos correría el ordinal y dejaría el Estudio
        # donde el corpus no lo tiene.
        if not concede:
            parrafo(doc, _ta.parrafo_cierre(tipo_asunto, False))
        elif not _ta.cierre_de(tipo_asunto)["efectos"]:
            parrafo(doc, _ta.parrafo_cierre(tipo_asunto, True,
                                            _calificacion_plural(cs)))

    # ═══════════════════════════════════════════════════════════════════════
    # SI EL CÓMPUTO DA EXTEMPORÁNEA, NO HAY FONDO
    # ═══════════════════════════════════════════════════════════════════════
    # David: «si oportunidad == Extemporánea, el flujo debe abortar el estudio
    # de fondo y generar automáticamente el sobreseimiento». La falla que
    # describe —declarar la extemporaneidad y luego conceder— no se corrige
    # avisando: se corrige no escribiendo el estudio.
    #
    # El apartado de fondo se sustituye por el de improcedencia, que es lo que
    # el corpus escribe: «{ordinal}. Extemporaneidad del recurso de revisión.
    # El presente medio de impugnación se interpuso de manera extemporánea».
    _extemp = (getattr(computo, "oportuna", None) is False
               and not getattr(computo, "en_cualquier_tiempo", False))
    if _extemp:
        _ex = _ta.extemporaneo_de(tipo_asunto)
        _avisos_bk.insert(0, (
            f"EL CÓMPUTO DA EXTEMPORÁNEA: el proyecto NO entra al fondo y "
            f"resuelve la improcedencia conforme al {_ex['fundamento']}. Si la "
            f"fecha de notificación o la de presentación están mal, corrígelas "
            f"y vuelve a generar: de esas dos fechas depende todo el asunto."))
        con_apartados.append(
            (_ex["rotulo"] + ".",
             (lambda c: lambda p: _texto_en(p, c))(_ex["considerando"])))
    else:
        # ── LA CUESTIÓN, FIJADA ANTES DE RESOLVERLA ────────────────────────
        # Roberto Lara Chagoyán, «Sobre la estructura de las sentencias en
        # México», § 3.2 y § 4.2: «la desgracia de muchas malas sentencias
        # comienza con el descuido del deber de fijar cuidadosamente la
        # cuestión… Una forma de mejorar los planteamientos es utilizar la
        # PREGUNTA EXPRESA», y su ejemplo de apartado: «TERCERO. Materia de la
        # revisión. Se constriñe a determinar si la parte quejosa logra, con
        # sus agravios, desvirtuar las razones por las que el Juez de Distrito
        # negó el amparo».
        #
        # POR QUÉ SE COMPONE Y NO SE PIDE. Se lo pedí al modelo en el prompt y
        # no lo hizo: medido, 5 de 6 problemas sin pregunta, y después del
        # arreglo seguían siendo cero. No era desobediencia: la instrucción
        # cae en el carácter 27.732 de un prompt de 32.076, entre decenas de
        # reglas, y una regla más en el último tercio se pierde. El pipeline YA
        # tiene las preguntas —las calcula `proponer`, y las calcula bien—, así
        # que se escriben. Lo que se puede componer no se pide.
        _cuestiones = [str(getattr(c, "problema", "") or "").strip()
                       for c in (criterios or [])
                       if str(getattr(c, "problema", "") or "").strip()]
        # El orden es el de prelación lógica, igual que en el estudio: primero
        # el principal, del que dependen los demás.
        _cuestiones = [q for c, q in sorted(
            zip(criterios or [], _cuestiones),
            key=lambda x: 0 if str(getattr(x[0], "jerarquia", "")).lower()
            == "principal" else 1)]
        if _cuestiones:
            def _materia_del_asunto(p, _qs=_cuestiones):
                _lineas = []
                for i_, q_ in enumerate(_qs, 1):
                    q_ = " ".join((q_ or "").split())
                    if not q_.startswith("¿"):
                        q_ = "¿" + q_.lstrip("¿")
                    if not q_.rstrip().endswith("?"):
                        q_ = q_.rstrip(" .;") + "?"
                    _lineas.append(f"{i_}. {q_}")
                # CADA PREGUNTA EN SU PÁRRAFO. Escritas con `_texto_en(p, …)`
                # una tras otra se pegaban todas al rótulo y salía un bloque de
                # 1.682 caracteres seguidos —«…cuestiones siguientes:1. ¿La
                # audiencia…?2. ¿La vista…?»—, que es exactamente lo contrario
                # de lo que la pregunta expresa viene a resolver: que se vea de
                # un golpe qué se va a decidir.
                _texto_en(p, _ta.materia_a_resolver(tipo_asunto), _lineas)
            con_apartados.append((_ta.rotulo_materia_de(tipo_asunto), _materia_del_asunto))

        con_apartados.append((_ta.rotulo_estudio_de(tipo_asunto).rstrip(".") + ".",
                              _estudio))

    # EL APARTADO DE EFECTOS ES DEL AMPARO. «Procede conceder el amparo y
    # protección de la Justicia Federal para el efecto de que la responsable
    # deje insubsistente la sentencia reclamada» no se escribe en una queja
    # fundada: ahí se revoca el auto y se ordena proveer de nuevo. Se emitía en
    # los cuatro tipos, así que era la segunda puerta —la de la rama FUNDADA—
    # por la que la fórmula del amparo entraba en un recurso. Arreglar sólo la
    # otra habría tapado la mitad.
    if _extemp:
        pass          # no hay efectos de una concesión que no existe
    elif concede and _ta.cierre_de(tipo_asunto)["efectos"]:
        def _efectos(p):
            if _efectos_escritos:
                _texto_en(p, _efectos_escritos[0], _efectos_escritos[1:])
                return
            _texto_en(p,
                      f"Consecuentemente, procede conceder el amparo y protección "
                      f"de la Justicia Federal a {datos.get('quejoso','')} para el "
                      f"efecto de que "
                      f"{_con_articulo(datos.get('responsable','')) or HUECO} "
                      f"deje insubsistente la sentencia reclamada y dicte otra en "
                      f"la que atienda los lineamientos de esta ejecutoria.")
        con_apartados.append(("Efectos.", _efectos))

    _emitir(con_apartados)

    # ── RESUELVE ──
    p_pe = parrafo(doc, "Por lo expuesto y fundado, se:", sangria=True)
    p_pe.paragraph_format.space_before = Pt(14)
    rotulo(doc, "Resuelve")
    _res = RESOLUTIVO.get(_ta.normalizar(tipo_asunto),
                          RESOLUTIVO["amparo_directo"])

    # ═══════════════════════════════════════════════════════════════════════
    # EL AMPARO EN REVISIÓN TIENE DOS PUNTOS, NO UNO
    # ═══════════════════════════════════════════════════════════════════════
    # Está medido en el corpus del propio tribunal, con esas palabras: «en
    # revisión el resolutivo tiene DOS puntos: PRIMERO decide sobre la
    # sentencia recurrida —confirma, modifica, revoca— y SEGUNDO reproduce el
    # sentido del amparo —ampara, no ampara, sobresee—. Sólo hay ÚNICO cuando
    # se desecha el recurso». Y estaba en `banco_plantillas.json` sin que nadie
    # lo leyera: el proyecto salía siempre con «ÚNICO. Se confirma la sentencia
    # recurrida», de modo que un amparo en revisión NUNCA amparaba.
    _hecho = False
    if _extemp:
        _ex = _ta.extemporaneo_de(tipo_asunto)
        _t = _ex["resolutivo"].replace("{quejoso}",
                                       str(datos.get("quejoso") or HUECO))
        _cab, _resto = _t.split(". ", 1) if ". " in _t else (_t, "")
        tramos(doc, [(_cab + ". ", {"bold": True}), (_resto, {})], sangria=False)
        _hecho = True
    elif _ta.normalizar(tipo_asunto) == "amparo_revision":
        # EL ESTUDIO ENTERO, NO SUS PRIMEROS SEIS MIL CARACTERES. `resolvio_a_quo`
        # es un barrido de expresiones regulares: cuesta lo mismo mirarlo todo,
        # y truncarlo sólo puede perder la frase que dice qué resolvió el
        # juzgado. Un resolutivo en hueco por no haber leído el párrafo 40 es
        # un precio absurdo por un ahorro que nadie iba a notar.
        _fuente_rama = " ".join([
            str(datos.get("antecedentes") or ""),
            str(datos.get("acto") or ""),
            " ".join(str(r.get("texto") or "")
                     for r in (estructura.resultandos or [])),
            str(estudio or "")])
        try:
            import fase_rama as _fr
            # LOS ANTECEDENTES MANDAN SOBRE EL ESTUDIO. En el estudio la
            # palabra «sobreseimiento» aparece dentro de las TESIS
            # TRANSCRITAS, y con eso el proyecto confirmaba un sobreseimiento
            # que nadie decretó —medido sobre el ARA 17/2025: con el estudio
            # entero da «sobresee», sin las tesis da «concede»—. Lo que hizo el
            # juzgado está en los antecedentes y en los resultandos, que es
            # donde el catálogo manda escribirlo.
            _antes_rama = " ".join([
                str(datos.get("antecedentes") or ""),
                " ".join(str(r.get("texto") or "")
                         for r in (estructura.resultandos or []))]).strip()
            _que_hizo = _fr.resolvio_a_quo(_fuente_rama, _antes_rama)
            # EL SENTIDO EN PLENITUD SE LEE DEL ESTUDIO, no del recurso. Que el
            # agravio sea fundado prueba que el juez no debió sobreseer, no que
            # el quejoso tenga razón en el fondo.
            _sent_amparo = _fr.sentido_en_plenitud(str(estudio or ""))
            _clave = _ta.rama_revision(
                _que_hizo,
                "fundado" if concede else "infundado",
                solo_efectos=_fr.solo_los_efectos(str(estudio or "")),
                violacion_procesal=_fr.hay_violacion_procesal(str(estudio or "")),
                sentido_amparo=_sent_amparo)
            _rama = _ta.RAMAS_REVISION[_clave]
            # ═══════════════════════════════════════════════════════════════
            # EL RESPALDO AMPARABA CONTRA EL ÓRGANO RECURRIDO
            # ═══════════════════════════════════════════════════════════════
            # `datos["responsable"]` en un recurso es el ÓRGANO RECURRIDO —la
            # propia carátula lo rotula así: «ÓRGANO RECURRIDO: EL JUZGADO
            # SEGUNDO DE DISTRITO»—, no la responsable originaria. Cuando
            # `responsable_originaria` no lograba leerla del resumen, el
            # respaldo escribía «La Justicia de la Unión ampara y protege a
            # Juan Pérez, contra el acto reclamado AL JUZGADO SEGUNDO DE
            # DISTRITO». Es un disparate: se ampara contra el acto de la
            # autoridad que emitió el acto reclamado del amparo indirecto, no
            # contra el juez que resolvió ese amparo.
            #
            # Este módulo ya tenía la doctrina escrita en `fase_rama`: «Y
            # CUANDO NO CONSTA, SE DEJA EL HUECO». Se aplica. El comodín se ve,
            # el linter lo cuenta y el aviso dice dónde buscarla.
            #
            # Y PASA POR `_con_articulo`, que era el otro defecto: el nombre
            # sale del resumen sin artículo y el punto resolutivo decía «contra
            # el acto reclamado a Director de Ingresos». El artículo se pone
            # aquí y `_contraer` hace el resto —«a el» → «al»—.
            _orig = _con_articulo(_fr.responsable_originaria(_antes_rama)
                                  or _fr.responsable_originaria(_fuente_rama))
            if not _orig:
                _orig = HUECO
                _avisos_bk.append(
                    "NO SE PUDO LEER LA AUTORIDAD RESPONSABLE ORIGINARIA y el "
                    "SEGUNDO punto resolutivo va con comodín. NO se puso el "
                    "órgano recurrido en su lugar: el amparo se concede o se "
                    "niega contra el acto de la autoridad que lo emitió en el "
                    "amparo indirecto, no contra el Juzgado de Distrito que lo "
                    "resolvió. Escríbela tú, está en la sentencia recurrida.")
            if _rama.get("aviso"):
                _avisos_bk.append(_rama["aviso"])
            for _pt in _rama["puntos"]:
                _txt = (_pt.replace("{HUECO}", HUECO)
                           .replace("{quejoso}", str(datos.get("quejoso") or HUECO))
                           .replace("{responsable_originaria}", _orig)
                           .replace("{expediente}",
                                    str(_datos_bk.get("expediente") or HUECO)))
                _txt = _contraer(_txt)
                _cab, _resto = _txt.split(". ", 1) if ". " in _txt else (_txt, "")
                tramos(doc, [(_cab + ". ", {"bold": True}), (_resto, {})],
                       sangria=False)
            _avisos_bk.append(
                f"RESOLUTIVO DE REVISIÓN, rama «{_clave}» "
                f"({_rama['fundamento']}). El a quo "
                f"{_que_hizo or 'no consta qué resolvió'}; el recurso resultó "
                f"{'fundado' if concede else 'infundado'}. Compruébalo: de esto "
                f"depende que se confirme, se revoque o se modifique.")
            # EL SEGUNDO PUNTO NO SALE DEL RECURSO. Cuando se levanta el
            # sobreseimiento y el tribunal asume jurisdicción, el sentido del
            # amparo lo decide el estudio, y si el estudio no lo dijo, quien
            # firma tiene que saber que ahí se puso el sentido por omisión.
            if _rama.get("plenitud") and _que_hizo == "sobresee":
                _avisos_bk.append(
                    "SE ASUME JURISDICCIÓN Y EL SEGUNDO RESOLUTIVO "
                    + (f"NIEGA el amparo porque el estudio concluye que procede "
                       f"negarlo." if _sent_amparo == "niega" else
                       f"CONCEDE el amparo, que es el sentido por omisión "
                       f"porque el estudio no dice cuál es." if not _sent_amparo
                       else "CONCEDE el amparo, como concluye el estudio.")
                    + " Que el agravio sea fundado sólo prueba que no debió "
                      "sobreseerse: el fondo se estudia por primera vez aquí y "
                      "el sentido es tuyo.")
            _hecho = True
        except Exception as _erm:
            print(f"   ⚠️ rama de revisión no determinada: {_erm}")

    if _hecho:
        pass
    elif _res.get("punto"):
        # Queja y revisión: NO amparan. Se califica el recurso o se resuelve
        # sobre la sentencia recurrida, que es lo que hacen los engroses.
        _cal = _res["calif"][0] if concede else _res["calif"][1]
        # LOS DOS DATOS QUE IDENTIFICAN LA SENTENCIA. Se leen de los
        # resultandos, que es donde el propio proyecto acaba de escribirlos, y
        # si no se pudieron leer salen en hueco con su aviso: un resolutivo que
        # revoca «la sentencia recurrida» a secas no dice cuál, y quien lo
        # ejecute tiene que ir a buscarla.
        _fecha_s = _expte_s = ""
        if "{fecha_sentencia}" in _res["punto"]:
            try:
                import fase_origen as _fo_r
                _res_txt = " ".join(str(r.get("texto") or "")
                                    for r in (estructura.resultandos or []))
                _fuente = f"{datos.get('antecedentes') or ''} {_res_txt}"
                _fecha_s = _fo_r.fecha_de(_fuente)
                _expte_s = _fo_r.numero_de(_fuente)
            except Exception as _eo:
                print(f"   ⚠️ no se pudo leer fecha/expediente del recurrido: {_eo}")
            if not _fecha_s or not _expte_s:
                _avisos_bk.append(
                    "EL RESOLUTIVO NO IDENTIFICA LA SENTENCIA POR COMPLETO: "
                    + ("falta su FECHA. " if not _fecha_s else "")
                    + ("falta el EXPEDIENTE de origen. " if not _expte_s else "")
                    + "No se pudo leer de los resultandos y sale en hueco. "
                      "Escríbelo: es lo que distingue esta sentencia de las "
                      "demás que dictó la misma Sala.")
        _texto = _res["punto"].format(
            calificacion=_cal,
            fecha_sentencia=_fecha_s or HUECO,
            expediente_origen=_expte_s or HUECO,
            responsable=_con_articulo(datos.get("responsable", "")) or HUECO)
        _texto = _contraer(_texto)
        tramos(doc, [("ÚNICO. ", {"bold": True}), (_texto, {})], sangria=False)
    else:
        formula = _AMPARA if concede else _NO_AMPARA
        tramos(doc, [("ÚNICO. ", {"bold": True}),
                     ("La Justicia de la Unión ", {}),
                     (formula, {"bold": True}),
                     (f" a {datos.get('quejoso','') or HUECO}, en contra de "
                      f"{esq.get('recurrido','la sentencia reclamada')}, dictada "
                      f"por {_con_articulo(datos.get('responsable','')) or HUECO}, "
                      f"precisada en el primer resultando de esta ejecutoria.", {})],
               sangria=False)

    parrafo(doc, _res["notif"], sangria=True)

    _bloque_firmas(doc, datos)
    # RED DE SEGURIDAD. Si alguna marca sobrevivió a todo lo anterior —porque el
    # modelo la escribió de una forma que no previmos—, se borra antes de
    # guardar. El andamio no sale al papel, y punto.
    _RX_RESTO = re.compile(r"\s*\[{1,2}[^\[\]]{0,120}\]{0,2}")
    for p in doc.paragraphs:
        if "[[" in p.text:
            entero = _RX_RESTO.sub("", p.text)
            if p.runs:
                p.runs[0].text = entero
                for r in p.runs[1:]:
                    r.text = ""
    doc.save(ruta_salida)
    _inyectar_notas(ruta_salida, notas)
    # Los avisos deterministas de la carátula viajan con el documento. Se
    # cuelgan de la estructura porque es lo que ya recorre el camino de vuelta.
    try:
        for _a in list(avisos_doc) + list(_avisos_bk):
            if _a not in estructura.avisos:
                estructura.avisos.append(_a)
    except Exception:
        pass
    return ruta_salida
