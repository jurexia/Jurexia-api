"""El circuito completo del adelanto: de dos PDF a un .docx en la mano.

Encadena lo que ya existe por separado:

    Fase 0   fase0_oportunidad.py   ficha y cómputo — SIN modelo, es aritmética
    Fase 1-3 fases123_pipeline.py   los dos resúmenes y los problemas jurídicos
    Fase 7   ensamblar_adelanto.py  relleno de la plantilla REAL del secretario

Y se detiene donde tiene que detenerse: el sentido del fallo y el criterio no
los pone la máquina. El .docx sale con esos huecos marcados, igual que el
adelanto de papel, y `huecos_pendientes()` los enumera para que nadie firme un
documento con un `*****` dentro.
"""

from __future__ import annotations

import asyncio
import datetime as _dt
import re
from dataclasses import dataclass, field
from typing import Optional

import ensamblar_adelanto as ens
import fase0_oportunidad as f0
import fase_partes as fpartes
import fase6_estudio as f6
import fase6_rag as f6rag
import marco_juridico as mjur
import fases123_pipeline as f123


@dataclass
class Encargo:
    """Lo que el secretario aporta. Todo esto lo sabe de memoria o lo lee de un
    sello; nada de esto justifica pagar el OCR de un expediente."""
    numero: str                       # «512/2026»
    encabezado: str                   # «AMPARO DIRECTO ADMINISTRATIVO: 512/2026»
    quejoso: str
    magistrado: str
    secretario: str
    notificacion: _dt.date
    presentacion: _dt.date
    # LA OMISIÓN ES LA REGLA GENERAL DE LA LEY DE AMPARO, no la de un
    # tribunal concreto: esto lo usan secretarios de toda la república.
    regla_surtimiento: str = "personal"
    # EL PLAZO LO PONE LA LEY SEGÚN EL TIPO. Cero significa «el que
    # corresponda»; se resuelve al computar, con `tipos_asunto`.
    plazo: int = 0
    # La excepción de plazo que el secretario haya declarado: en la queja,
    # «suspension» (dos días) u «omision_tramite» (en cualquier tiempo).
    excepcion_plazo: str = ""
    # Los días que el secretario declara inhábiles y el calendario federal no
    # trae: un acuerdo de suspensión de labores de su tribunal, una
    # contingencia, un no laborable local.
    dias_inhabiles_extra: list = field(default_factory=list)
    # LAS DOS FECHAS DE LA SESIÓN, que el secretario sí sabe y el sistema no
    # puede deducir de ningún documento: el proyecto se escribe ANTES de que la
    # sesión ocurra. Salían como dos comodines de asteriscos en los cinco
    # proyectos —el hallazgo más repetido de la última medición—, y era un
    # hueco honesto pero evitable: basta preguntarlo. Vacías = siguen en hueco,
    # que es mejor que inventarlas. ISO.
    fecha_lista: str = ""
    fecha_sesion: str = ""
    responsable: Optional[str] = None
    es_recurso: bool = False
    # LA HERRAMIENTA NO ES DE UN TRIBUNAL, ES DE TODOS. Estos tres campos son
    # lo que impide que un secretario de otro circuito firme «Resolución del
    # Tercer Tribunal Colegiado… del Vigésimo Segundo Circuito» sin verlo: en
    # modo `generado` el documento se escribe entero con ESTOS datos y no hay
    # plantilla ajena de la que heredar identidad.
    # El tipo decide el ESQUELETO del documento: los recursos no llevan
    # «Existencia del acto reclamado», la queja hace el cómputo en prosa y cada
    # uno rotula la dispensa a su manera. Medido por tipo, no supuesto.
    tipo_asunto: str = "amparo_directo"
    tribunal: str = ""                # «Primer Tribunal Colegiado… del Décimo…»
    ciudad: str = ""                  # «Mérida, Yucatán»
    modo: str = "generado"            # plantilla | generado
    plantilla: str = ""               # el .docx del propio tribunal
    coleccion_estatal: str = ""       # «leyes_queretaro», para el RAG del fondo
    # LA MATERIA, DECLARADA. Decide el silo del RAG y el filtro del sondeo de
    # precedentes. Vacía = se deduce del tribunal y del encabezado, como antes.
    materia: str = ""


@dataclass
class Resultado:
    ruta: str
    computo: f0.Computo
    fases: f123.Fases123
    huecos: list[str] = field(default_factory=list)
    avisos: list[str] = field(default_factory=list)
    # Se guardan para poder REENSAMBLAR cuando llegue el criterio, sin releer
    # los PDF ni volver a pagar los resúmenes.
    encargo: Optional["Encargo"] = None
    estudio: str = ""
    advertencias: str = ""
    partes: Optional["fpartes.Partes"] = None
    # Las partes estructurales ya escritas, para no volver a pedirlas.
    estructura: object = None


    @property
    def listo_para_el_secretario(self) -> bool:
        """El documento está completo HASTA donde puede estarlo sin criterio."""
        return not self.avisos


# ═══ EL RELOJ ═════════════════════════════════════════════════════════════
# David, 30-ago-2026: «¿cómo aceleramos la generación de la sentencia sin
# perder calidad? mide cuánto tarda». Lo primero es medir por fase: sin eso se
# optimiza lo que se supone lento, que casi nunca es lo que lo es.
import time as _time
from contextlib import contextmanager

TIEMPOS: dict = {}


@contextmanager
def cronometrar(nombre: str):
    t0 = _time.perf_counter()
    try:
        yield
    finally:
        dt = _time.perf_counter() - t0
        TIEMPOS[nombre] = round(dt, 1)
        print(f"   ⏱️  {nombre}: {dt:.1f}s")


def reloj_resumen(total: float = 0.0) -> str:
    if not TIEMPOS:
        return ""
    partes = " · ".join(f"{k} {v}s" for k, v in TIEMPOS.items())
    return f"{partes}" + (f" · TOTAL {total:.1f}s" if total else "")


async def generar(cliente, e: Encargo, texto_acto: str, texto_conceptos: str,
                  ruta_salida: str, texto_autos: str = "") -> Resultado:
    """El circuito entero. `cliente` es el AsyncOpenAI de main.py.

    `texto_autos` son las constancias del expediente, si el secretario las
    subió. NO son fuente de derecho y no se confunden con el acervo: son los
    documentos del caso —contrato colectivo, reglamento, peritajes— que no
    están en ninguna base pública y sin los cuales hay asuntos que no se pueden
    resolver. En el ADL 382/2024 el motor dijo «falta el texto de las cláusulas
    40 y 83» y ese texto estaba en un PDF que el secretario había subido y que
    el pipeline no leía.
    """
    avisos: list[str] = []

    # ── Fase 0 — aritmética, sin modelo ──────────────────────────────────
    # LA REGLA POR DEFECTO NO PUEDE SER LA DE OTRA MATERIA. `tja_qro_boletin`
    # —Boletín Jurisdiccional, surte al tercer día hábil— es del Tribunal de
    # Justicia Administrativa de Querétaro, y por ser el valor por omisión se
    # aplicó a un amparo LABORAL contra un laudo de la Junta Federal. Los laudos
    # se notifican PERSONALMENTE (artículo 742 de la Ley Federal del Trabajo) y
    # el propio proyecto se contradecía: los antecedentes decían que el actuario
    # notificó en persona y el cómputo hablaba de boletín.
    #
    # Un plazo mal contado invalida la sentencia, así que aquí no se hereda un
    # valor por omisión de otra materia: si la materia es laboral y nadie
    # declaró otra cosa, se cuenta personal y SE AVISA.
    # ── El plazo y el tipo, del catálogo ─────────────────────────────────
    import tipos_asunto as _ta
    _tipo = _ta.normalizar(getattr(e, "tipo_asunto", "")) or "amparo_directo"
    e.tipo_asunto = _tipo
    # `es_recurso` DEJA DE SER UN CAMPO APARTE. Era independiente del tipo y
    # podía contradecirlo —un amparo directo marcado como recurso escribía
    # «agravios» donde van conceptos de violación—. Lo dice el tipo y punto.
    e.es_recurso = _tipo != "amparo_directo"
    _pl = _ta.plazo_de(_tipo, getattr(e, "excepcion_plazo", ""))
    if _pl.get("aviso"):
        avisos.append(_pl["aviso"])
    if _pl["en_cualquier_tiempo"]:
        # NO ES UN PLAZO LARGO: ES QUE NO HAY PLAZO. Contar días aquí y
        # declarar extemporaneidad sería inventar una causa de improcedencia.
        e.plazo = 0
        avisos.append(
            f"Este recurso procede EN CUALQUIER TIEMPO ({_pl['fundamento']}), "
            f"así que no se computa plazo ni puede declararse extemporáneo.")
    elif not e.plazo:
        e.plazo = _pl["dias"]

    _mat = fp_materia(e)
    if _mat == "laboral" and e.regla_surtimiento in ("tja_qro_boletin", "lista"):
        e.regla_surtimiento = "personal"
        avisos.append(
            "El cómputo se hizo con notificación PERSONAL, que es como se "
            "notifican los laudos (artículo 742 de la Ley Federal del "
            "Trabajo). Venía declarada la regla del Boletín Jurisdiccional del "
            "Tribunal de Justicia Administrativa de Querétaro, que es de otra "
            "materia. Confírmalo contra la constancia de notificación.")
    # EL CERO VIAJA HASTA EL CÓMPUTO. `e.plazo or 15` lo convertía en quince
    # días: el «en cualquier tiempo» que se acababa de declarar se perdía en el
    # camino, y por eso hacía falta corregirlo después escribiendo sobre
    # `c.oportuna` —una @property de sólo lectura—, que reventaba con
    # AttributeError. Se pasa el plazo que es, y el cómputo sabe qué hacer con
    # el cero: `sin_plazo`, sin vencimiento y sin extemporaneidad posible.
    _plazo_computo = 0 if _pl["en_cualquier_tiempo"] else (e.plazo or 15)
    c = f0.computar(e.notificacion, e.presentacion, e.regla_surtimiento,
                    _plazo_computo, e.responsable,
                    getattr(e, "dias_inhabiles_extra", None))
    avisos.extend(c.avisos)
    if c.oportuna is False:
        avisos.append("EL CÓMPUTO DA EXTEMPORÁNEA. Compruébalo antes de seguir: "
                      "si es correcto, el asunto no se resuelve en el fondo.")

    # ── La autoridad, leída del acto ─────────────────────────────────────
    # No se le pregunta al secretario lo que está en el documento que ya subió.
    # Cada campo que se le pide y podría deducirse es un minuto suyo y una
    # ocasión de equivocarse: el adelanto vale por lo que le ahorra.
    if not (e.responsable or "").strip():
        import fase_autoridad as _fa
        # EL TIPO FILTRA QUÉ ÓRGANO PUEDE SER. Sin él, en la queja se quedaba
        # con el juez del juicio natural que el auto recurrido nombra por
        # dentro, en vez de con el Juzgado de Distrito que lo dictó.
        leida = _fa.de_texto(texto_acto, e.tipo_asunto)
        if leida:
            e.responsable = leida
            print(f"   ⚖️ autoridad responsable leída del acto: «{leida[:70]}»")
        else:
            avisos.append(
                "No se pudo leer la autoridad responsable del acto reclamado, y "
                "sin ella la competencia, los efectos y el resolutivo salen con "
                "hueco. Escríbela en el encargo.")

    # ── Las constancias, acotadas ────────────────────────────────────────
    # Un expediente son cien mil caracteres y el prompt no los aguanta. Cuando
    # no cabe entero se conserva lo NORMATIVO —las cláusulas, los preceptos del
    # reglamento—, que es lo que un acervo público no puede darnos; el relato
    # de los hechos ya viene por el acto reclamado y por la demanda.
    autos = _acotar_autos(texto_autos)

    # ── Fases 1-3 — lectura ──────────────────────────────────────────────
    # La ficha de partes se hace AQUÍ, con los documentos delante, y viaja al
    # estudio. Sin ella el redactor resuelve los sujetos por proximidad, y en un
    # juicio con reconvención y tercero interesado la proximidad miente.
    # LA ESTRUCTURA ARRANCABA A CIEGAS, Y ESO SÍ SE PERDÍA. El comentario que
    # había aquí decía que corriéndola en paralelo con las fases «no se pierde
    # nada, porque la competencia y la existencia sólo necesitan los datos del
    # encargo». Es cierto de la competencia; no de los RESULTANDOS, que también
    # los escribe esta llamada y que tienen que individualizar el acto y nombrar
    # al tercero interesado. Sin el acto ni la ficha de partes delante, el
    # modelo no podía más que la perífrasis, y salía:
    #
    #     «promovió demanda de amparo CONTRA EL ACTO RECLAMADO PRECISADO EN LOS
    #      ANTECEDENTES»
    #     «LA PERSONA A QUIEN RESULTA TAL CARÁCTER fue emplazada»
    #
    # Y con ella se llevaba el {expediente} del considerando SEGUNDO, que se lee
    # de los resultandos ya escritos: la evasión y el asterisco eran el mismo
    # defecto. Lo señaló David como dos hallazgos separados; es uno.
    #
    # SE CONSERVA EL PARALELO donde de verdad lo hay: la ficha de partes y la
    # estructura van encadenadas —la segunda necesita la primera— pero las dos
    # juntas siguen corriendo a la vez que las fases de lectura, así que la
    # espera sigue siendo la del más lento y no la suma.
    async def _partes_y_estructura():
        _p = await fpartes.fichar(cliente, texto_acto, texto_conceptos,
                                  e.tipo_asunto)
        _est = None
        if (e.modo or "").lower() == "generado":
            import documento_generado as _dg
            _est = await _dg.redactar_estructura(
                cliente, _datos_estructura(e, acto=texto_acto, partes=_p))
        return _p, _est

    with cronometrar("fases1-3+partes+estructura"):
        f, (partes, estructura_previa) = await asyncio.gather(
            f123.correr(cliente, texto_acto, texto_conceptos, e.es_recurso,
                        e.tipo_asunto),
            _partes_y_estructura())
    avisos.extend(f.avisos)
    avisos.extend(partes.avisos)

    # ── Fase 7 — el documento ────────────────────────────────────────────
    relleno = ens.Relleno(
        encabezado=e.encabezado, numero_asunto=e.numero, quejoso=e.quejoso,
        magistrado=e.magistrado, secretario=e.secretario,
        oportunidad=f0.parrafo_oportunidad(c),
        antecedentes=f.parrafos_antecedentes(),
        resumen_acto=f.parrafos_acto(),
        resumen_conceptos=f.parrafos_conceptos(),
        problemas=f.parrafos_problemas(),
        presentacion=f0.fecha_en_letra(e.presentacion)
                     if getattr(e, 'presentacion', None) else '',
        responsable=getattr(e, 'responsable', '') or '',
        es_recurso=e.es_recurso,
    )
    estructura = None
    if (e.modo or "").lower() == "generado":
        with cronometrar("estructura+docx"):
            # YA NO ES UNA TAREA. La estructura se resuelve arriba, encadenada
            # tras la ficha de partes; aquí llega hecha.
            estructura = estructura_previa
            ruta, av_gen, estructura = await _componer_generado(
                cliente, e, relleno, c, ruta_salida,
                estructura_previa=estructura, acto=texto_acto, partes=partes)
        avisos.extend(av_gen)
    else:
        with cronometrar("ensamblado"):
            ruta = ens.ensamblar(e.plantilla, relleno, ruta_salida)
            # El documento se lee antes de entregarlo: lo que quedó de la
            # plantilla no se ve leyendo por encima, se ve contándolo.
            avisos.extend(ens.residuo_de_plantilla(ruta, e.numero, e.plantilla))

    # Las constancias cuelgan de las fases porque son lo único que se serializa
    # entero al guardar la sesión: colgarlas del Resultado las perdería en
    # cuanto la petición siguiente cayera en el otro worker de gunicorn.
    f.autos = autos
    # LOS TEXTOS DE ORIGEN, para poder comprobar después que nada del proyecto
    # viene de fuera del asunto. Se guarda un extracto: comprobar la
    # contaminación no justifica duplicar el expediente entero en la sesión.
    f.fuentes = [(texto_acto or "")[:120000], (texto_conceptos or "")[:120000]]
    if autos:
        print(f"   📁 constancias del expediente: {len(autos)} caracteres")

    # EL ADELANTO TAMBIÉN SE COMPRUEBA. La revisión de contaminación sólo
    # corría en `_terminar()`, es decir, al final del camino largo: quien pide
    # únicamente el adelanto —que es la mayoría— se llevaba los resultandos con
    # sus nombres, expedientes y cantidades sin que nadie los contrastara con
    # los documentos que subió. Y el adelanto es precisamente donde van los
    # datos duros del asunto: quién promovió, contra qué, cuándo y por cuánto.
    _r0 = Resultado(ruta=ruta, computo=c, fases=f, encargo=e, partes=partes)
    for _a in _revisar_contaminacion(_r0, e):
        if _a not in avisos:
            avisos.append(_a)

    return Resultado(ruta=ruta, computo=c, fases=f, encargo=e, partes=partes,
                     huecos=ens.huecos_pendientes(ruta), avisos=avisos,
                     estructura=estructura)


# ═══════════════════════════════════════════════════════════════════════════
# La segunda mitad: el criterio del secretario entra AQUÍ y sólo aquí
# ═══════════════════════════════════════════════════════════════════════════
#
# El proceso se parte en dos a propósito, porque así es como David lo describió:
# la máquina lee y ordena, él decide, la máquina redacta la demostración. Entre
# `generar()` y `resolver()` hay una persona, y ese es el punto del diseño.
#
#   generar()   →  adelanto con los problemas jurídicos planteados
#   consultar() →  lo que el acervo tiene sobre esos problemas, para que decida
#   resolver()  →  la sentencia, con su criterio dentro


async def consultar(qdrant, embed_juris, embed_leyes,
                    r: Resultado) -> f6.Material:
    """Lo que el acervo dice sobre los problemas del caso.

    Se le enseña ANTES de pedirle el criterio: decidir el sentido sin ver la
    jurisprudencia obligatoria del tema es exactamente el error que este
    utillaje existe para evitar.
    """
    # Los problemas de la Fase 3 son DICCIONARIOS —pregunta, resolvió, combate,
    # impedimento—, no cadenas. Con datos de prueba sintéticos nunca se notó;
    # con el primer caso real, `'dict' object has no attribute 'strip'`.
    problemas = ([r.fases.problema_global] if r.fases.problema_global else [])
    for p in (r.fases.problemas or []):
        pregunta = p.get("pregunta", "") if isinstance(p, dict) else str(p)
        if pregunta:
            problemas.append(pregunta)
        # EL ACERVO SE CONSULTA POR LAS DOS. El impedimento técnico es una
        # cuestión jurídica que hay que fundar —la inoperancia se razona con
        # tesis, no se declara— pero consultarlo SÓLO a él inclinaba la
        # balanza antes de que nadie decidiera: el RAG volvía cargado de
        # material para inoperar y vacío de material para entrar al fondo. Un
        # buscador que sólo busca razones para no contestar acaba
        # encontrándolas.
        if isinstance(p, dict):
            for _clave, _pre in (("impedimento", "inoperancia"),
                                 ("apoyo", "sustento")):
                _x = p.get(_clave)
                if isinstance(_x, dict) and _x.get("explicacion"):
                    problemas.append(
                        f"¿{str(_x.get('motivo') or _pre).capitalize()}: "
                        f"{_x['explicacion']}?")
    coleccion = (r.encargo.coleccion_estatal if r.encargo else "") or None
    # LA LEY DEL ESTADO NO PINTA NADA EN UN LABORAL FEDERAL. En el ADL 382/2024
    # —IMSS contra un enfermero, ante la Junta Federal— el marco jurídico salió
    # citando el artículo 142 de la LEY ORGÁNICA MUNICIPAL DE QUERÉTARO, que
    # regula el recurso de inconformidad contra multas y licencias de comercio,
    # y el Código de Procedimientos Civiles del Estado, que se traía sólo para
    # decir que no rige. Una ejecutoria no cita leyes impertinentes para
    # explicar que no aplican.
    #
    # Se busca en el acervo estatal sólo cuando la controversia puede regirse
    # por ley local. En laboral la rige la Ley Federal del Trabajo, salvo el
    # burocrático estatal, que se reconoce porque el patrón es el propio Estado
    # o un municipio.
    if fp_materia(r.encargo) == "laboral" and not _burocratico_estatal(r):
        coleccion = None

    # EL SONDEO DE PRECEDENTE VA EN PARALELO al material. Cuesta menos de dos
    # segundos —se mide— y responde una pregunta que hasta ahora nadie hacía:
    # cómo resolvieron otros colegiados este mismo problema. Se le enseña al
    # redactor junto con el material, pero SEPARADO de él, porque un precedente
    # de otro tribunal no funda: orienta.
    material, sondeo = await asyncio.gather(
        f6rag.material_del_caso(qdrant, embed_juris, embed_leyes,
                                problemas, coleccion, fp_materia(r.encargo)),
        _sondear_precedente(qdrant, embed_leyes, r, problemas))
    material.sondeo = sondeo
    material.materia = fp_materia(r.encargo)
    # Y EL TIPO, para que la prosa del estudio nombre a las partes con las
    # figuras de ESTE recurso y no con las del amparo directo.
    material.tipo_asunto = r.encargo.tipo_asunto
    material.entidad = _entidad_de(coleccion)
    if sondeo is not None:
        for a in (sondeo.avisos or []):
            if a not in r.avisos:
                r.avisos.append(a)
    return material


_ENTIDADES = {
    "aguascalientes": "Aguascalientes", "bajacalifornia": "Baja California",
    "bajacaliforniasur": "Baja California Sur", "campeche": "Campeche",
    "chiapas": "Chiapas", "chihuahua": "Chihuahua", "cdmx": "Ciudad de México",
    "coahuila": "Coahuila", "colima": "Colima", "durango": "Durango",
    "guanajuato": "Guanajuato", "guerrero": "Guerrero", "hidalgo": "Hidalgo",
    "jalisco": "Jalisco", "edomex": "Estado de México", "mexico": "Estado de México",
    "michoacan": "Michoacán", "morelos": "Morelos", "nayarit": "Nayarit",
    "nuevoleon": "Nuevo León", "oaxaca": "Oaxaca", "puebla": "Puebla",
    "queretaro": "Querétaro", "quintanaroo": "Quintana Roo",
    "sanluispotosi": "San Luis Potosí", "sinaloa": "Sinaloa", "sonora": "Sonora",
    "tabasco": "Tabasco", "tamaulipas": "Tamaulipas", "tlaxcala": "Tlaxcala",
    "veracruz": "Veracruz", "yucatan": "Yucatán", "zacatecas": "Zacatecas",
}


def _entidad_de(coleccion: str) -> str:
    """«leyes_queretaro» → «Querétaro». Sin adivinar: si no la conozco, vacío."""
    c = (coleccion or "").strip().lower().replace("leyes_", "").replace("_", "")
    return _ENTIDADES.get(c, "")


# Cuánto del expediente cabe. Generoso: es material que no está en ningún otro
# sitio y su ausencia ya costó un asunto.
MAX_AUTOS = 24000
_RX_CLAUSULA = re.compile(
    r"(cl[áa]usula|art[íi]culo|reglamento|contrato\s+colectivo|"
    r"fracci[óo]n|condiciones\s+generales\s+de\s+trabajo)", re.I)


def _acotar_autos(texto: str) -> str:
    """El expediente, y si no cabe, su parte normativa."""
    t = " ".join((texto or "").split())
    if len(t) <= MAX_AUTOS:
        return t
    # Se trocea en párrafos y se conservan los que traen norma, en su orden.
    trozos = re.split(r"(?<=[.;])\s+(?=[A-ZÁÉÍÓÚ])", t)
    con_norma = [x for x in trozos if _RX_CLAUSULA.search(x)]
    fuera, total = [], 0
    for x in (con_norma or trozos):
        if total + len(x) > MAX_AUTOS:
            break
        fuera.append(x)
        total += len(x)
    return " ".join(fuera)


_RX_BUROCRATICO = re.compile(
    r"trabajadores?\s+al\s+servicio\s+del\s+estado|burocr[áa]tic|"
    r"tribunal\s+de\s+conciliaci[óo]n\s+y\s+arbitraje\s+del\s+estado|"
    r"ley\s+de\s+los\s+trabajadores\s+del\s+estado", re.I)


def _burocratico_estatal(r) -> bool:
    """¿Es un laboral que SÍ se rige por ley local? Sólo el burocrático."""
    f = getattr(r, "fases", None)
    texto = " ".join(str(getattr(f, k, "") or "") for k in
                     ("antecedentes", "resumen_acto", "problema_global"))
    e = getattr(r, "encargo", None)
    texto += " " + str(getattr(e, "encabezado", "") or "")
    texto += " " + str(getattr(e, "responsable", "") or "")
    return bool(_RX_BUROCRATICO.search(texto))


def fp_materia(e) -> str:
    """La materia del asunto. La DECLARADA manda sobre la deducida.

    David: «en función de la materia el RAG es selectivo (como en laboral) y
    eso lo determina un campo seleccionado por el secretario». Hasta ahora la
    materia se deducía del nombre del tribunal y del encabezado, y eso acierta
    casi siempre —un colegiado de trabajo no ve otra cosa— pero falla justo
    donde más cuesta: en un tribunal MIXTO «en materias administrativa y
    civil», donde el nombre no decide, y en un asunto laboral que llega a un
    colegiado administrativo, que es cuando el silo laboral hace falta.

    La deducción se conserva como respaldo: quien no declare materia sigue
    teniendo el comportamiento de siempre.
    """
    declarada = str(getattr(e, "materia", "") or "").strip().lower()
    if declarada:
        return declarada
    try:
        import fase_precedente as fp
    except Exception:
        return ""
    return fp.materia_de(getattr(e, "encabezado", ""),
                         getattr(e, "tribunal", ""),
                         getattr(e, "tipo_asunto", ""))


async def _sondear_precedente(qdrant, embed, r: Resultado, problemas: list):
    """El acervo de colegiados, sondeado por el problema, no por el escrito.

    El fraseo de la demanda envenena la búsqueda —es el fallo ya medido de
    HyDE—, así que se sondea con el problema jurídico que la fase 3 normalizó
    por concepto. Si algo falla, se devuelve None y el redactor sigue: el
    sondeo mejora la sentencia, no la condiciona.

    EL EMBEBEDOR ES EL DE 1536, NO EL DE JURISPRUDENCIA. Aquí hay dos modelos
    distintos conviviendo: `jurisprudencia_nacional_v3` se indexó con
    text-embedding-3-large (3072 dimensiones) y `sentencias_holdings` con
    text-embedding-3-small (1536). Le pasé el de jurisprudencia por costumbre y
    Qdrant devolvió «expected dim: 3072, got 1536»; como el sondeo captura sus
    propios errores para no tumbar la sentencia, en producción se habría visto
    sencillamente como «no hay precedentes». Lo cazó la prueba local, que es
    justamente para lo que sirve correrla antes de desplegar.
    """
    try:
        import fase_precedente as fp
    except Exception:
        return None
    if not (qdrant and problemas):
        return None
    e = r.encargo
    # NO SE PIDEN: se leen. La materia está en el encabezado y en el nombre del
    # tribunal; el circuito, en ese mismo nombre.
    materia = fp_materia(e)
    if not materia:
        return None
    # UNO POR PROBLEMA, EN PARALELO. Sondeaba `problemas[0]` y ya: el acervo
    # decía cómo se resuelve la cuestión principal y callaba sobre las demás,
    # que es donde el secretario más agradece la señal —un accesorio que el
    # 90% de los tribunales declara inoperante se resuelve en dos renglones—.
    # Son N búsquedas contra la misma colección; corriendo a la vez, la espera
    # es la de la más lenta.
    _circ = fp.circuito_de(getattr(e, "tribunal", ""))
    _txt = [p if isinstance(p, str) else str((p or {}).get("pregunta") or p)
            for p in problemas]
    try:
        _todos = await asyncio.gather(*[
            fp.sondear(qdrant, embed, t, materia, circuito=_circ)
            for t in _txt], return_exceptions=True)
    except Exception as exc:
        print(f"   ⚠️ sondeo de precedente omitido: {exc}")
        return None
    _sondeos = [x if not isinstance(x, BaseException) else None for x in _todos]
    if not any(_sondeos):
        return None
    s = next(x for x in _sondeos if x)
    # Los demás viajan colgados del primero: `Material.sondeo` es lo que lee el
    # estudio y no se le cambia la forma, pero la predicción de cada problema
    # tiene que llegar a la pantalla.
    s.por_problema = [
        {"problema": t, "prediccion": fp.prediccion(x) if x else {}}
        for t, x in zip(_txt, _sondeos)]
    # SE DEJA CONSTANCIA AUNQUE VAYA BIEN. Hasta ahora este paso sólo hablaba
    # cuando fallaba, y tras desplegarlo no había manera de saber desde los
    # registros si había corrido: el silencio significaba «no falló», no
    # «funcionó». Es el mismo defecto que el aviso que nadie veía. Una línea.
    _pred = [d["prediccion"].get("frase", "—") for d in s.por_problema]
    print(f"   ⚖️ jurimetría por problema: " + " · ".join(_pred[:4]))
    print(f"   ⚖️ precedente[{materia}]: "
          f"{sum(s.distribucion.values())} sentencias del tema · "
          f"{len(s.moldes)} moldes · {len(s.razonados)} con razón escrita"
          + (f" · avisos: {len(s.avisos)}" if s.avisos else ""))
    return s


async def resolver(cliente, r: Resultado, criterios: list[f6.Criterio],
                   material: f6.Material, ruta_salida: str,
                   marco: str = "", qdrant=None, contexto: str = "") -> Resultado:
    """La sentencia: el mismo documento, ahora con el estudio de fondo dentro.

    Se REENSAMBLA desde la plantilla en vez de editar el adelanto, porque el
    ensamblador trabaja sobre los formatos de la plantilla original y aplicarlo
    dos veces sobre su propia salida duplica bloques.
    """
    if r.encargo is None:
        raise ValueError("El resultado no trae el encargo: no se puede reensamblar.")
    e = r.encargo

    # EL MARCO SE ESCRIBE A LA VEZ QUE EL ESTUDIO. Son dos llamadas
    # independientes —la del marco sólo mira el material constitucional, la del
    # estudio mira el caso— y ponerlas en paralelo hace que el marco no cueste
    # un segundo de espera. Que APAREZCA ya no depende de que el modelo del
    # estudio se acuerde de escribirlo: lo coloca el compositor.
    tarea_marco = None
    if (e.modo or "").lower() == "generado" and (marco or "").strip():
        import documento_generado as _dg2
        tarea_marco = asyncio.create_task(_dg2.redactar_marco(
            cliente, marco,
            [p for p in (r.fases.problemas or [])], e.es_recurso, e.tipo_asunto))

    with cronometrar("estudio de fondo"):
        estudio, advertencias, avisos = await f6.redactar(
            cliente, r.fases.resumen_acto, r.fases.resumen_conceptos,
            criterios, material, e.es_recurso, r.partes, marco, contexto)

    return await _terminar(cliente, r, e, criterios, material, estudio,
                           advertencias, avisos, tarea_marco, ruta_salida, qdrant, marco)


async def resolver_en_vivo(cliente, r: Resultado, criterios: list[f6.Criterio],
                           material: f6.Material, ruta_salida: str,
                           marco: str = "", qdrant=None, contexto: str = ""):
    """La sentencia, viéndose escribir. Rinde trozos y, al final, el Resultado."""
    e = r.encargo
    avisos: list[str] = []
    tarea_marco = None
    if (e.modo or "").lower() == "generado" and (marco or "").strip():
        import documento_generado as _dg3
        tarea_marco = asyncio.create_task(_dg3.redactar_marco(
            cliente, marco, [p for p in (r.fases.problemas or [])],
            e.es_recurso, e.tipo_asunto))

    estudio = advertencias = ""
    t0 = _time.perf_counter()
    async for paso in f6.redactar_en_vivo(
            cliente, r.fases.resumen_acto, r.fases.resumen_conceptos,
            criterios, material, e.es_recurso, r.partes, marco, contexto):
        if paso.get("tipo") == "texto":
            yield paso
        else:
            estudio = paso.get("estudio", "")
            advertencias = paso.get("advertencias", "")
            avisos.extend(paso.get("avisos", []))
    TIEMPOS["estudio de fondo"] = round(_time.perf_counter() - t0, 1)

    yield {"tipo": "componiendo"}
    res = await _terminar(cliente, r, e, criterios, material, estudio,
                          advertencias, avisos, tarea_marco, ruta_salida, qdrant, marco)
    yield {"tipo": "listo", "resultado": res}


def _revisar_contaminacion(r, e) -> list:
    """Que nada del proyecto sea de otro asunto.

    David: «asegúrate de que los formatos de salida no estén contaminados con
    datos que no correspondan al asunto que proyecta el secretario». Se
    comprueba sin modelo: todo nombre propio, número de expediente y cantidad
    del proyecto debe estar en los documentos que él subió o en lo que tecleó.
    Lo que no esté viene de otra parte.
    """
    try:
        import contaminacion as _c
    except Exception:
        return []
    fuentes = list(getattr(getattr(r, "fases", None), "fuentes", []) or [])
    autos = str(getattr(getattr(r, "fases", None), "autos", "") or "")
    if autos:
        fuentes.append(autos)
    enc = {k: getattr(e, k, "") for k in
           ("numero", "encabezado", "quejoso", "responsable", "magistrado",
            "secretario", "tribunal", "ciudad")}
    try:
        with open(r.ruta, "rb"):
            pass
        import docx as _dx
        d = _dx.Document(r.ruta)
        texto = "\n".join(p.text for p in d.paragraphs)
        for tb in d.tables:
            for fila in tb.rows:
                texto += "\n" + " | ".join(c.text for c in fila.cells)
    except Exception:
        return []
    return _c.revisar(texto, fuentes, enc)


async def _terminar(cliente, r, e, criterios, material, estudio,
                    advertencias, avisos, tarea_marco, ruta_salida, qdrant=None,
                    marco: str = ""):
    """De la salida del modelo al documento entregado.

    Vive fuera de `resolver()` porque la versión en vivo hace exactamente lo
    mismo cuando el flujo termina, y tener dos copias de esto es tener dos
    sitios donde se rompe la congruencia.
    """
    relleno = ens.Relleno(
        encabezado=e.encabezado, numero_asunto=e.numero, quejoso=e.quejoso,
        magistrado=e.magistrado, secretario=e.secretario,
        oportunidad=f0.parrafo_oportunidad(r.computo),
        antecedentes=r.fases.parrafos_antecedentes(),
        resumen_acto=r.fases.parrafos_acto(),
        resumen_conceptos=r.fases.parrafos_conceptos(),
        problemas=r.fases.parrafos_problemas(),
        estudio=f6.parrafos(estudio),
        tesis=material.tesis,
        normas=material.normas,
        calificaciones=[c.sentido for c in criterios],
        presentacion=f0.fecha_en_letra(e.presentacion)
                     if getattr(e, 'presentacion', None) else '',
        responsable=getattr(e, 'responsable', '') or '',
        es_recurso=e.es_recurso,
    )
    # ═══ LOS ARTÍCULOS QUE DE VERDAD CITÓ ══════════════════════════════
    # Antes se buscaban por parecido ANTES de escribir, cuatro por problema, y
    # si el estudio acababa citando otros se quedaban sin texto y sin nota al
    # pie. Ahora se lee qué citó y se piden ESOS por número: `articulo_num`
    # está indexado en las 34 colecciones de leyes. No es adivinar lo que hará
    # falta, es traer lo que hizo falta.
    try:
        if qdrant is None:
            raise RuntimeError("sin cliente de Qdrant")
        import fase_normas as _fn
        with cronometrar("artículos citados"):
            # EL FUERO DE LA AUTORIDAD decide si el acervo del estado siquiera
            # se consulta. Se lee de lo que ya se sabe del asunto —la autoridad
            # responsable y el acto—, no del estudio, que aún puede no
            # nombrarla.
            _quien = " ".join(str(x) for x in (
                getattr(e, "autoridad_responsable", ""),
                getattr(e, "acto_reclamado", ""),
                getattr(material, "acto_reclamado", ""),
                " ".join(getattr(material, "autoridades", None) or []),
            ) if x)
            _fed = _fn.autoridad_es_federal(_quien)
            if _fed:
                print("   ⚖️ autoridad del fuero FEDERAL: no se consulta el "
                      "acervo estatal salvo que la cita nombre una ley local")
            _extra = await _fn.recuperar(
                qdrant, estudio,
                (e.coleccion_estatal or "") if hasattr(e, "coleccion_estatal") else "",
                fuero_federal=_fed)
        if _extra:
            _ya = {(str(n_.get("articulo")), str(n_.get("cuerpo_legal") or
                                                 n_.get("fuente") or ""))
                   for n_ in (material.normas or [])}
            nuevos = [x for x in _extra
                      if (x["articulo"], x["cuerpo_legal"]) not in _ya]
            material.normas = list(material.normas or []) + nuevos
            print(f"   ⚖️ artículos citados recuperados: {len(nuevos)} nuevos "
                  f"de {len(_extra)} hallados")
    except Exception as _ea:
        avisos.append(f"No se pudieron recuperar los artículos citados: {_ea}")

    marco_escrito = ""
    if tarea_marco is not None:
        with cronometrar("marco escrito"):
            try:
                marco_escrito = await tarea_marco
                import documento_generado as _dg4
                avisos.extend(_dg4.revisar_marco(marco_escrito, marco or ""))
            except Exception as _em:
                avisos.append(f"No se pudo escribir el marco jurídico: {_em}")
    if (e.modo or "").lower() == "generado":
        with cronometrar("recomposición"):
            ruta, av_gen, _est = await _componer_generado(
                cliente, e, relleno, r.computo, ruta_salida,
                estructura_previa=getattr(r, "estructura", None),
                # Las constancias ya leídas viajan con las fases; si la
                # estructura hubiera que rehacerla, que no sea a ciegas.
                acto=(getattr(getattr(r, "fases", None), "fuentes", []) or [""])[0],
                partes=getattr(r, "partes", None),
                marco_escrito=marco_escrito)
        avisos.extend(av_gen)
    else:
        with cronometrar("ensamblado"):
            ruta = ens.ensamblar(e.plantilla, relleno, ruta_salida)
    _, aviso_efectos = ens.formula_resolutivo(relleno.calificaciones)
    if aviso_efectos:
        avisos.append(aviso_efectos)
    # Deduplicado: el aviso del nombre se dispara una vez por párrafo donde
    # aparece, y el secretario no necesita leer tres veces lo mismo.
    for a in ens.avisos_ensamblado:
        if a not in avisos:
            avisos.append(a)
    for a in ens.residuo_de_plantilla(ruta, e.numero, e.plantilla):
        if a not in avisos:
            avisos.append(a)
    # LA CONGRUENCIA VA LA PRIMERA. Es el único aviso de esta lista que no
    # describe algo mejorable sino algo que no se puede firmar, y el secretario
    # tiene que verlo antes que los otros trece.
    incongruente = ens.revisar_congruencia(ruta, relleno.calificaciones,
                                           e.tipo_asunto)
    for a in reversed(incongruente):
        if a not in avisos:
            avisos.insert(0, a)

    # LA CALIDAD DEL FONDO, MEDIDA SOBRE EL DOCUMENTO ENTREGADO. No es una
    # opinión: son las cinco medidas que salieron de contar los defectos de los
    # engroses reales —densidad, exhaustividad, congruencia interna, promesa
    # cumplida y duplicación—. El secretario ve el número, no un adjetivo.
    try:
        import calidad_estudio as _ce
        from docx import Document as _Doc
        _txt = "\n".join(p.text for p in _Doc(ruta).paragraphs)
        _m = _ce.medir(_txt)
        _d = _m["densidad"]
        if _d["palabras"] > 400 and _d["pct_propio"] < 0.70:
            avisos.append(
                f"EL ESTUDIO VIVE DE LA CITA: sólo el {_d['pct_propio']*100:.0f}% "
                f"es razonamiento propio ({_d['transcritas']} palabras "
                f"transcritas de {_d['palabras']}). La referencia de los "
                f"engroses es el 45%, así que esto no es un desastre —pero el "
                f"objetivo es el 70%.")
        for _r in _m["remisiones_rotas"]:
            avisos.append(f"REMISIÓN A UN APARTADO QUE NO EXISTE: «{_r}». Es el "
                          f"defecto que se caza leyendo el resolutivo en voz alta.")
        if _m["procedencia_contradice"]:
            avisos.append("LA PROCEDENCIA CONTRADICE EL FALLO: dice improcedente "
                          "en un asunto que se resuelve por el fondo.")
        if _m["promesa"]["rota"]:
            avisos.append("SE PROMETIÓ NO TRANSCRIBIR Y SE TRANSCRIBIÓ: el "
                          "documento dice que es innecesario reproducir el acto "
                          "y luego lo reproduce.")
        _dup = _ce.duplicacion_interna(_ce.estudio_de(_txt))
        if _dup:
            avisos.append(
                f"{len(_dup)} PASAJE(S) REPETIDO(S) dentro del estudio: «"
                f"{_dup[0][:90]}…». Un pasaje repetido no refuerza; delata que "
                f"se escribió por trozos.")
        # EL MARCO NORMATIVO DE OTRA VÍA. Acotado al considerando de
        # competencia y al de procedencia: fuera de ahí una cita de otra ley
        # puede ser legítima —un criterio análogo, una remisión— y prohibirla
        # empobrecería el estudio.
        import tipos_asunto as _ta_v
        for _pat, _porque in _ta_v.preceptos_ajenos(e.tipo_asunto, _txt):
            avisos.insert(0, f"PRECEPTO DE OTRA VÍA EN EL MARCO: {_porque}.")

        # EL RESULTANDO QUE NO INFORMA. Va el primero de la lista porque de ese
        # dato dependen el inciso del 97, la rama del 93 y que quien lea el
        # proyecto sepa de qué va el asunto.
        for _q, _d in _ta_v.resultando_evasivo(_txt, e.tipo_asunto):
            avisos.insert(0, f"{_q}" + (f": «…{_d[:130]}…»" if _d else "."))

        # EL LINTER. Frases cortadas, comillas huérfanas y el marcador genérico
        # donde debería ir el nombre.
        try:
            import linter_juridico as _lj
            for _que, _donde in _lj.revisar(_txt):
                avisos.append(f"SINTAXIS — {_que}"
                              + (f": «…{_donde[:110]}…»" if _donde else ""))
        except Exception as _el:
            print(f"   ⚠️ linter no disponible: {_el}")

        for _e in _ce.estadistica_en_el_texto(_txt):
            avisos.append(f"LA CIFRA DEL ACERVO SE COLÓ EN LA SENTENCIA: «{_e[:110]}». "
                          f"El criterio no se vota: quítala.")
        _ex = _m["exhaustividad"]
        if _ex.get("contesta_todo") is False:
            avisos.append(
                f"QUEDAN PLANTEAMIENTOS SIN RESPUESTA: se anuncian "
                f"{_ex['planteamientos_anunciados']} y se emiten "
                f"{_ex['calificaciones_emitidas']} calificaciones, sin decir que "
                f"se estudian conjuntamente. Es omisión de estudio.")
    except Exception as _ec:
        print(f"   ⚠️ no se pudo medir la calidad del estudio: {_ec}")

    # NADA DEL PROYECTO PUEDE SER DE OTRO ASUNTO. Se comprueba sobre el .docx ya
    # escrito, que es lo que el secretario va a leer, y contra los documentos
    # que él subió.
    # SE COMPRUEBA DOS VECES —el adelanto y la sentencia—, así que hay que
    # unificar lo que digan las dos. Sin esto el proyecto salía con dos avisos
    # que NO PUEDEN SER CIERTOS A LA VEZ: uno afirmaba que la cantidad
    # «no consta en las fuentes» y el otro que las fuentes no se pueden leer.
    # Quien lee eso no sabe a cuál hacer caso, y con razón.
    r_final = Resultado(ruta=ruta, computo=r.computo, fases=r.fases, encargo=e)
    for a in _revisar_contaminacion(r_final, e):
        if a not in avisos:
            avisos.append(a)

    _todos = list(r.avisos) + avisos
    _vistos, _limpios = set(), []
    for a in _todos:
        _k = str(a)[:70]                    # el mismo aviso con un nombre más
        if _k in _vistos:                   # no es un aviso nuevo
            continue
        _vistos.add(_k)
        _limpios.append(a)
    # LA AUSENCIA DE PRUEBA NO ES PRUEBA, y tampoco al revés: si una pasada SÍ
    # pudo comprobar y encontró algo, el «no se pudo comprobar» de la otra
    # sobra y sólo resta credibilidad a lo que sí se halló.
    _hallo = any(str(a).startswith(("NOMBRES QUE NO", "EXPEDIENTES QUE NO",
                                    "CANTIDADES QUE NO")) for a in _limpios)
    if _hallo:
        _limpios = [a for a in _limpios
                    if not str(a).startswith("No se pudo comprobar la contaminación")]

    # LA RAMA DE LA REVISIÓN, POR EL MISMO MOTIVO. El adelanto se compone antes
    # de que exista el estudio, así que no puede saber qué resolvió el a quo y
    # anota la rama «sin_determinar». Cuando la sentencia sí lo determina, el
    # proyecto salía con LOS DOS avisos: «el a quo no consta qué resolvió» y,
    # dos líneas más abajo, «el a quo sobresee». El segundo es el bueno —tiene
    # el estudio delante— y el primero sólo siembra duda sobre él.
    if any("RESOLUTIVO DE REVISIÓN, rama" in str(a)
           and "sin_determinar" not in str(a) for a in _limpios):
        _limpios = [a for a in _limpios
                    if "RESOLUTIVO DE REVISIÓN, rama «sin_determinar»" not in str(a)
                    and not str(a).startswith("NO CONSTA QUÉ RESOLVIÓ EL JUZGADO")]

    return Resultado(ruta=ruta, computo=r.computo, fases=r.fases, encargo=e,
                     partes=r.partes, estudio=estudio, advertencias=advertencias,
                     huecos=ens.huecos_pendientes(ruta),
                     avisos=_limpios)




# ── Los ANTECEDENTES ya se generan ────────────────────────────────────────
#
# Medidos sobre 199 apartados reales antes de escribir una línea de prompt:
# 645 palabras en 17 párrafos de 37 —crónica de trámite, frases cortas—, con
# verbos de procedimiento (dictó 186, admitió 112, interpuso 82) y arranques
# fijos («Por auto de», «En proveído de», «Seguido el juicio»).
#
# NO son el resumen del acto, aunque salgan de la misma sentencia: el resumen
# cuenta lo que la responsable RESOLVIÓ y por qué; los antecedentes, lo que
# PASÓ en el juicio de origen. Uno es razonamiento, el otro es crónica. Por eso
# se piden con otro prompt y sobre el documento ENTERO —el recorte del resumen
# se queda con el estudio de fondo, donde el trámite ya no está—.


def _datos_estructura(e: Encargo, antecedentes: str = "", acto: str = "",
                     partes=None) -> dict:
    """Lo que la estructura necesita.

    Decía «TODO sale del encargo» y por eso se lanzaba a ciegas. Del encargo
    salen la competencia y la existencia; los RESULTANDOS necesitan el acto
    —para individualizar la sentencia recurrida con su fecha y su expediente—
    y la ficha de partes —para nombrar al tercero interesado en vez de escribir
    «la persona a quien resulta tal carácter»—.
    """
    import fase0_oportunidad as _f0
    _terceros = ""
    if partes is not None:
        _terceros = str(getattr(partes, "tercero_interesado", "") or "")
    return {
        # LA CABEZA DEL ACTO, que es donde se identifica: fecha, órgano,
        # expediente y toca. No el documento entero —eso ya lo leen las fases—
        # sino lo justo para individualizarlo sin adivinar.
        "acto": (acto or "")[:6000],
        "tercero": _terceros,
        # EL NÚMERO DEL PROPIO ASUNTO. Sin él, `fase_origen.numero_de` no puede
        # descartarlo y devuelve el toca de ESTE recurso como si fuera el
        # expediente de origen: el dict tenía `encabezado`, que es otra cosa.
        "numero": e.numero,
        "tribunal": e.tribunal or "",
        "ciudad": e.ciudad or "",
        "encabezado": e.encabezado,
        "quejoso": e.quejoso,
        "responsable": e.responsable or "",
        "magistrado": e.magistrado,
        "secretario": e.secretario,
        "presentacion": _f0.fecha_en_letra(e.presentacion),
        # EL TIPO DE ASUNTO, que faltaba y por eso el prompt de estructura
        # escribía siempre «una sentencia de amparo directo» y pedía identificar
        # el acto por «sala, toca y expediente» aunque fuera una queja.
        "tipo_asunto": getattr(e, "tipo_asunto", "amparo_directo"),
        "es_recurso": e.es_recurso,
        # Las dos fechas de la sesión, en letra como todo lo demás del cuerpo.
        # Vacías si el secretario no las declaró: entonces siguen en hueco.
        # SE COMPRUEBA QUE LA FECHA SE PUDO LEER, no que el campo venga lleno.
        # Escrito como estaba, un «15/03/2026» —que es como se teclea una fecha
        # en España y en México— pasaba el `if`, `_fecha_iso` devolvía None y
        # `fecha_en_letra(None)` tumbaba el adelanto con un 500 DESPUÉS de
        # treinta y cinco segundos de trabajo ya pagado. Lo introduje yo hoy.
        "fecha_lista": _f0.fecha_en_letra(_fecha_iso(getattr(e, "fecha_lista", "")))
                       if _fecha_iso(getattr(e, "fecha_lista", "")) else "",
        "fecha_sesion": _f0.fecha_en_letra(_fecha_iso(getattr(e, "fecha_sesion", "")))
                        if _fecha_iso(getattr(e, "fecha_sesion", "")) else "",
        "antecedentes": antecedentes,
    }


def _fecha_iso(x):
    """La fecha ISO del formulario, o None si no se puede leer.

    Se devuelve None —y el hueco se queda— en vez de una fecha aproximada: una
    fecha de sesión equivocada en el resultando es peor que un asterisco, que
    al menos se ve.
    """
    import datetime as _dt
    import re as _re
    t = str(x or "").strip()
    if not t:
        return None
    try:
        return _dt.date.fromisoformat(t[:10])
    except Exception:
        pass
    # «15/03/2026» y «15-03-2026», que es como se teclea una fecha aquí.
    m = _re.match(r"^(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})$", t)
    if m:
        try:
            return _dt.date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except ValueError:
            return None
    return None


async def _componer_generado(cliente, e: Encargo, relleno, computo,
                             ruta_salida: str, estructura_previa=None,
                             marco_escrito: str = "", acto: str = "",
                             partes=None):
    """El documento escrito entero. Devuelve (ruta, avisos, estructura)."""
    import documento_generado as dg
    import fase0_oportunidad as _f0

    # LA SEGUNDA LLAMADA TAMBIÉN NECESITA EL ACTO Y LAS PARTES. Se arregló la
    # de arriba y ésta se quedó igual: cuando `estructura_previa` es None
    # —porque se recompone el documento sin haber pasado por el adelanto— la
    # estructura volvía a escribirse a ciegas, y con ella la perífrasis.
    datos = _datos_estructura(e, "\n".join(relleno.antecedentes or []),
                              acto=acto, partes=partes)
    # LA ESTRUCTURA SE ESCRIBE UNA VEZ. El resolver recompone el documento
    # entero, y volver a pedirla al modelo son treinta segundos por nada: no
    # depende del estudio ni del criterio, sólo del asunto.
    est = estructura_previa or await dg.redactar_estructura(cliente, datos)
    ruta = dg.componer(
        datos, est, computo, _f0.fecha_en_letra, ruta_salida,
        antecedentes=relleno.antecedentes,
        resumen_acto=relleno.resumen_acto,
        resumen_conceptos=relleno.resumen_conceptos,
        problemas=relleno.problemas,
        estudio=relleno.estudio,
        calificaciones=relleno.calificaciones,
        tesis=relleno.tesis,
        marco_escrito=marco_escrito,
        # El tipo decide el esqueleto: los recursos no llevan «Existencia del
        # acto reclamado» y la queja hace el cómputo en prosa.
        tipo_asunto=(getattr(e, "tipo_asunto", "") or
                     ("amparo_revision" if e.es_recurso else "amparo_directo")),
        normas=getattr(relleno, "normas", None))
    avisos = list(est.avisos)
    if not e.tribunal:
        avisos.append(
            "No se indicó el TRIBUNAL que resuelve: la competencia y la "
            "fórmula de apertura salen incompletas. Es el dato que hace que "
            "esto sirva fuera de un solo circuito.")
    return ruta, avisos, est


def resumen_legible(r: Resultado) -> str:
    """Lo que se le enseña al secretario cuando termina el proceso."""
    lineas = [
        f"Adelanto generado: {r.ruta.split('/')[-1]}",
        f"Oportunidad: {'en tiempo' if r.computo.oportuna else 'EXTEMPORÁNEA'} "
        f"· plazo del {r.computo.inicio} al {r.computo.vencimiento} "
        f"({r.computo.plazo} días hábiles)",
        f"Antecedentes: {len(r.fases.antecedentes.split())} palabras "
        f"en {len(r.fases.parrafos_antecedentes())} párrafos",
        f"Resumen del acto: {len(r.fases.resumen_acto.split())} palabras",
        f"Resumen de conceptos: {len(r.fases.resumen_conceptos.split())} palabras",
        f"Problemas jurídicos: {len(r.fases.problemas)}"
        + (" (+ el global)" if r.fases.problema_global else ""),
    ]
    if r.avisos:
        lineas.append("\nAVISOS:")
        lineas += [f"  · {a}" for a in r.avisos]
    if r.huecos:
        lineas.append("\nPENDIENTE DE TU CRITERIO:")
        lineas += [f"  · {h[:100]}" for h in r.huecos]
    return "\n".join(lineas)
