"""Ensambla el adelanto sobre la plantilla REAL del secretario.

NO SE CONSTRUYE UN WORD NUEVO. Se abre el `.docx` del propio tribunal y se
rellenan sus huecos. La diferencia no es estética: en ese archivo viven la
cabecera, los estilos, los márgenes, la página de síntesis y las dos tablas de
calendario. Un documento generado desde cero se parece; éste ES el suyo.

La estructura, leída del ADA 240/2026 y confirmada por David:

    AMPARO DIRECTO ADMINISTRATIVO: 240/2026     ← número del asunto
    QUEJOSO: …
    MAGISTRADO: …                                ← ponente
    SECRETARIO DE TRIBUNAL: …
    V I S T O, para resolver …
    R E S U L T A N D O:
      PRIMERO. Presentación de la demanda        ┐
      SEGUNDO. Derechos humanos vulnerados       │
      TERCERO. Parte tercera interesada          │ machote + datos de la ficha
      CUARTO. Trámite del juicio                 │
      QUINTO. Turno del asunto                   │
      SEXTO. Verificación de la sesión remota    ┘
    C O N S I D E R A N D O:
      PRIMERO. Competencia                       ← machote por tipo de asunto
      SEGUNDO. Existencia del acto reclamado     ← machote
      TERCERO. Legitimación y oportunidad        ← FASE 0: el cómputo
      CUARTO. Sentencia reclamada y conceptos    ← «Es innecesario transcribir…»
      QUINTO. Antecedentes                       ← de la sentencia reclamada
      SEXTO. Estudio                             ← el corazón:
          · resumen de la sentencia reclamada        438 palabras, PASADO
          · resumen de los conceptos o agravios      472 palabras, PRESENTE
          · problemas jurídicos (o el global)
          · … y aquí entra el criterio del secretario
    Por lo expuesto y fundado, se resuelve:
    ÚNICO. La Justicia de la Unión ***** a …     ← el sentido lo pone él
    (página de síntesis: TEMA, OPORTUNIDAD y los dos calendarios)

CÓMO SE RELLENA SIN ROMPER EL FORMATO: cada párrafo de Word es una lista de
«runs», y el formato vive en los runs, no en el párrafo. Si se borra el párrafo
y se escribe otro, se pierde la fuente, el interlineado y la sangría. Aquí se
conserva SIEMPRE el primer run —que lleva el formato bueno—, se le cambia el
texto y se eliminan los demás. Para añadir párrafos se CLONA uno vecino, por lo
mismo.
"""

from __future__ import annotations

import copy
import re
from dataclasses import dataclass, field
from typing import Iterable, Optional

try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ImportError:                                     # pragma: no cover
    docx = None


# ═══════════════════════════════════════════════════════════════════════════
# Manipulación cuidadosa del .docx
# ═══════════════════════════════════════════════════════════════════════════

def texto_de(p) -> str:
    return "".join(n.text or "" for n in p._p.iter(qn("w:t")))


def escribir(p, texto: str) -> None:
    """Cambia el texto conservando el formato del párrafo.

    Se conserva el PRIMER run y se borran los demás: el formato —fuente,
    tamaño, negrita, sangría— vive en el run, no en el párrafo. Vaciar el
    párrafo y escribir de nuevo lo dejaría en el estilo por defecto.
    """
    runs = p.runs
    if not runs:
        p.add_run(texto)
        return
    runs[0].text = texto
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def escribir_tramos(p, tramos) -> None:
    """El párrafo con VARIOS formatos dentro: [(texto, {"bold":…, "italic":…}), …].

    `escribir()` mete todo en el primer run y hereda su formato, que es por qué
    el resolutivo salía ENTERO en negrita: el primer run de la plantilla es
    «ÚNICO.» y va en negrita. Aquí el primer run se usa como MOLDE —de él salen
    la fuente, el tamaño y el color— y se clona uno por tramo.
    """
    runs = p.runs
    if not runs:
        for texto, fmt in tramos:
            r = p.add_run(texto)
            r.bold = fmt.get("bold", False)
            r.italic = fmt.get("italic", False)
        return

    molde = runs[0]._element
    nuevos = []
    for texto, fmt in tramos:
        e = copy.deepcopy(molde)
        nuevos.append(e)
        molde.addprevious(e)
        r = Run(e, p)
        r.text = texto
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
    for r in list(p.runs):
        if r._element not in nuevos:
            r._element.getparent().remove(r._element)


# El rubro de una tesis se escribe entre comillas y EN NEGRITA; así lo hace él
# en los 5 casos del ADC 174-2026 y en todos los engroses revisados. La cita se
# reconoce por las comillas tipográficas y por ir en versales, que es como la
# publica la Corte.
_RX_RUBRO = re.compile(r"[“\"]([A-ZÁÉÍÓÚÑ][^”\"]{18,}?)[”\"]")


def tramos_con_rubro(texto: str):
    """Parte el párrafo en tramos para que el rubro salga en negrita."""
    tramos, i = [], 0
    for m in _RX_RUBRO.finditer(texto):
        if m.start() > i:
            tramos.append((texto[i:m.start()], {}))
        tramos.append((texto[m.start():m.end()], {"bold": True}))
        i = m.end()
    if not tramos:
        return None
    if i < len(texto):
        tramos.append((texto[i:], {}))
    return tramos


def clonar_tras(p, texto: str, modelo=None):
    """Un párrafo nuevo DESPUÉS de `p`, con el formato de `modelo`.

    EL MODELO NO ES UN DETALLE. Antes se clonaba siempre el párrafo ancla, y
    el ancla es el ENCABEZADO —«SEXTO. Estudio…»—, que va en negrita y con otra
    sangría. Resultado: el estudio entero salía en negrita y con el sangrado
    del título, y el secretario tenía que arreglar el formato a mano, que es
    justo el tiempo que este programa pretende ahorrarle.

    Medido en la plantilla de David hay TRES formatos distintos:
        encabezado  negrita, sangría 540385
        rótulo      negrita, SIN sangría
        cuerpo      sin negrita, sangría 457200
    """
    fuente = modelo if modelo is not None else p
    nuevo = copy.deepcopy(fuente._p)
    p._p.addnext(nuevo)
    q = Paragraph(nuevo, p._parent)
    escribir(q, texto)
    return q


# ═══════════════════════════════════════════════════════════════════════════
# Notas al pie DE VERDAD
# ═══════════════════════════════════════════════════════════════════════════
#
# Las citas a página NO van entre paréntesis en el cuerpo. David lo pidió así y
# es la convención del oficio: al pie, «Cfr. página 7, párrafo 3».
#
# `docx_generator_tcc.py` las finge con un superíndice y una lista al final del
# documento. Aquí se hacen de verdad, escribiendo en `word/footnotes.xml`, que
# la plantilla ya tiene —trae tres notas propias—, y copiando el formato de las
# suyas para que las nuevas no desentonen.

_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _parte_notas(doc):
    """La parte `footnotes.xml`, o None si la plantilla no tiene notas.

    OJO: python-docx la entrega como `Part` GENÉRICO, no como parte XML — no
    tiene `.element`, sólo `.blob` con los bytes. Hay que analizarla y volver a
    escribirla a mano.
    """
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/footnotes"):
            return rel.target_part
    return None


def _leer_notas(parte):
    from lxml import etree
    return etree.fromstring(parte.blob)


def _guardar_notas(parte, raiz) -> None:
    from lxml import etree
    parte._blob = etree.tostring(raiz, xml_declaration=True,
                                 encoding="UTF-8", standalone=True)


def _siguiente_id(raiz) -> int:
    ids = [int(n.get(qn("w:id"))) for n in raiz.findall(qn("w:footnote"))
           if n.get(qn("w:id")) is not None]
    return max(ids) + 1 if ids else 1


def _estilo_referencia(doc) -> str:
    """El styleId que la plantilla usa para la llamada de nota al pie."""
    try:
        for st in doc.styles:
            if (st.name or "").strip().lower() in ("footnote reference",
                                                   "ref. de nota al pie",
                                                   "referencia de nota al pie"):
                return st.style_id
    except Exception:
        pass
    return "FootnoteReference"


def anadir_nota(doc, parrafo, texto: str) -> bool:
    """Cuelga una nota al pie del final de `parrafo`. Devuelve si pudo.

    Se clona una nota EXISTENTE de la plantilla para heredar su estilo —tamaño,
    fuente, sangría— en vez de fabricar una desde cero, que saldría con el
    formato por defecto y cantaría al lado de las suyas.
    """
    parte = _parte_notas(doc)
    if parte is None:
        return False
    raiz = _leer_notas(parte)
    modelo = None
    for n in raiz.findall(qn("w:footnote")):
        tipo = n.get(qn("w:type"))
        if tipo in (None, "normal") and n.findall(qn("w:p")):
            modelo = n
            break
    if modelo is None:
        return False

    nid = _siguiente_id(raiz)
    nueva = copy.deepcopy(modelo)
    nueva.set(qn("w:id"), str(nid))
    # Se deja UN párrafo y se le escribe el texto MANIPULANDO EL XML, no con
    # `Paragraph`: la parte se analiza con lxml plano y sus elementos no son
    # los `CT_P` de python-docx, así que `.runs` no existe ahí.
    ps = nueva.findall(qn("w:p"))
    for extra in ps[1:]:
        nueva.remove(extra)
    textos = ps[0].findall(f".//{qn('w:t')}") if ps else []
    if textos:
        textos[0].text = texto
        textos[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for sobra in textos[1:]:
            sobra.text = ""
    raiz.append(nueva)
    _guardar_notas(parte, raiz)

    # La llamada en el cuerpo: un run con estilo de referencia Y con superíndice
    # explícito. El estilo por nombre —«Refdenotaalpie»— sólo se aplica si la
    # plantilla lo define; cuando no existe, Word ignora la referencia y el
    # número sale como texto corrido pegado a la palabra, que es lo que David vio
    # como «caracteres incorrectos en lugar del número». El superíndice directo
    # no depende de ningún estilo.
    r = parrafo.add_run()
    rpr = r._element.get_or_add_rPr()
    # El styleId REAL de la plantilla, leído de ella: «FootnoteReference». El
    # nombre que yo puse —«Refdenotaalpie»— es el rótulo español de Word y NO
    # existe como identificador ahí, así que Word lo ignoraba en silencio.
    estilo = OxmlElement("w:rStyle")
    estilo.set(qn("w:val"), _estilo_referencia(doc))
    rpr.append(estilo)
    va = OxmlElement("w:vertAlign")
    va.set(qn("w:val"), "superscript")
    rpr.append(va)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(nid))
    r._element.append(ref)
    return True


# ── El formato de la TRANSCRIPCIÓN de una tesis ──────────────────────────
#
# Medido sobre 1,676 transcripciones reales del corpus, y es DISTINTO del
# cuerpo en cuatro cosas a la vez:
#
#                    cuerpo (17,715 párr.)   transcripción (1,676)
#     sangría izq.   ninguna (91%)           709 twips (48%)  ← 1.25 cm
#     1ª línea       709 twips (55%)         ninguna (80%)
#     interlineado   1.5 (75%)               1.0 (48%)
#     tamaño         14 pt (76%)             13 pt (72%)
#     cursiva        no (97%)                SÍ (100%)
#
# Clonarla del cuerpo —que es lo que se hacía— la deja sin sangrar, a 14 puntos
# y con interlineado y medio: se lee como si fuera prosa del tribunal y no como
# lo que es, texto ajeno transcrito.
SANGRIA_CITA = 709          # twips = 1.25 cm

# La sangría de PRIMERA LÍNEA. El corpus la usa en el 63-70% de los párrafos
# (709 twips). Quitarla de TODO el documento fue una sobrecorrección mía: él se
# quejó de los RESÚMENES, no del estudio, y me lo dijo con todas las letras
# —«quitaste toda la sangría en el estudio de fondo, yo no pedí eso»—. Se
# conserva donde el corpus la pone y se hereda de la plantilla.
SANGRIA_PARRAFO = None      # None = no se toca, hereda del modelo
TAMANO_CITA = 13            # puntos


def _sangrar(p) -> None:
    """La sangría del cuerpo. Con SANGRIA_PARRAFO=None no se toca nada."""
    if SANGRIA_PARRAFO is None:
        return
    from docx.shared import Twips
    p.paragraph_format.first_line_indent = Twips(SANGRIA_PARRAFO)


def _aplicar_formato_cita(p) -> None:
    """Sangra, achica e inclina un párrafo ya escrito."""
    from docx.shared import Pt, Twips
    pf = p.paragraph_format
    pf.left_indent = Twips(SANGRIA_CITA)
    pf.first_line_indent = Twips(0)
    pf.line_spacing = 1.0
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(TAMANO_CITA)


def formato_cita(modelo_cuerpo):
    """El modelo del cuerpo, ajustado a las medidas de la transcripción."""
    if modelo_cuerpo is None:
        return None
    from docx.shared import Pt, Twips
    nuevo = copy.deepcopy(modelo_cuerpo._p)
    q = Paragraph(nuevo, modelo_cuerpo._parent)
    pf = q.paragraph_format
    pf.left_indent = Twips(SANGRIA_CITA)
    pf.first_line_indent = Twips(0)
    pf.line_spacing = 1.0
    for r in q.runs:
        r.italic = True
        r.font.size = Pt(TAMANO_CITA)
    return q


def modelos_de_formato(doc) -> dict:
    """Un párrafo de muestra por cada formato, tomado de la propia plantilla.

    HAY QUE LLAMARLA ANTES DE BORRAR NADA: los modelos son los párrafos que el
    secretario ya escribió, y al limpiar la sección para rellenarla se van con
    ella.
    """
    rotulo = cuerpo = None
    for p in doc.paragraphs:
        t = texto_de(p).strip()
        if not t or not p.runs:
            continue
        negrita = bool(p.runs[0].bold)
        palabras = len(t.split())
        # Rótulo: negrita, corto y sin numeral ordinal delante.
        if rotulo is None and negrita and 2 <= palabras <= 9 \
                and not re.match(r"^(PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|S[ÉE]PTIMO)\.", t):
            rotulo = p
        # Cuerpo: prosa larga sin negrita.
        if cuerpo is None and not negrita and palabras > 25:
            cuerpo = p
        if rotulo is not None and cuerpo is not None:
            break
    # El párrafo VACÍO también es un modelo, y es el que faltaba.
    #
    # David lo dijo mirando el documento: «sin espacios». No era una propiedad
    # de espaciado: en su plantilla hay UN PÁRRAFO VACÍO entre cada párrafo de
    # texto —la secuencia real es T · T · T ·— y yo los escribía pegados.
    vacio = next((p for p in doc.paragraphs if not texto_de(p).strip()), None)
    return {"rotulo": rotulo, "cuerpo": cuerpo, "vacio": vacio}


# Las marcas que deja el pipeline: [[p.7 §3]] o [[p.7]].
# El modelo no siempre ancla un párrafo suelto: cuando la idea abarca varios
# escribe «[[p.39 §2-3]]», y a veces encadena dos marcas para dos páginas. El
# patrón original sólo admitía «§3» y esas marcas se quedaban CRUDAS dentro de
# la sentencia — visibles, con corchetes, en el documento que se firma.
# Y no siempre es UNA cita: cuando la idea abarca dos páginas el modelo escribe
# «[[p.33 §1; p.34 §§1-2]]» —dos referencias dentro de la misma marca, con «§§»
# para el plural—. El patrón anterior exigía una sola y esas marcas se quedaban
# CRUDAS, con corchetes, justo donde debía ir el número de la nota al pie. Ahora
# se captura la marca ENTERA y se despieza aparte.
_MARCA_CITA = re.compile(r"\s*\[\[([^\[\]]{2,120})\]\]")
_UNA_CITA = re.compile(
    r"p{1,2}\.?\s*(\d{1,4})(?:\s*[-–]\s*\d{1,4})?"
    r"(?:\s*§+\s*(\d{1,3}(?:\s*[-–]\s*\d{1,3})?))?", re.I)


def _citas_de(marca: str) -> list[tuple[str, str]]:
    """Las (página, párrafo) que hay dentro de una marca, sean una o varias."""
    fuera = []
    for trozo in re.split(r"[;,]", marca):
        m = _UNA_CITA.search(trozo)
        if m:
            fuera.append((m.group(1), m.group(2) or ""))
    return fuera


def _texto_de_nota(pagina: str, parrafo: Optional[str]) -> str:
    if not parrafo:
        return f"Cfr. página {pagina}."
    p = re.sub(r"\s*[-–]\s*", " a ", parrafo)
    return (f"Cfr. página {pagina}, párrafos {p}." if " a " in p
            else f"Cfr. página {pagina}, párrafo {p}.")


def _normaliza_rubro(x: str) -> str:
    return re.sub(r"[^A-ZÁÉÍÓÚÑ0-9]+", " ", (x or "").upper()).strip()


def tesis_del_rubro(texto: str, tesis: list) -> Optional[dict]:
    """La tesis del acervo cuyo rubro se cita en este párrafo, si alguna.

    Se compara NORMALIZANDO —fuera comillas, acentos de puntuación y espacios—
    porque el modelo reproduce el rubro con comillas tipográficas y a veces
    corta la coletilla final.
    """
    m = _RX_RUBRO.search(texto or "")
    if not m:
        return None
    citado = _normaliza_rubro(m.group(1))
    if len(citado) < 25:
        return None
    for t in (tesis or []):
        real = _normaliza_rubro(t.get("rubro", ""))
        if not real:
            continue
        if real.startswith(citado[:70]) or citado.startswith(real[:70]):
            return t
    return None


# Una marca de cita por párrafo es lo normal; VEINTIDÓS seguidas al final de uno
# solo significa que el modelo devolvió el apartado entero sin saltos de línea y
# todas las llamadas se apilaron en el último punto. Se ve como
# «recurrida.³⁴⁵⁶⁷⁸⁹¹⁰…» y es lo primero que salta a la vista en el papel.
MAX_CITAS_POR_PARRAFO = 3


def _partir_si_es_un_bloque(textos: list[str]) -> list[str]:
    """Si un apartado llegó como UN párrafo con muchas citas, se parte por ellas.

    Cada marca cierra la idea que anota, así que cortar tras la marca deja
    párrafos que corresponden a una afirmación y su fuente — que es justamente
    como se lee un engrose.
    """
    fuera: list[str] = []
    for t in textos:
        if len(_MARCA_CITA.findall(t)) <= MAX_CITAS_POR_PARRAFO:
            fuera.append(t)
            continue
        trozos, resto = [], t
        while True:
            m = _MARCA_CITA.search(resto)
            if not m:
                break
            trozos.append(resto[:m.end()].strip())
            resto = resto[m.end():].lstrip()
        if resto.strip():
            trozos.append(resto.strip())
        fuera.extend(x for x in trozos if x)
    return fuera


def clonar_bloque(ancla, textos, modelo, modelo_vacio, doc=None, tesis=None,
                  modelo_cita=None):
    """Varios párrafos con su separador vacío detrás, como los escribe él.

    Y con las marcas de cita convertidas en NOTAS AL PIE. Si el documento no
    admite notas, la marca se retira del texto: nunca se deja un «[[p.7 §3]]»
    a la vista.
    """
    for t in _partir_si_es_un_bloque(list(textos)):
        citas = [c for m in _MARCA_CITA.finditer(t) for c in _citas_de(m.group(1))]
        limpio = _MARCA_CITA.sub("", t)
        ancla = clonar_tras(ancla, limpio, modelo)
        _sangrar(ancla)

        # El rubro va en NEGRITA dentro del párrafo, como él lo escribe.
        tramos = tramos_con_rubro(limpio)
        if tramos:
            escribir_tramos(ancla, tramos)

        if doc is not None:
            for pagina, parrafo in citas:
                anadir_nota(doc, ancla, _texto_de_nota(pagina, parrafo))

        # Y la tesis se completa desde el acervo: la LOCALIZACIÓN a pie de
        # página —«Semanario Judicial…, Novena Época, tomo XXXI, página 830»— y
        # el TEXTO ÍNTEGRO en párrafo aparte y en cursiva. Citar sólo el rubro
        # deja al lector sin poder comprobar qué dice la tesis.
        hallada = tesis_del_rubro(limpio, tesis) if tesis else None

        # ESTRUCTURA DE LA CITA, dictada por David con un ejemplo:
        #
        #     …la tesis orientadora de rubro y texto siguientes:
        #     «RUBRO.»                    ← párrafo aparte
        #     Texto de la tesis…          ← cursiva, sangrado
        #     ÓRGANO EMISOR.              ← y ahí la nota al pie
        #
        # El corpus mete el rubro DENTRO de la prosa 430 veces frente a 203 que
        # lo ponen aparte, así que la moda dice lo contrario. Pero cuando se
        # transcribe el texto ÍNTEGRO debajo, embeber el rubro en una frase que
        # sigue después deja la cita partida por la mitad. Su forma es la buena
        # y es quien firma.
        if hallada and (hallada.get("texto") or "").strip():
            m_r = _RX_RUBRO.search(limpio)
            if m_r:
                antes = limpio[:m_r.start()].rstrip(" ,;:")
                # Lo que el modelo escribió DESPUÉS del rubro —«…, establece que
                # la pensión debe…»— es prosa que continúa la frase. Se conserva
                # como párrafo propio detrás de la transcripción, en vez de
                # perderse: es argumento del estudio, no relleno de la cita.
                cola = limpio[m_r.end():].lstrip(" ,;:.")
                # Se rehace el anuncio y se tira la coleta que venía después del
                # rubro («, registro X, reconoce que…»): la dice el texto.
                # EL ANUNCIO SE LIMPIA ENTERO, no sólo «de rubro:». El prompt
                # le enseña al modelo a cerrar con «de rubro y texto
                # siguientes:» —así se cita aquí— y el ensamblador volvía a
                # añadirlo, dejando «…de rubro y texto siguientes: de rubro y
                # texto siguientes:». Lo vio David en el ADC 380/2025.
                anuncio = _RX_COLA_ANUNCIO.sub("", antes).rstrip(" ,;:")
                # SI NO HAY ANUNCIO, EL ANUNCIO YA ESTÁ ARRIBA. El modelo suele
                # escribir «…de registro 2018735, de rubro y texto siguientes:»
                # en un párrafo y el rubro en el SIGUIENTE. Al procesar el del
                # rubro, `antes` viene vacío y aquí se escribía la etiqueta a
                # secas, dejando «de rubro y texto siguientes:» dos veces
                # seguidas. Es el defecto que vio David, y no estaba en la
                # coleta sino en el párrafo sin anuncio.
                if anuncio:
                    escribir(ancla, f"{anuncio} de rubro y texto siguientes:")
                    _sangrar(ancla)
                    ancla.paragraph_format.keep_with_next = True
                    if modelo_vacio is not None:
                        ancla = clonar_tras(ancla, "", modelo_vacio)
                        ancla.paragraph_format.keep_with_next = True
                    ancla = clonar_tras(ancla, m_r.group(0), modelo_cita or modelo)
                escribir_tramos(ancla, [(m_r.group(0), {"bold": True, "italic": True})])
                _aplicar_formato_cita(ancla)
                for r_ in ancla.runs:
                    r_.bold = True
                ancla.paragraph_format.keep_with_next = True
                _cola_pendiente = cola if len(cola.split()) > 6 else ""

        if hallada:
            cuerpo = (hallada.get("texto") or "").strip()
            if cuerpo:
                # El rubro y su transcripción son UNA cita partida en dos
                # párrafos. Sin esto Word los separa en el salto de página y la
                # tesis aparece dos hojas después de lo que anuncia.
                ancla.paragraph_format.keep_with_next = True
                if modelo_vacio is not None:
                    ancla = clonar_tras(ancla, "", modelo_vacio)
                ancla.paragraph_format.keep_with_next = True   # el vacío también
                ancla = clonar_tras(ancla, cuerpo, modelo_cita or modelo)
                escribir_tramos(ancla, [(cuerpo, {"italic": True})])
                # El formato va DESPUÉS de escribir: `escribir_tramos` clona el
                # primer run como molde y heredaría el tamaño del cuerpo.
                _aplicar_formato_cita(ancla)

                # El ÓRGANO EMISOR cierra la cita, y de él cuelga la nota con la
                # localización. Así lo cierra él: «CUARTO TRIBUNAL COLEGIADO DEL
                # VIGÉSIMO SEGUNDO CIRCUITO.» y el pie detrás.
                organo = (hallada.get("instancia") or "").strip()
                if organo and organo.upper() not in cuerpo.upper()[-160:]:
                    if modelo_vacio is not None:
                        ancla = clonar_tras(ancla, "", modelo_vacio)
                    ancla = clonar_tras(ancla, organo.upper() + ".",
                                        modelo_cita or modelo)
                    escribir_tramos(ancla, [(organo.upper() + ".", {"italic": True})])
                    _aplicar_formato_cita(ancla)
                if doc is not None and hallada.get("localizacion"):
                    anadir_nota(doc, ancla, " " + hallada["localizacion"].strip())
                # Y detrás, la prosa que el modelo había puesto tras el rubro.
                if locals().get("_cola_pendiente"):
                    if modelo_vacio is not None:
                        ancla = clonar_tras(ancla, "", modelo_vacio)
                    ancla = clonar_tras(ancla, _cola_pendiente[0].upper()
                                        + _cola_pendiente[1:], modelo)
                    _sangrar(ancla)
                    _cola_pendiente = ""
            elif doc is not None and hallada.get("localizacion"):
                anadir_nota(doc, ancla, " " + hallada["localizacion"].strip())

        if modelo_vacio is not None:
            ancla = clonar_tras(ancla, "", modelo_vacio)
    return ancla


def buscar(doc, patron: str, desde: int = 0):
    """El primer párrafo, a partir del índice `desde`, que empieza por `patron`.

    EL `desde` NO ES UN LUJO. Los ordinales se repiten: hay un «SEXTO.
    Verificación de la sesión» en el RESULTANDO y un «SEXTO. Estudio» en el
    CONSIDERANDO, y el primero va ANTES que el «QUINTO. Antecedentes». Buscar
    «^SEXTO\.» a secas devuelve el equivocado y deja un rango invertido que no
    borra nada — pasó, y los antecedentes viejos se quedaron en el documento.
    """
    rx = re.compile(patron, re.I)
    for i, p in enumerate(doc.paragraphs):
        if i >= desde and rx.match(texto_de(p).strip()):
            return p
    return None


def indice_de(doc, p) -> int:
    for i, q in enumerate(doc.paragraphs):
        if q._p is p._p:
            return i
    return 0


def borrar_entre(doc, ini, fin) -> int:
    """Vacía los párrafos ENTRE dos marcas ya localizadas, sin tocarlas.

    Recibe los párrafos, no patrones: quien llama ya se aseguró de que `fin`
    va DESPUÉS de `ini`. Con patrones sueltos el rango puede salir invertido y
    entonces no borra nada en silencio, que es peor que fallar.
    """
    if ini is None or fin is None:
        return 0
    cuerpo = ini._p.getparent()
    hijos = list(cuerpo.iterchildren())
    try:
        a, b = hijos.index(ini._p), hijos.index(fin._p)
    except ValueError:
        return 0
    n = 0
    for el in hijos[a + 1:b]:
        if el.tag == qn("w:p"):
            cuerpo.remove(el)
            n += 1
    return n


# ═══════════════════════════════════════════════════════════════════════════
# Lo que se rellena
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class Relleno:
    """Todo lo que el ensamblador necesita saber del asunto."""
    # Cabecera
    encabezado: str = ""                 # «AMPARO DIRECTO ADMINISTRATIVO: 240/2026»
    quejoso: str = ""
    magistrado: str = ""
    secretario: str = ""

    # Fase 0 — el párrafo del cómputo, ya redactado
    oportunidad: str = ""

    # QUINTO. Antecedentes — de la sentencia reclamada
    antecedentes: list[str] = field(default_factory=list)

    # SEXTO. Estudio — los tres apartados, en este orden
    resumen_acto: list[str] = field(default_factory=list)
    resumen_conceptos: list[str] = field(default_factory=list)
    problemas: list[str] = field(default_factory=list)

    # Fase 6 — el estudio de fondo, ya con el criterio del secretario dentro.
    # Vacío mientras no lo haya dictado: entonces el documento es un ADELANTO,
    # con su hueco, y no una sentencia a medio resolver.
    estudio: list[str] = field(default_factory=list)

    # Las tesis del acervo (registro, rubro, texto, localizacion). El ensamblador
    # TRANSCRIBE de aquí, no de lo que escriba el modelo: la transcripción de una
    # tesis es dato verificado, y hacérsela redactar es invitarle a alterarla.
    # Las normas del acervo, para poner su texto al pie cuando se citan.
    normas: list = field(default_factory=list)
    tesis: list[dict] = field(default_factory=list)

    # Lo que decide el secretario. Mientras esté vacío se deja el hueco
    # marcado, igual que en el adelanto de papel.
    sentido: str = ""                    # la FÓRMULA literal, si ya se tiene
    calificaciones: list[str] = field(default_factory=list)   # o las del fondo
    # Datos del asunto que la plantilla trae del SUYO. Si no se pasan, no se
    # inventan: se dejan en hueco, que se ve, en vez de firmar los de otro.
    presentacion: str = ""               # «quince de abril de dos mil veinticinco»
    responsable: str = ""                # «Primera Sala Civil del Tribunal…»
    tema: str = ""

    es_recurso: bool = False             # revisión → «agravios», no «conceptos»
    numero_asunto: str = ""              # «512/2026», para el VISTO y el resolutivo


HUECO = "*********"

# El hueco se escribe con el número de asteriscos que a cada plantilla le tocó.
# Buscar una longitud fija deja restos a la vista en el resolutivo.
_RX_HUECO = re.compile(r"\*{3,}")

# La fórmula que el resolutivo ya trae escrita, para poder pisarla cuando el
# sentido del asunto es el contrario.
_RX_FORMULA_AMPARO = re.compile(
    r"no\s+ampara\s+ni\s+protege|ampara\s+y\s+protege", re.I)

# La coleta con que el modelo cierra el anuncio de una cita. Se le quita entera
# antes de poner la fórmula canónica, porque si no se escribe dos veces.
# Cubre lo que de verdad escribe: «de rubro y texto siguientes:», «de rubro y
# texto:», «cuyo rubro y texto son los siguientes:», «de rubro:», «de rubro y
# registro:», y la coma que a veces queda colgando delante.
_RX_COLA_ANUNCIO = re.compile(
    r"\s*,?\s*(?:cuy[oa]s?\s+|de\s+|del\s+)?"
    r"rubro(?:\s+y\s+(?:texto|registro|contenido))?"
    r"(?:\s+(?:es|son|se\s+transcribe[n]?|siguientes?|los\s+siguientes?|"
    r"el\s+siguiente|la\s+siguiente))*"
    r"\s*[:.]?\s*$", re.I)

# Lo que el ensamblador quiere decir sobre el documento que acaba de escribir.
avisos_ensamblado: list[str] = []

# ── La calificación de los conceptos NO ES el resolutivo ──────────────────
#
# «ineficaces» califica los conceptos de violación; el resolutivo dice si la
# Justicia de la Unión ampara o no. Meter la primera donde va el segundo produce
# «La Justicia de la Unión ineficaz a María Fernanda Ruiz», que no significa
# nada y que es exactamente la línea que se firma.
# En MINÚSCULAS y en negrita: así lo escribe él —«La Justicia de la Unión **no
# ampara ni protege** a…»—. Las versales son de otro tribunal, no del suyo.
_AMPARA = "ampara y protege"
_NO_AMPARA = "no ampara ni protege"


def formula_resolutivo(calificaciones: list[str]) -> tuple[str, str]:
    """(fórmula, aviso). Basta un concepto fundado para que se conceda.

    Con calificaciones mixtas la fórmula es la de concesión, pero el resolutivo
    real lleva además los EFECTOS, y esos no los escribe una tabla: se avisa
    para que los ponga quien firma.
    """
    cs = [str(c or "").strip().lower() for c in calificaciones if str(c or "").strip()]
    if not cs:
        return "", ""
    hay_fundado = any(c.startswith("fundad") for c in cs)
    if not hay_fundado:
        return _NO_AMPARA, ""
    if len(set(cs)) > 1:
        return _AMPARA, ("El resolutivo concede porque hay conceptos fundados, "
                         "pero la calificación es mixta: los EFECTOS de la "
                         "concesión los tienes que redactar tú.")
    return _AMPARA, ""


def _inicio_considerando(doc) -> int:
    """Dónde empieza el CONSIDERANDO. Los ordinales se repiten entre el
    RESULTANDO y el CONSIDERANDO, y sin este corte se toma el equivocado."""
    for i, q in enumerate(doc.paragraphs):
        t = re.sub(r"\s+", "", texto_de(q)).upper()
        if t.startswith("CONSIDERANDO"):
            return i
    return 0


# Un nombre tachado NO es un nombre. Las plantillas vienen anonimizadas y su
# quejoso es «*********»; los MISMOS asteriscos son los huecos que el
# secretario rellena. Tomarlos por un nombre y sustituirlos en todo el texto
# convertía cada hueco en el nombre del quejoso: en el ADC 380/2025 salió «de
# conformidad con el Parte Quejosa 6/2026, del Pleno del Órgano de Parte
# Quejosa». Con un nombre real ahí iría el de una persona, dentro de una frase
# que no habla de ella. Se firma tal cual.
_RX_TACHADO = re.compile(r"^[\s*_\-–—.·•xX]+$")


# La línea «QUEJOSO:» viene tachada, pero el RESOLUTIVO y el RESULTANDO no: la
# plantilla `amparo_directo` es el asunto de una persona real y la nombra
# cuatro veces. La anonimización rompió el mecanismo por los dos lados —tachó
# la referencia que se usaba para sustituir y dejó el nombre donde de verdad
# está—, así que el nombre de un litigante ajeno viajaba dentro de CADA
# proyecto que salía. Se busca donde de verdad aparece.
_RX_NOMBRE_RESOLUTIVO = re.compile(
    r"(?:no\s+)?ampara(?:\s+y|\s+ni)\s+protege\s+a\s+"
    r"([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ.\s]{5,70}?)\s*,", re.I)
_RX_NOMBRE_PROMOVIO = re.compile(
    r"promovid[oa]\s+por\s+([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ.\s]{5,70}?)\s*,")
# «…, Juan Carlos Quevedo Regadado, por propio derecho, promovió…»: aquí el
# nombre va ANTES del verbo, y con la otra grafía del apellido.
_RX_NOMBRE_DERECHO = re.compile(
    r"([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ.\s]{5,70}?)\s*,\s*por\s+(?:su\s+)?propio\s+derecho")
_RX_NOMBRE_ACTOR = re.compile(
    r"(?:del|el|la)\s+(?:actor|actora|demandad[oa]|apelante)\s+"
    r"([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ.\s]{5,70}?)\s*,")
_NO_ES_NOMBRE = ("la parte", "el quejoso", "la quejosa", "quien", "la suscrita",
                 "el suscrito", "la parte citada", "propio derecho")


def nombres_de_plantilla(ruta: str) -> list[str]:
    """Los nombres de personas reales que trae la plantilla dentro.

    Se leen del resolutivo, del resultando y de la síntesis, que es donde la
    anonimización no llegó. Devuelve TODAS las grafías: esta plantilla escribe
    el mismo apellido de dos maneras —«Regadado» y «Regalado»—, y sustituir
    sólo una deja la otra firmada.
    """
    d = docx.Document(ruta)
    fuera: list[str] = []
    for parte in _partes_con_texto(d):
        for p in _parrafos_todos(parte):
            t = texto_de(p)
            for rx in (_RX_NOMBRE_RESOLUTIVO, _RX_NOMBRE_PROMOVIO,
                       _RX_NOMBRE_DERECHO, _RX_NOMBRE_ACTOR):
                for m in rx.finditer(t):
                    nombre = " ".join(m.group(1).split())
                    if _RX_TACHADO.match(nombre):
                        continue
                    if any(x in nombre.lower() for x in _NO_ES_NOMBRE):
                        continue
                    if len(nombre.split()) < 2:      # un nombre lleva apellido
                        continue
                    if nombre not in fuera:
                        fuera.append(nombre)
    return fuera


def _quejoso_de_plantilla(ruta: str) -> str:
    """El quejoso que trae la plantilla, para poder sustituirlo en todo el texto.

    Si la línea «QUEJOSO:» viene tachada, se busca en el resolutivo.
    """
    d = docx.Document(ruta)
    for p in d.paragraphs:
        t = texto_de(p).strip()
        m = re.match(r"^QUEJOS[OA]\s*:\s*(.+?)\.?$", t, re.I)
        if m:
            nombre = m.group(1).strip()
            if not _RX_TACHADO.match(nombre):
                return nombre
    nombres = nombres_de_plantilla(ruta)
    return nombres[0] if nombres else ""


# ═══ EL NÚMERO QUE SE QUEDÓ DE LA PLANTILLA ════════════════════════════════
# Cada plantilla precargada ES UNA SENTENCIA REAL y arrastra su propio número
# de expediente: amparo_directo lleva 125/2026, queja lleva 143/2026. Aparece
# en los DOS encabezados y dos o tres veces más en el cuerpo —en la celda del
# rubro y en el resultando del registro y alta—. El ensamblador cambiaba la
# línea «AMPARO DIRECTO CIVIL: …» y dejaba todo lo demás, así que el proyecto
# del ADC 380/2025 salió encabezado «AMPARO DIRECTO CIVIL 125/2026». Lo vio
# David, y es la clase de error que invalida una sentencia entera.
#
# LA TRAMPA: en el cuerpo hay números con la misma forma que NO son
# expedientes —«2a./J. 58/2010», «P./J. 3/2013» son claves de tesis—. Por eso
# no se barre el patrón: se barre LA CADENA EXACTA que trae el encabezado, y
# aun así se comprueba que no vaya precedida de la marca de una clave.
_RX_CLAVE_TESIS = re.compile(r"(?:[JP]\.|/J\.|[0-9]a\.)\s*$")


def _partes_con_texto(doc):
    """Cuerpo, encabezados y pies. Lo que se olvida son los dos últimos."""
    yield doc
    for sec in doc.sections:
        for parte in (sec.header, sec.footer,
                      getattr(sec, "first_page_header", None),
                      getattr(sec, "first_page_footer", None),
                      getattr(sec, "even_page_header", None),
                      getattr(sec, "even_page_footer", None)):
            if parte is not None:
                yield parte


def _parrafos_todos(parte):
    for p in parte.paragraphs:
        yield p
    for t in parte.tables:
        for fila in t.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    yield p


def _expediente_de_plantilla(doc) -> str:
    """El número propio de la plantilla. La verdad está en el encabezado."""
    for sec in doc.sections:
        for parte in (sec.first_page_header, sec.header, sec.even_page_header):
            if parte is None:
                continue
            for p in _parrafos_todos(parte):
                m = re.search(r"\b\d{1,5}/20\d{2}\b", texto_de(p))
                if m:
                    return m.group(0)
    return ""


def _sustituir_en_parrafo(p, viejo: str, nuevo: str) -> int:
    """Cambia respetando el formato. Si el número viene partido entre runs,
    se junta en el primero —perder el reparto de runs no cambia cómo se ve—."""
    entero = texto_de(p)
    if viejo not in entero:
        return 0
    # No se toca una clave de tesis: «2a./J. 58/2010» no es un expediente.
    veces = 0
    for m in re.finditer(re.escape(viejo), entero):
        if not _RX_CLAVE_TESIS.search(entero[max(0, m.start() - 10):m.start()]):
            veces += 1
    if not veces:
        return 0
    for run in p.runs:                      # el caso normal: cabe en un run
        if viejo in run.text:
            run.text = run.text.replace(viejo, nuevo)
    if viejo not in texto_de(p):
        return veces
    if p.runs:                              # venía partido entre varios
        p.runs[0].text = entero.replace(viejo, nuevo)
        for run in p.runs[1:]:
            run.text = ""
    return veces


def _sanear_expediente(doc, numero: str) -> int:
    """Borra del documento el expediente que traía la plantilla."""
    viejo = _expediente_de_plantilla(doc)
    if not viejo or not numero or viejo == numero:
        return 0
    n = 0
    for parte in _partes_con_texto(doc):
        for p in _parrafos_todos(parte):
            n += _sustituir_en_parrafo(p, viejo, numero)
    if n:
        avisos_ensamblado.append(
            f"Se sustituyeron {n} apariciones del expediente {viejo} que traía "
            f"la plantilla por {numero}.")
    return n


def ensamblar(ruta_plantilla: str, r: Relleno, ruta_salida: str) -> str:
    """Rellena la plantilla y guarda el adelanto. Devuelve la ruta escrita."""
    avisos_ensamblado.clear()
    doc = docx.Document(ruta_plantilla)
    # ANTES QUE NADA: quitar el número de expediente de la plantilla. Si se hace
    # después, los párrafos nuevos ya escritos lo llevarían dentro.
    _num = re.search(r"\b\d{1,5}/20\d{2}\b", r.encabezado or "")
    if _num:
        _sanear_expediente(doc, _num.group(0))
    q = "agravios" if r.es_recurso else "conceptos de violación"
    # PRIMERO los modelos de formato: luego se borra el contenido viejo y con
    # él se irían los ejemplos de los que se copia el estilo.
    mod = modelos_de_formato(doc)

    # ── Cabecera ─────────────────────────────────────────────────────────
    # Aparece DOS veces: al inicio y en la página de síntesis. Se cambian las
    # dos, que es justo lo que se olvida al hacerlo a mano.
    # LOS DOS PUNTOS SON OBLIGATORIOS EN EL PATRÓN. Sin ellos, «^MAGISTRAD»
    # casa con «Magistrado Instructor adscrito a la Segunda Sección…» —que es
    # la AUTORIDAD RESPONSABLE, dentro del resultando— y la sobrescribe. Pasó.
    for patron, valor in (
        (r"^AMPARO\s+[\w\s]+:", r.encabezado),
        (r"^QUEJOS[OA]\s*:", f"QUEJOSO: {r.quejoso}."),
        (r"^MAGISTRADO(?:\s+PONENTE)?\s*:", f"MAGISTRADO: {r.magistrado}."),
        (r"^SECRETARIO(?:\s+DE\s+TRIBUNAL)?\s*:", f"SECRETARIO DE TRIBUNAL: {r.secretario}."),
    ):
        if not valor:
            continue
        rx = re.compile(patron, re.I)
        for p in doc.paragraphs:
            if rx.match(texto_de(p).strip()):
                escribir(p, valor)

    # ── TERCERO. Legitimación y oportunidad ──────────────────────────────
    if r.oportunidad:
        # NO basta con «el párrafo siguiente»: entre el de legitimación y el
        # cómputo puede haber un vacío, y entonces se escribía en él y el
        # cómputo del ASUNTO ANTERIOR sobrevivía debajo. En el papel salían dos
        # párrafos de oportunidad seguidos, con fechas distintas, en el mismo
        # documento que se firma. Se busca el que HABLA de oportunidad y los
        # demás se retiran.
        ini = buscar(doc, r"^TERCERO\.\s*Legitimaci")
        fin = buscar(doc, r"^CUARTO\.",
                     desde=indice_de(doc, ini) + 1) if ini is not None else None
        if ini is not None:
            i0 = indice_de(doc, ini) + 1
            i1 = indice_de(doc, fin) if fin is not None else len(doc.paragraphs)
            candidatos = [q for q in doc.paragraphs[i0:i1]
                          if re.search(r"(presentaci[óo]n de la demanda result|"
                                       r"oportun|surti[óo] efectos|plazo para la "
                                       r"promoci[óo]n)", texto_de(q), re.I)]
            if candidatos:
                escribir(candidatos[0], r.oportunidad)
                for sobra in candidatos[1:]:
                    sobra._p.getparent().remove(sobra._p)
            else:
                sig = ini._p.getnext()
                while sig is not None and sig.tag != qn("w:p"):
                    sig = sig.getnext()
                if sig is not None:
                    escribir(Paragraph(sig, ini._parent), r.oportunidad)

    # ── QUINTO. Antecedentes ─────────────────────────────────────────────
    if r.antecedentes:
        p = buscar(doc, r"^QUINTO\.\s*Antecedentes")
        fin = buscar(doc, r"^SEXTO\.\s*Estudio", desde=indice_de(doc, p) + 1) if p else None
        borrar_entre(doc, p, fin)
        if p is not None:
            ancla = clonar_bloque(p, r.antecedentes, mod["cuerpo"], mod["vacio"], doc)

    # ── SEXTO. Estudio ───────────────────────────────────────────────────
    # EL ORDINAL DEL ESTUDIO NO ES FIJO. En amparo directo es el SEXTO; en un
    # recurso de queja el considerando de estudio es el QUINTO, porque el
    # RESULTANDO tiene tres apartados y no seis. Buscar «^SEXTO. Estudio» a
    # secas deja el estudio fuera del documento en toda la familia de recursos.
    p = None
    for _ord in ("SEXTO", "QUINTO", "SÉPTIMO", "SEPTIMO", "CUARTO", "OCTAVO"):
        cand = buscar(doc, rf"^{_ord}\.\s*Estudio")
        if cand is not None and indice_de(doc, cand) > _inicio_considerando(doc):
            p = cand
            break
    if p is not None and (r.resumen_acto or r.resumen_conceptos or r.problemas
                          or r.estudio):
        fin = buscar(doc, r"^Por lo expuesto", desde=indice_de(doc, p) + 1)
        borrar_entre(doc, p, fin)
        _m_ord = re.match(r"^([A-ZÉ]+)\.", texto_de(p))
        escribir(p, f"{_m_ord.group(1) if _m_ord else 'SEXTO'}. Estudio de los {q}.")
        ancla = p
        bloques: list[tuple[str, list[str]]] = [
            ("Consideraciones relevantes de la sentencia recurrida"
             if r.estudio else
             "Consideraciones relevantes de la resolución reclamada.", r.resumen_acto),
            (q.capitalize() + ".", r.resumen_conceptos),
            # «Problema jurídico a resolver» y «Solución» son los rótulos del
            # ENGROSE; «Problemas jurídicos» era el del adelanto. Al llegar el
            # criterio, el documento deja de ser adelanto y toma los suyos.
            ("Problema jurídico a resolver" if r.estudio else "Problemas jurídicos.",
             r.problemas),
            # El estudio va SIN rótulo propio: abre con su encabezado ordinal y
            # la calificación —«Los conceptos son ineficaces»—, que es como
            # arranca el 40% de los engroses y lo que el lector busca primero.
            ("Solución", r.estudio),
        ]
        for rotulo, parrafos in bloques:
            if not parrafos:
                continue
            if rotulo:
                ancla = clonar_tras(ancla, rotulo, mod["rotulo"])
                if mod["vacio"] is not None:
                    ancla = clonar_tras(ancla, "", mod["vacio"])
            ancla = clonar_bloque(ancla, parrafos, mod["cuerpo"], mod["vacio"],
                                  doc, tesis=r.tesis)

    # ── El NOMBRE de la parte, allí donde la plantilla lo arrastra ───────
    #
    # El VISTO y el resolutivo traen el quejoso y el número del asunto ANTERIOR,
    # porque la plantilla es un adelanto real. Dejar ahí el nombre de otro es
    # más peligroso que dejar un hueco: un hueco se ve, un nombre equivocado
    # se firma. Se sustituye por el del asunto en curso.
    quejoso_viejo = _quejoso_de_plantilla(ruta_plantilla)

    # La plantilla puede traer el MISMO nombre escrito de dos maneras —en este
    # tribunal, «Larrañaga» en el proemio y «Larragaña» en el resolutivo—. Se
    # sustituye la forma canónica y la otra sobrevive intacta hasta la firma.
    # No se corrige automáticamente: un apellido no se arregla por parecido.
    if r.quejoso:
        _apellidos = [w for w in re.findall(r"[\wÁÉÍÓÚÑáéíóúñ]{5,}", r.quejoso)]
        for p_ in doc.paragraphs:
            t_ = texto_de(p_)
            for ap in _apellidos:
                for cand in re.findall(r"\b[\wÁÉÍÓÚÑáéíóúñ]{5,}\b", t_):
                    if cand.lower() == ap.lower():
                        continue
                    if (sorted(cand.lower()) == sorted(ap.lower())
                            and cand.lower() not in (x.lower() for x in _apellidos)):
                        avisos_ensamblado.append(
                            f"La plantilla escribe «{cand}» donde el quejoso es "
                            f"«{ap}». Compruébalo: se firma tal cual.")
                        break

    # TODAS las grafías, y en todas las partes. La plantilla escribe el mismo
    # apellido de dos maneras y en el encabezado, el cuerpo y la síntesis.
    if r.quejoso:
        _ajenos = nombres_de_plantilla(ruta_plantilla)
        if quejoso_viejo and quejoso_viejo not in _ajenos:
            _ajenos.insert(0, quejoso_viejo)
        _cambiados = 0
        for parte in _partes_con_texto(doc):
            for p in _parrafos_todos(parte):
                t = texto_de(p)
                nuevo = t
                for viejo in _ajenos:
                    if viejo.lower() in nuevo.lower():
                        nuevo = re.sub(re.escape(viejo), r.quejoso.title(),
                                       nuevo, flags=re.I)
                if nuevo != t:
                    escribir(p, nuevo)
                    _cambiados += 1
        if _cambiados:
            avisos_ensamblado.append(
                f"La plantilla nombraba a {', '.join(_ajenos)}: se sustituyó en "
                f"{_cambiados} párrafos por el quejoso de este asunto. "
                f"Compruébalo, que es un nombre y se firma.")
    if r.numero_asunto:
        # NO TODOS LOS NÚMEROS DEL RESOLUTIVO SON EL DEL AMPARO. La frase dice
        # «…contra la sentencia dictada el X, por la Sala Y, EN EL TOCA CIVIL
        # Z», y sustituir todas las cifras escribía el número del amparo en el
        # sitio del toca: en el ADC 380/2025 el resolutivo acabó diciendo «en el
        # toca civil 380/2025», que es falso y además verosímil. Se sustituye
        # sólo el número del juicio de amparo, no el que va tras «toca» o
        # «expediente», que pertenecen al asunto de origen.
        rx_num = re.compile(
            r"(?<!toca\s)(?<!toca\scivil\s)(?<!toca\sfamiliar\s)"
            r"(?<!expediente\s)\b\d{1,4}\s*/\s*\d{4}\b")

        def _solo_amparo(texto: str) -> str:
            fuera, i = [], 0
            for m in re.finditer(r"\b\d{1,4}\s*/\s*\d{4}\b", texto):
                previo = texto[max(0, m.start() - 34):m.start()].lower()
                fuera.append(texto[i:m.start()])
                if re.search(r"(?:toca|expediente|juicio de origen)"
                             r"(?:\s+\w+){0,2}\s*$", previo):
                    fuera.append(m.group(0))          # es del origen: intacto
                else:
                    fuera.append(r.numero_asunto)
                i = m.end()
            fuera.append(texto[i:])
            return "".join(fuera)

        for patron in (r"^V\s?I\s?S\s?T", r"^ÚNICO\.", r"^PRIMERO\.\s*(?:Se |La Justicia)"):
            q = buscar(doc, patron)
            if q is not None:
                escribir(q, _solo_amparo(texto_de(q)))

    # ── RESULTANDO PRIMERO y RESOLUTIVO: los datos del asunto ────────────
    # La plantilla los trae del SUYO y sólo se cambiaba el nombre, así que el
    # proyecto afirmaba que la demanda se presentó el «seis de febrero de dos
    # mil veintiséis» —fecha del asunto de Quevedo— bajo el nombre de la
    # quejosa de éste. Un dato falso y verosímil en el resultando es peor que
    # un hueco: el hueco se ve.
    if r.presentacion:
        q = buscar(doc, r"^PRIMERO\.\s*Presentaci")
        if q is not None:
            t_q = texto_de(q)
            nuevo_q, n_q = re.subn(
                r"(escrito\s+presentado\s+el\s+)([^,]{6,70})(,)",
                lambda m: m.group(1) + r.presentacion + m.group(3), t_q, count=1)
            if n_q and nuevo_q != t_q:
                escribir(q, nuevo_q)

    # ── El sentido y la síntesis: huecos si el secretario no ha decidido ──
    p = buscar(doc, r"^ÚNICO\.|^PRIMERO\.\s*(?:Se |La Justicia)")
    if p is not None:
        formula = r.sentido or formula_resolutivo(r.calificaciones)[0]
        # Sólo se toca el resolutivo de AMPARO. En un recurso la fórmula es otra
        # —se confirma, se revoca, se declara infundado— y depende de lo que se
        # recurrió: dejar el hueco es más honesto que rellenarlo por analogía.
        # EL RESOLUTIVO DE UN RECURSO NO LLEVA LA FÓRMULA DEL AMPARO. En queja
        # se escribe «Es fundado el recurso de queja» y el hueco quedaba a la
        # vista —«Es ********** el recurso»— porque sólo se rellenaba cuando el
        # párrafo decía «la Justicia de la Unión».
        t_res = texto_de(p)
        # La plantilla no siempre pone NUEVE asteriscos: en la de queja hay
        # diez, y `replace(HUECO, …)` dejaba uno suelto —«Es fundado* el
        # recurso»—. Se sustituye la RACHA, no una longitud concreta.
        if _RX_HUECO.search(t_res) and re.search(r"\brecurso\b", t_res, re.I):
            calif = (r.calificaciones or [""])[0].strip().lower()
            palabra = {"fundado": "fundado", "infundado": "infundado",
                       "inoperante": "inoperante", "ineficaz": "ineficaz"}.get(calif)
            if palabra:
                entero = _RX_HUECO.sub(palabra, t_res, count=1)
                m_o = re.match(r"^(ÚNICO\.|PRIMERO\.|SEGUNDO\.)", entero)
                tr = []
                resto = entero
                if m_o:
                    tr.append((m_o.group(1), {"bold": True}))
                    resto = entero[m_o.end():]
                i_p = resto.find(palabra)
                if i_p >= 0:
                    tr += [(resto[:i_p], {}), (palabra, {"bold": True}),
                           (resto[i_p + len(palabra):], {})]
                else:
                    tr.append((resto, {}))
                escribir_tramos(p, [x for x in tr if x[0]])
                avisos_ensamblado.append(
                    "Revisa el resolutivo del recurso: los EFECTOS —qué se "
                    "revoca y qué debe hacer el juez— los redactas tú.")

        if formula and "justicia de la unión" in texto_de(p).lower():
            # El resolutivo NO va entero en negrita. Medido en su engrose: sólo
            # «ÚNICO.» y la fórmula. `escribir()` metía todo en el primer run,
            # que es el del ordinal, y heredaba su negrita: el resolutivo entero
            # salía resaltado y había que deshacerlo a mano.
            # LA PLANTILLA NO TRAE UN HUECO AQUÍ: TRAE LA FÓRMULA DEL ASUNTO
            # ANTERIOR. `_RX_HUECO.sub` no encontraba nada que sustituir y la
            # negativa heredada sobrevivía intacta, pasara lo que pasara con el
            # estudio. Así salió el ADC 380/2025: cuatro efectos de concesión
            # redactados y un resolutivo que negaba el amparo. Es la
            # incongruencia que ningún proyecto puede llevar.
            entero = texto_de(p)
            if _RX_HUECO.search(entero):
                entero = _RX_HUECO.sub(formula, entero, count=1)
            else:
                entero, _n = _RX_FORMULA_AMPARO.subn(formula, entero, count=1)
                if _n:
                    avisos_ensamblado.append(
                        f"El resolutivo de la plantilla decía lo contrario y se "
                        f"corrigió a «{formula}». Comprueba que concuerda con el "
                        f"estudio antes de firmar.")
            m_ord = re.match(r"^(ÚNICO\.|PRIMERO\.|SEGUNDO\.)", entero)
            tramos = []
            resto = entero
            if m_ord:
                tramos.append((m_ord.group(1), {"bold": True}))
                resto = entero[m_ord.end():]
            i = resto.find(formula)
            if i >= 0:
                tramos += [(resto[:i], {}), (formula, {"bold": True}),
                           (resto[i + len(formula):], {})]
            else:
                tramos.append((resto, {}))
            escribir_tramos(p, [t for t in tramos if t[0]])

    # VA DESPUÉS DE LA FÓRMULA, no antes. Puesto delante, el hueco que se
    # abre aquí para la fecha del acto era el primero del párrafo y la
    # sustitución de la fórmula lo tomaba por suyo: el resolutivo salió
    # diciendo «contra la sentencia dictada el ampara y protege». Lo cazó el
    # verificador de congruencia recién escrito, que para eso está.
    # El resolutivo nombra la fecha del acto, la Sala y el toca. La Sala la
    # sabemos por el encargo; la fecha y el toca no viajan como campo, y ahí se
    # abre hueco antes que dejar los del asunto anterior.
    p_res = buscar(doc, r"^ÚNICO\.|^PRIMERO\.\s*(?:Se |La Justicia)")
    if p_res is not None and "justicia de la unión" in texto_de(p_res).lower():
        t_r = texto_de(p_res)
        original = t_r
        if r.responsable:
            t_r = re.sub(r"(,\s*por\s+la\s+)([^,]{10,120})(,\s*en\s+el\s+toca)",
                         lambda m: m.group(1) + r.responsable + m.group(3),
                         t_r, count=1)
        t_r, n_f = re.subn(r"(sentencia\s+dictada\s+el\s+)([^,]{6,70})(,)",
                           lambda m: m.group(1) + HUECO + m.group(3), t_r, count=1)
        t_r, n_t = re.subn(r"((?:toca|expediente)\s+(?:civil|familiar|penal)?\s*)"
                           r"\d{1,5}\s*/\s*\d{4}",
                           lambda m: m.group(1) + HUECO, t_r, count=1)
        if t_r != original:
            escribir(p_res, t_r)
            if n_f or n_t:
                avisos_ensamblado.append(
                    "En el resolutivo se abrieron huecos donde iban la fecha del "
                    "acto y el toca: la plantilla traía los del asunto anterior "
                    "y no viajan como dato del encargo. Rellénalos.")

            # El quejoso se sustituye, pero el APODERADO, el acto reclamado y la
            # autoridad siguen siendo los de la plantilla. Un hueco se ve; un
            # nombre de otro asunto, no, y aquí se firma.
            avisos_ensamblado.append(
                "Revisa el resolutivo: la plantilla arrastra el apoderado, la "
                "autoridad y el acto del asunto anterior.")
    p = buscar(doc, r"^TEMA\s*:")
    if p is not None and r.tema:
        escribir(p, f"TEMA: {r.tema}")

    # ÚLTIMA RED. Si alguna marca sobrevivió a todo lo anterior —una forma que
    # el patrón no previó— se retira antes de guardar. Un «[[p.33 §1]]» con
    # corchetes dentro de una sentencia firmada es el peor final posible, y más
    # vale perder la referencia que publicarla así.
    huerfanas = 0
    for p in doc.paragraphs:
        t = texto_de(p)
        if "[[" in t:
            huerfanas += 1
            escribir(p, re.sub(r"\s*\[\[[^\]]*\]\]", "", t))
    if huerfanas:
        avisos_ensamblado.append(
            f"{huerfanas} marcas de cita no se pudieron convertir en nota al pie "
            f"y se retiraron. Esas afirmaciones quedaron sin su referencia.")

    doc.save(ruta_salida)
    return ruta_salida


# Lo que el ensamblador quiere decir sobre el documento que acaba de escribir.
# Vive fuera porque `ensamblar()` devuelve una ruta y cambiar su firma obligaría
# a tocar el endpoint y el frontend por un aviso.



# ═══ EL DOCUMENTO SE LEE ANTES DE ENTREGARLO ═══════════════════════════════
# David, 30-ago-2026: «lee bien los documentos que saca el redactor sólo para
# verificar inconsistencias». Encontró dos —el expediente de la plantilla en el
# encabezado y el anuncio de cita repetido— y el barrido destapó tres más. Un
# residuo de plantilla no se ve leyendo por encima: se ve contándolo.
_RX_EXPEDIENTE = re.compile(r"\b\d{1,5}/(?:19|20)\d{2}\b")
# Lo que tiene forma de expediente y no lo es: claves de tesis y acuerdos
# generales del Pleno, que son cita legítima y no residuo.
_RX_NO_ES_EXPEDIENTE = re.compile(
    r"(?:[JP]\.|/J\.|[0-9]a\.|Acuerdo\s+General|acuerdo\s+general|"
    r"diverso|tesis|jurisprudencia)\s*$", re.I)


def _numeros_de(ruta: str) -> set:
    """Los expedientes que contiene un .docx, encabezados y tablas incluidos."""
    doc = docx.Document(ruta)
    out = set()
    for parte in _partes_con_texto(doc):
        for p in _parrafos_todos(parte):
            out |= set(_RX_EXPEDIENTE.findall(texto_de(p)))
    return out


def residuo_de_plantilla(ruta: str, numero: str = "",
                         ruta_plantilla: str = "") -> list[str]:
    """Lo que quedó del asunto ANTERIOR dentro del proyecto entregado.

    No decide por el secretario: se lo enseña. Un número de otro expediente en
    la carátula es exactamente el error que invalida una sentencia, y es
    invisible si nadie lo cuenta.
    """
    doc = docx.Document(ruta)
    partes = list(_partes_con_texto(doc))
    entero = []
    for parte in partes:
        for p in _parrafos_todos(parte):
            entero.append(texto_de(p))
    texto = "\n".join(entero)
    fuera: list[str] = []

    # 1. Números de expediente ajenos al asunto.
    propio = {numero.strip()} if numero else set()
    # LA SEÑAL PRECISA: un número que está en el proyecto Y EN LA PLANTILLA, y
    # que no es el de este asunto, viene de la plantilla. Los demás —el toca de
    # origen, el juicio natural— salieron de los documentos del caso y son
    # legítimos; señalarlos todos ahoga el aviso que importa.
    de_plantilla = _numeros_de(ruta_plantilla) if ruta_plantilla else None
    ajenos: dict[str, str] = {}
    for m in _RX_EXPEDIENTE.finditer(texto):
        num = m.group(0)
        if num in propio:
            continue
        if de_plantilla is not None and num not in de_plantilla:
            continue
        if _RX_NO_ES_EXPEDIENTE.search(texto[max(0, m.start() - 22):m.start()]):
            continue
        ajenos.setdefault(num, re.sub(r"\s+", " ",
                                      texto[max(0, m.start() - 60):m.end() + 30]))
    if ajenos:
        fuera.append("EXPEDIENTES DE LA PLANTILLA QUE SIGUEN EN EL PROYECTO: " + "; ".join(f"{k} («…{v}…»)"
                                                 for k, v in list(ajenos.items())[:8]))

    # 2. El anuncio de cita escrito dos veces.
    dobles = len(re.findall(r"(?:de\s+rubro\s+y\s+texto\s+siguientes\s*:\s*){2,}",
                            texto, re.I))
    if dobles:
        fuera.append(f"{dobles} citas anuncian «de rubro y texto siguientes:» "
                     f"DOS VECES seguidas.")

    # 3. PÁRRAFOS ENTEROS QUE SOBREVIVEN IDÉNTICOS. La señal más fuerte y la
    #    más general: si un párrafo largo del proyecto está palabra por palabra
    #    en la plantilla, no lo escribió este asunto. Así salió la página de
    #    síntesis del ADC 380/2025, que seguía contando una disolución de
    #    copropiedad ajena bajo el rótulo «Antecedentes».
    if ruta_plantilla:
        try:
            plan = docx.Document(ruta_plantilla)
        except Exception:
            plan = None
        if plan is not None:
            suyos = set()
            for parte in _partes_con_texto(plan):
                for p in _parrafos_todos(parte):
                    t = " ".join(texto_de(p).split())
                    if len(t.split()) >= 25:
                        suyos.add(t)
            # NO TODO LO CALCADO ES DEFECTO. «Querétaro, Querétaro. Resolución
            # del Tercer Tribunal Colegiado…» abre TODAS las sentencias de este
            # tribunal y debe repetirse. Lo que no puede repetirse es un
            # párrafo con datos del otro asunto dentro: un nombre propio, un
            # número de expediente o una fecha escrita en letra.
            # El nombre propio NO sirve de marca: «Tercer Tribunal Colegiado»
            # son tres mayúsculas seguidas y aparece en la fórmula de apertura.
            # Lo que sí distingue un asunto de otro es su expediente y su fecha
            # escrita en letra, que es como las escribe una sentencia.
            _rx_dato = re.compile(
                r"\b\d{1,5}/(?:19|20)\d{2}\b"
                r"|\bde\s+dos\s+mil\s+\w+"
                r"|\bexpediente\s+\d"
                r"|\btoca\s+(?:civil|familiar|penal)?\s*\d", re.I)
            calcados = []
            for parte in partes:
                for p in _parrafos_todos(parte):
                    t = " ".join(texto_de(p).split())
                    if len(t.split()) >= 25 and t in suyos and _rx_dato.search(t):
                        calcados.append(t)
            if calcados:
                fuera.append(
                    f"{len(calcados)} párrafos siguen IGUALES a los de la "
                    f"plantilla, palabra por palabra: no los escribió este "
                    f"asunto. El primero empieza «{calcados[0][:150]}…»")
    return fuera


# ═══ LA INCONGRUENCIA QUE NO PUEDE PASAR ═══════════════════════════════════
# David, 30-ago-2026, sobre el ADC 380/2025: «El proyecto contiene
# consideraciones para conceder el amparo, precisa efectos y,
# contradictoriamente, termina negando el amparo. La incongruencia no debe
# pasar por ningún proyecto».
#
# Es el único defecto de esta lista que invalida por sí solo: un resolutivo que
# no dice lo que dice el estudio no es un proyecto mejorable, es un proyecto
# que no se puede firmar. Se comprueba SIEMPRE y contra el documento final, no
# contra lo que el pipeline creía estar escribiendo.
_RX_CONCEDE = re.compile(r"(?<!no\s)ampara\s+y\s+protege", re.I)
_RX_NIEGA = re.compile(r"no\s+ampara\s+ni\s+protege", re.I)
_RX_EFECTOS = re.compile(
    r"(?:como\s+(?:primer|segundo|tercer|cuarto|quinto)\s+efecto"
    r"|efectos?\s+de\s+la\s+concesi[óo]n"
    r"|deber[áa]\s+dejar\s+insubsistente"
    r"|dejar[áa]\s+insubsistente)", re.I)


_ORDINALES_CONCEPTO = ("primer", "segundo", "tercer", "cuarto", "quinto",
                       "sexto", "séptimo", "octavo", "noveno")


def _cuenta_conceptos(texto: str) -> int:
    """El ordinal más alto que se nombra como concepto o agravio."""
    alto = 0
    for i, o in enumerate(_ORDINALES_CONCEPTO, 1):
        if re.search(rf"\b{o}\s+(?:concepto|agravio)", texto, re.I):
            alto = i
    return alto


def revisar_congruencia(ruta: str, calificaciones=None,
                        tipo_asunto: str = "") -> list[str]:
    """El resolutivo tiene que decir lo que dice el estudio. Sin excepción."""
    doc = docx.Document(ruta)
    texto = "\n".join(texto_de(p) for parte in _partes_con_texto(doc)
                      for p in _parrafos_todos(parte))
    fuera: list[str] = []

    # ═══════════════════════════════════════════════════════════════════════
    # LA FÓRMULA DE OTRO TIPO DE ASUNTO
    # ═══════════════════════════════════════════════════════════════════════
    # POR QUÉ ESTO LLEGÓ A UN SECRETARIO SIN UN SOLO AVISO. Esta función corre
    # SIEMPRE sobre el .docx terminado y es la última red antes de la firma,
    # pero sólo conocía dos fórmulas: «ampara y protege» y «no ampara ni
    # protege». Una queja que cerraba diciendo «lo procedente es negar el
    # amparo solicitado» y decretaba «Es infundado el recurso de queja» pasaba
    # limpia: las dos frases son congruentes ENTRE SÍ para quien sólo sabe de
    # amparos, y ninguna de las dos activaba el par que se vigilaba.
    #
    # Se comprueba primero, porque una fórmula ajena no es un matiz de estilo:
    # es una resolución que no existe en ese recurso.
    if tipo_asunto:
        try:
            import tipos_asunto as _ta_cg
            _aj = _ta_cg.cierre_ajeno(tipo_asunto, texto)
        except Exception:
            _aj = []
        # LA LEY QUE NO GOBIERNA EL RECURSO, en el considerando que funda la
        # jurisdicción. Acotado a ese párrafo: en el resto del documento la Ley
        # de Amparo puede aparecer legítimamente —el cómputo cita su artículo
        # 19 para los inhábiles— y una alarma que salta siempre no se lee.
        try:
            _leyes = _ta_cg.prohibido_en_competencia(tipo_asunto, texto)
        except Exception:
            _leyes = []
        if _leyes:
            fuera.append(
                "LEY AJENA EN EL CONSIDERANDO DE COMPETENCIA: la jurisdicción "
                "se funda en una ley que no gobierna este recurso. Es lo que "
                "pasa cuando se cuela la plantilla de otro tipo de asunto, y "
                "está en el primer considerando que se lee.")
        if _aj:
            _n = _ta_cg.vocabulario_de(tipo_asunto)["nombre"]
            fuera.append(
                f"FÓRMULA DE OTRO TIPO DE ASUNTO EN {_n.upper()}: aparece "
                f"{', '.join('«' + a.replace(chr(92) + 'b', '') + '»' for a in _aj)}. "
                f"En este recurso no hay amparo que negar ni conceder: se coló "
                f"la plantilla del amparo directo y el proyecto no se sostiene.")

    concede = bool(_RX_CONCEDE.search(texto))
    niega = bool(_RX_NIEGA.search(texto))
    efectos = bool(_RX_EFECTOS.search(texto))

    if concede and niega:
        fuera.append("EL DOCUMENTO CONCEDE Y NIEGA A LA VEZ: aparecen «ampara y "
                     "protege» y «no ampara ni protege» en el mismo proyecto.")

    if niega and not concede and efectos:
        fuera.append("INCONGRUENCIA QUE INVALIDA: el estudio fija EFECTOS de la "
                     "concesión —«deberá dejar insubsistente…»— y el resolutivo "
                     "NIEGA el amparo. Un resolutivo que no dice lo que dice el "
                     "estudio no se puede firmar.")

    # EL ESTUDIO NO PUEDE ANALIZAR MÁS CONCEPTOS DE LOS QUE HAY. Medido en el
    # ADL 382/2024: el resumen sintetizaba CINCO y la solución analizaba SEIS,
    # atribuyendo al sexto lo que el quinto ya decía. Un concepto que nadie
    # planteó y que la sentencia contesta es una incongruencia por exceso.
    _i = texto.find("Conceptos de violación")
    if _i < 0:
        _i = texto.find("Agravios")
    _j = texto.find("Solución")
    if 0 <= _i < _j:
        _en_resumen = _cuenta_conceptos(texto[_i:_j])
        _en_estudio = _cuenta_conceptos(texto[_j:])
        if _en_resumen and _en_estudio > _en_resumen:
            fuera.append(
                f"EL ESTUDIO ANALIZA {_en_estudio} PLANTEAMIENTOS Y EL RESUMEN "
                f"SINTETIZA {_en_resumen}. Contestar uno que nadie planteó es "
                f"incongruencia por exceso: comprueba si el «{_ORDINALES_CONCEPTO[_en_estudio-1]}» "
                f"existe o si repite lo que ya dijo otro.")

    cs = [str(c or "").strip().lower() for c in (calificaciones or []) if str(c or "").strip()]
    if cs:
        hay_fundado = any(c.startswith("fundad") for c in cs)
        if hay_fundado and niega and not concede:
            fuera.append(f"El secretario calificó «{cs[0]}» y el resolutivo "
                         f"NIEGA el amparo. Basta un concepto fundado para "
                         f"conceder.")
        if not hay_fundado and concede and not niega:
            fuera.append(f"Ningún concepto se calificó de fundado "
                         f"({', '.join(sorted(set(cs)))}) y el resolutivo "
                         f"CONCEDE el amparo.")
    return fuera


def huecos_pendientes(ruta: str) -> list[str]:
    """Los párrafos que siguen sin rellenar. Nada sale con un `*****` dentro."""
    doc = docx.Document(ruta)
    fuera = []
    for p in doc.paragraphs:
        t = texto_de(p).strip()
        if HUECO in t or re.search(r"\*{4,}", t):
            fuera.append(t[:120])
        if re.match(r"^(TEMA|OPORTUNIDAD)\s*:\s*$", t):
            fuera.append(t)
    return fuera
